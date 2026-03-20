# saisie_a_faire_pdfdiff_extractor.py

#region saisie_a_faire_pdfdiff_extractor.py Template+Completed SAISIE PDFs â†’ CSV + generate k138_values.csv + optional K138 autofill hook
# Version 0.4.0 (2026/02/05)
# Add OCR fallback for image-based PDFs; fix extraction placement and duplicates; interim CSVs only in temp.

import re
import csv
import sys
import subprocess
import configparser
import tempfile
import importlib.util
import os
import io
import zipfile
import html
import hashlib
import json
import shutil
from datetime import datetime
from dataclasses import dataclass
from pathlib import Path
from typing import List, Tuple, Dict, Optional

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog

import fitz  # PyMuPDF

# ---- OCR (optional, for image-based PDFs) ----
try:
    from PIL import Image, ImageTk
    HAVE_PIL = True
except ImportError:
    HAVE_PIL = False
    Image = None
    ImageTk = None
try:
    import pytesseract
    # Try to find Tesseract - check common locations and PATH
    TESSERACT_EXE = os.environ.get("TESSERACT_CMD", "")
    if not TESSERACT_EXE:
        # Check common Windows installation path
        common_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]
        for path in common_paths:
            if os.path.exists(path):
                TESSERACT_EXE = path
                break
    if TESSERACT_EXE:
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_EXE
    # Verify Tesseract is actually accessible
    try:
        pytesseract.get_tesseract_version()
        HAVE_TESSERACT = True
    except Exception:
        HAVE_TESSERACT = False
except Exception:
    HAVE_TESSERACT = False
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    HAVE_DND = True
except Exception:
    DND_FILES = None
    TkinterDnD = None
    HAVE_DND = False

try:
    from barcode import Code128
    from barcode.writer import ImageWriter
    HAVE_BARCODE = True
except Exception:
    HAVE_BARCODE = False

# ======================== =
# Configuration
# ======================== =
# templates_dir: Set via GUI/config; contains static templates and reference .txt files.
# temp_dir: Where interim CSVs/temp data are stored (never in working case or templates).
# Output K138 PDF always goes to working_dir = parent of input Saisie PDF.

CONFIG_FILE = Path("radiance_copilot.cfg")
APP_INSTANCE_LOCK = Path(tempfile.gettempdir()) / "radiance_copilot.instance.lock"
# User-facing output preference: hide CSV mentions from Instruction/Feedback panel.
HIDE_CSV_IN_FEEDBACK = True
CONCISE_PROGRESS_LOGS = True
_NOISY_PROGRESS_SUBSTRINGS = (
    "active case folder:",
    "input file:",
    "source file already in case folder",
    "copied source file into case folder",
    "could not copy source file into case folder",
    "configurations folder:",
    "case storage root:",
    "case source file:",
    "searching for templates",
    "[ok] found saisie template:",
    "[ok] input is image",
    "[ok] input is word",
    "[ok] loaded notice text from file",
    "[warn] no notice text file found",
    "processing:",
    "output folder:",
    "[ok] extraction mode:",
    "internal case folder:",
    "[info] saved to hidden folder:",
    "[info] updated:",
    "============================================================",
)


def _pid_is_running(pid: int) -> bool:
    """Return True if a process with PID appears alive."""
    if pid <= 0:
        return False
    if os.name != "nt":
        try:
            os.kill(pid, 0)
            return True
        except Exception:
            return False
    try:
        flags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
        out = subprocess.run(
            ["tasklist", "/FI", f"PID eq {pid}", "/FO", "CSV", "/NH"],
            capture_output=True,
            text=True,
            timeout=2,
            creationflags=flags,
        )
        txt = (out.stdout or "").strip()
        if (not txt) or ("No tasks are running" in txt):
            return False
        return str(pid) in txt
    except Exception:
        return False


def acquire_single_instance_lock() -> Optional[Path]:
    """
    Prevent multiple concurrent Radiance UI instances.
    Returns lock path when acquired, or None if another live instance exists.
    """
    lock = APP_INSTANCE_LOCK
    try:
        if lock.exists():
            prev_txt = (lock.read_text(encoding="utf-8", errors="ignore") or "").strip()
            prev_pid = int(prev_txt) if prev_txt.isdigit() else 0
            if prev_pid and _pid_is_running(prev_pid):
                return None
            try:
                lock.unlink()
            except Exception:
                pass
        lock.write_text(str(os.getpid()), encoding="utf-8")
        return lock
    except Exception:
        # Fail open to avoid blocking app launch due lock IO issues.
        return lock


def release_single_instance_lock(lock: Optional[Path]) -> None:
    if not lock:
        return
    try:
        if not lock.exists():
            return
        cur_txt = (lock.read_text(encoding="utf-8", errors="ignore") or "").strip()
        if cur_txt == str(os.getpid()):
            lock.unlink()
    except Exception:
        pass


def show_single_instance_warning() -> None:
    """Show non-fatal warning when another app instance is already running."""
    try:
        probe = tk.Tk()
        probe.withdraw()
        messagebox.showwarning(
            "Radiance Copilot",
            "Radiance is already running.\n\n"
            "Close the existing window first, then launch again.",
            parent=probe,
        )
        probe.destroy()
    except Exception:
        pass
REQUIRED_HELPER_MODULES = {
    "K138": "fill_k138_notice",
    "Saisie d'interet": "fill_saisie_interet",
}


def is_helper_module_available(module_name: str) -> bool:
    """Return True when helper module can be imported in current runtime."""
    try:
        return importlib.util.find_spec(module_name) is not None
    except Exception:
        return False


def helper_module_hint(module_name: str) -> str:
    """
    Return a short hint showing where the helper script is expected in dev mode.
    In bundled mode, helper may be inside executable (no .py file on disk).
    """
    script_name = f"{module_name}.py"
    app_candidate = Path(_app_dir()) / script_name
    cwd_candidate = Path.cwd() / script_name
    if app_candidate.resolve() == cwd_candidate.resolve():
        return str(app_candidate)
    return f"{app_candidate} (or bundled in executable)"


def _app_dir() -> str:
    """Return directory containing the app script/executable."""
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def resource_path(rel: str) -> str:
    """
    Resolve bundled resource path for both dev mode and PyInstaller onefile mode.
    """
    base = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
    return str(base / rel)


def resolve_asset_path(*relative_candidates: str) -> str:
    """Return first existing asset path from candidate relative paths."""
    for rel in relative_candidates:
        if not rel:
            continue
        p = Path(resource_path(rel))
        if p.exists():
            return str(p)
    return ""


def get_temp_dir() -> Path:
    """Return temp directory for interim files. Prefer AppData on Windows."""
    if sys.platform == "win32":
        appdata = os.environ.get("LOCALAPPDATA", "")
        if appdata:
            tmp = Path(appdata) / "RadianceCopilot" / "tmp"
            tmp.mkdir(parents=True, exist_ok=True)
            return tmp
    return Path(tempfile.gettempdir())


# Fixed location for K138 form â€“ always printed at "at" / Ã  (Dmitry: ciblage/lieu_interception is extracted for analysis only)
K138_SEIZURE_LOCATION_FIXED = "MONTREAL POSTAL FACILITY, ETC LEO-BLANCHETTE / POSTAL CUSTOMS"


def get_hidden_data_dir(project_root: Path) -> Path:
    """
    Return hidden folder for technical/runtime data (not client-visible outputs).
    Preferred global location:
      - Windows: %LOCALAPPDATA%/Radiance/.extracted_data
      - Other OS: ~/Radiance/.extracted_data
    """
    _ = project_root  # kept for backward-compatible signature
    if sys.platform == "win32":
        appdata = os.environ.get("LOCALAPPDATA", "")
        if appdata:
            hidden = Path(appdata) / "Radiance" / ".extracted_data"
        else:
            hidden = Path.home() / "Radiance" / ".extracted_data"
    else:
        hidden = Path.home() / "Radiance" / ".extracted_data"

    hidden.mkdir(parents=True, exist_ok=True)

    # Make technical folder hidden on Windows Explorer.
    if sys.platform == "win32":
        try:
            subprocess.run(
                ["attrib", "+h", str(hidden)],
                check=False,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
            )
        except Exception:
            pass
    return hidden


def _is_transient_upload_path(path: Path) -> bool:
    """Return True when path points to temporary upload/runtime folders."""
    try:
        txt = str(path.resolve()).replace("\\", "/").lower()
    except Exception:
        txt = str(path).replace("\\", "/").lower()
    markers = [
        "/radiance_web_uploads/",
        "/radiancecopilot/tmp/",
        "/localcache/local/radiancecopilot/tmp/",
        "/appdata/local/temp/",
        "/tmp/",
    ]
    return any(m in txt for m in markers)


def _preferred_case_output_dir() -> Optional[Path]:
    """
    Pick stable output directory for cases when source file comes from temp uploads.
    Priority:
      1) env RADIANCE_OUTPUT_DIR
      2) config paths.case_output_folder
      3) config paths.saisie_folder
      4) ./working_case_example
    """
    env_dir = clean_value(os.environ.get("RADIANCE_OUTPUT_DIR", ""))
    candidates: List[Path] = []
    if env_dir:
        candidates.append(Path(env_dir))
    cfg_case = get_config_path("paths", "case_output_folder")
    if cfg_case:
        candidates.append(cfg_case)
    cfg_saisie = get_config_path("paths", "saisie_folder")
    if cfg_saisie:
        candidates.append(cfg_saisie)
    default_working = Path("working_case_example")
    if default_working.exists():
        candidates.append(default_working)

    for cand in candidates:
        try:
            rc = cand.resolve()
            if rc.exists() and rc.is_dir() and (not _is_transient_upload_path(rc)):
                return rc
        except Exception:
            continue
    return None


def _resolve_working_directory_for_source(selected_file: Path) -> Path:
    """Resolve output folder from source path, redirecting transient upload paths when possible."""
    src_parent = selected_file.resolve().parent
    if not _is_transient_upload_path(src_parent):
        return src_parent
    preferred = _preferred_case_output_dir()
    if preferred:
        return preferred
    return src_parent


def detect_working_directory(selected_file: Path) -> Tuple[Path, str]:
    """
    Resolve active case folder from selected SAISIE file.
    Case folder name is based on the file stem (one folder per selected SAISIE file).
    """
    wd = _resolve_working_directory_for_source(selected_file)
    case_name = selected_file.stem
    print(f"Active case folder: {wd}")
    return wd, case_name


def _dot_hidden_folder_name(name: str) -> str:
    """
    Build a safe hidden folder name from working directory name.
    Example: "12345 2025-12-11 ABXXX..." -> ".12345 2025-12-11 ABXXX..."
    """
    raw = str(name or "").strip().lstrip(".")
    raw = re.sub(r'[<>:"/\\|?*]+', "_", raw)
    raw = re.sub(r"\s+", " ", raw).strip(" .")
    if not raw:
        raw = "case"
    # Keep path lengths reasonable on Windows.
    if len(raw) > 120:
        raw = raw[:120].rstrip(" .")
    return f".{raw}"


def _folder_case_type_label(form_type: str) -> str:
    t = clean_value(form_type or "").lower()
    if "cannabis" in t:
        return "Cannabis"
    if "knife" in t or "arm" in t:
        return "Knives"
    if "other" in t or "stupefiant" in t:
        return "Other"
    return clean_value(form_type or "") or "Other"


def _soi_label_from_text(raw: str) -> str:
    t = clean_value(raw or "").upper()
    if not t:
        return "No SoI"
    if re.search(r"(?:\bX\b|âœ“|âœ”|\bOUI\b|\bYES\b|\bTRUE\b|\b1\b)", t):
        return "SoI (Seizure of Importance)"
    return "No SoI"


def build_internal_case_meta(top: Optional[Dict[str, str]], values: Optional[Dict[str, str]]) -> Dict[str, str]:
    """
    Build metadata used to name hidden internal case folder:
    '<badge> <date> <inventory> - <type> - <SoI/No SoI>'
    """
    top = top or {}
    values = values or {}

    agent = re.sub(r"\D", "", values.get("seizing_officer", "") or "")
    if agent:
        agent = agent[-5:] if len(agent) > 5 else agent.zfill(5)

    dt = parse_first_date(clean_value(top.get("DATE / HEURE INTERCEPTION:", "")))
    if dt:
        y, m, d = dt
        case_date = f"{y:04d}-{m:02d}-{d:02d}"
    else:
        nd = clean_value(values.get("notice_date", ""))
        case_date = nd if re.fullmatch(r"\d{4}-\d{2}-\d{2}", nd or "") else ""

    inv_raw = clean_value(values.get("description_inventory", "") or "")
    inventory = _normalize_inventory_number(inv_raw)
    if not inventory:
        fallback_inv = re.sub(r"[^A-Za-z0-9-]+", "", (inv_raw or "").upper())
        if len(fallback_inv) >= 6:
            inventory = fallback_inv
    case_type = _folder_case_type_label(values.get("form_type", "") or "")
    soi_raw = _top_first_match(top, ["ENVERGURE", "SOI"])
    soi = _soi_label_from_text(soi_raw)

    return {
        "agent_id": agent,
        "case_date": case_date,
        "inventory_number": inventory,
        "case_type": case_type,
        "soi": soi,
    }


def _build_internal_folder_title(working_dir: Path, case_meta: Optional[Dict[str, str]]) -> str:
    """
    Preferred title:
      '<badge> <date> <inventory> - <type> - <SoI/No SoI>'
    Fallback:
      '<working_dir_name>'
    """
    case_meta = case_meta or {}
    agent = clean_value(case_meta.get("agent_id", ""))
    case_date = clean_value(case_meta.get("case_date", ""))
    inv_raw = clean_value(case_meta.get("inventory_number", "") or "")
    inventory = _normalize_inventory_number(inv_raw)
    if not inventory:
        fallback_inv = re.sub(r"[^A-Za-z0-9-]+", "", (inv_raw or "").upper())
        if len(fallback_inv) >= 6:
            inventory = fallback_inv
    case_type = clean_value(case_meta.get("case_type", ""))
    soi = clean_value(case_meta.get("soi", ""))

    head_parts = [x for x in (agent, case_date, inventory) if x]
    if head_parts:
        head = " ".join(head_parts)
        tail = []
        if case_type:
            tail.append(case_type)
        if soi:
            tail.append(soi)
        if tail:
            return f"{head} - " + " - ".join(tail)
        return head
    return working_dir.resolve().name


def _read_case_key_marker(case_dir: Path) -> str:
    try:
        marker = case_dir / ".case_key"
        if marker.exists():
            return marker.read_text(encoding="utf-8").strip()
    except Exception:
        pass
    return ""


def _write_case_key_marker(case_dir: Path, case_key: str) -> None:
    try:
        (case_dir / ".case_key").write_text(case_key, encoding="utf-8")
    except Exception:
        pass


def _find_internal_case_root_by_key(hidden_root: Path, case_key: str) -> Optional[Path]:
    try:
        for p in hidden_root.iterdir():
            if not p.is_dir():
                continue
            if _read_case_key_marker(p) == case_key:
                return p
    except Exception:
        pass
    return None


def ensure_case_structure(
    working_dir: Path,
    selected_file: Path,
    case_meta: Optional[Dict[str, str]] = None,
) -> Dict[str, Path]:
    """
    Ensure per-case folder structure exists.

    Client-visible structure (flat):
      <working_dir>/
          K138.pdf
          Agenda_<file_stem>.pdf

    Internal hidden structure (developer-only):
      <global_hidden_root>/.<case_name>/
          values_latest.json
          k138/
              k138_latest.pdf
          agenda/
              agenda_latest.pdf
              agenda_latest.docx
              barcode.png
    """
    selected_file = selected_file.resolve()
    working_dir = working_dir.resolve()

    # Client-visible outputs must stay in the working directory root (flat layout).
    case_root = working_dir
    case_root.mkdir(parents=True, exist_ok=True)
    legacy_nested_case_root = working_dir / selected_file.stem

    # Client-visible outputs (flat in working directory).
    client_k138_pdf = case_root / "K138.pdf"
    legacy_client_k138_pdf = case_root / f"K138_{selected_file.stem}.pdf"
    client_agenda_pdf = case_root / f"Agenda_{selected_file.stem}.pdf"
    client_agenda_docx = case_root / f"Agenda_{selected_file.stem}.docx"

    hidden_root = get_hidden_data_dir(working_dir)
    if sys.platform == "win32":
        # Keep .extracted_data hidden, but make dot-prefixed internal case folders visible to developers.
        try:
            for child in hidden_root.iterdir():
                if child.is_dir() and child.name.startswith("."):
                    subprocess.run(
                        ["attrib", "-h", str(child)],
                        check=False,
                        stdout=subprocess.DEVNULL,
                        stderr=subprocess.DEVNULL,
                        creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
                    )
        except Exception:
            pass
    internal_title = _build_internal_folder_title(working_dir, case_meta)
    hidden_case_name = _dot_hidden_folder_name(internal_title)
    case_key = hashlib.sha1(
        f"{str(working_dir.resolve())}|{selected_file.stem}".encode("utf-8")
    ).hexdigest()[:12]
    desired_internal_case_root = hidden_root / hidden_case_name
    existing_internal_case_root = _find_internal_case_root_by_key(hidden_root, case_key)
    internal_case_root = existing_internal_case_root or desired_internal_case_root

    # Rename existing internal folder to metadata-based title when we have better case metadata.
    if (
        case_meta
        and existing_internal_case_root
        and existing_internal_case_root.resolve() != desired_internal_case_root.resolve()
    ):
        rename_target = desired_internal_case_root
        if rename_target.exists() and _read_case_key_marker(rename_target) != case_key:
            rename_target = hidden_root / _dot_hidden_folder_name(f"{internal_title} {case_key}")
        try:
            existing_internal_case_root.rename(rename_target)
            internal_case_root = rename_target
        except Exception:
            internal_case_root = existing_internal_case_root

    internal_k138_dir = internal_case_root / "k138"
    internal_agenda_dir = internal_case_root / "agenda"
    internal_case_root.mkdir(parents=True, exist_ok=True)
    internal_k138_dir.mkdir(parents=True, exist_ok=True)
    internal_agenda_dir.mkdir(parents=True, exist_ok=True)
    _write_case_key_marker(internal_case_root, case_key)
    if sys.platform == "win32":
        try:
            subprocess.run(
                ["attrib", "-h", str(internal_case_root)],
                check=False,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
            )
        except Exception:
            pass

    internal_values = internal_case_root / "values_latest.json"
    internal_k138_latest = internal_k138_dir / "k138_latest.pdf"
    internal_agenda_latest_pdf = internal_agenda_dir / "agenda_latest.pdf"
    internal_agenda_latest_docx = internal_agenda_dir / "agenda_latest.docx"
    internal_barcode_png = internal_agenda_dir / "barcode.png"

    # Migrate old hidden runtime layout in working directory:
    # .extracted_data/case_runtime/<stem_hash>/...
    legacy_runtime_root = hidden_root / "case_runtime"
    legacy_internal_case_root = legacy_runtime_root / f"{selected_file.stem}_{case_key}"
    legacy_internal_values = legacy_internal_case_root / "values_latest.json"
    legacy_internal_k138_latest = legacy_internal_case_root / "k138" / "k138_latest.pdf"
    legacy_internal_agenda_latest_pdf = legacy_internal_case_root / "agenda" / "agenda_latest.pdf"
    legacy_internal_agenda_latest_docx = legacy_internal_case_root / "agenda" / "agenda_latest.docx"
    legacy_internal_barcode_png = legacy_internal_case_root / "agenda" / "barcode.png"
    if legacy_internal_values.exists() and (not internal_values.exists()):
        try:
            shutil.copy2(legacy_internal_values, internal_values)
        except Exception:
            pass
    if legacy_internal_k138_latest.exists() and (not internal_k138_latest.exists()):
        try:
            shutil.copy2(legacy_internal_k138_latest, internal_k138_latest)
        except Exception:
            pass
    if legacy_internal_agenda_latest_pdf.exists() and (not internal_agenda_latest_pdf.exists()):
        try:
            shutil.copy2(legacy_internal_agenda_latest_pdf, internal_agenda_latest_pdf)
        except Exception:
            pass
    if legacy_internal_agenda_latest_docx.exists() and (not internal_agenda_latest_docx.exists()):
        try:
            shutil.copy2(legacy_internal_agenda_latest_docx, internal_agenda_latest_docx)
        except Exception:
            pass
    if legacy_internal_barcode_png.exists() and (not internal_barcode_png.exists()):
        try:
            shutil.copy2(legacy_internal_barcode_png, internal_barcode_png)
        except Exception:
            pass

    # Migrate from older project-root hidden runtime to working-dir hidden runtime.
    old_project_hidden_root = get_hidden_data_dir(Path(__file__).parent)
    if old_project_hidden_root.resolve() != hidden_root.resolve():
        old_runtime_root = old_project_hidden_root / "case_runtime"
        old_case_root = old_runtime_root / f"{selected_file.stem}_{case_key}"
        old_values = old_case_root / "values_latest.json"
        old_k138 = old_case_root / "k138" / "k138_latest.pdf"
        old_agenda_pdf = old_case_root / "agenda" / "agenda_latest.pdf"
        old_agenda_docx = old_case_root / "agenda" / "agenda_latest.docx"
        old_barcode = old_case_root / "agenda" / "barcode.png"
        if old_values.exists() and (not internal_values.exists()):
            try:
                shutil.copy2(old_values, internal_values)
            except Exception:
                pass
        if old_k138.exists() and (not internal_k138_latest.exists()):
            try:
                shutil.copy2(old_k138, internal_k138_latest)
            except Exception:
                pass
        if old_agenda_pdf.exists() and (not internal_agenda_latest_pdf.exists()):
            try:
                shutil.copy2(old_agenda_pdf, internal_agenda_latest_pdf)
            except Exception:
                pass
        if old_agenda_docx.exists() and (not internal_agenda_latest_docx.exists()):
            try:
                shutil.copy2(old_agenda_docx, internal_agenda_latest_docx)
            except Exception:
                pass
        if old_barcode.exists() and (not internal_barcode_png.exists()):
            try:
                shutil.copy2(old_barcode, internal_barcode_png)
            except Exception:
                pass

    # Legacy visible structures from older releases:
    # 1) working_dir/k138 + working_dir/agenda
    # 2) working_dir/<file_stem>/k138 + working_dir/<file_stem>/agenda
    # 3) working_dir/<file_stem>/values_latest.json
    legacy_k138_dirs: List[Path] = [case_root / "k138"]
    legacy_agenda_dirs: List[Path] = [case_root / "agenda"]
    legacy_values_files: List[Path] = [case_root / "values_latest.json"]
    if legacy_nested_case_root != case_root:
        legacy_k138_dirs.append(legacy_nested_case_root / "k138")
        legacy_agenda_dirs.append(legacy_nested_case_root / "agenda")
        legacy_values_files.append(legacy_nested_case_root / "values_latest.json")

    # Migrate legacy values_latest.json to hidden runtime.
    for legacy_values in legacy_values_files:
        if not legacy_values.exists():
            continue
        try:
            if not internal_values.exists():
                shutil.copy2(legacy_values, internal_values)
            legacy_values.unlink(missing_ok=True)
        except Exception:
            pass

    # Migrate legacy internal artifacts (agenda/k138 latest + barcode) to hidden runtime.
    for legacy_k138_dir in legacy_k138_dirs:
        legacy_k138_latest = legacy_k138_dir / "k138_latest.pdf"
        if legacy_k138_latest.exists():
            try:
                if not internal_k138_latest.exists():
                    shutil.copy2(legacy_k138_latest, internal_k138_latest)
                legacy_k138_latest.unlink(missing_ok=True)
            except Exception:
                pass

    for legacy_agenda_dir in legacy_agenda_dirs:
        legacy_agenda_latest_pdf = legacy_agenda_dir / "agenda_latest.pdf"
        if legacy_agenda_latest_pdf.exists():
            try:
                if not internal_agenda_latest_pdf.exists():
                    shutil.copy2(legacy_agenda_latest_pdf, internal_agenda_latest_pdf)
                legacy_agenda_latest_pdf.unlink(missing_ok=True)
            except Exception:
                pass

        legacy_agenda_latest_docx = legacy_agenda_dir / "agenda_latest.docx"
        if legacy_agenda_latest_docx.exists():
            try:
                if not internal_agenda_latest_docx.exists():
                    shutil.copy2(legacy_agenda_latest_docx, internal_agenda_latest_docx)
                legacy_agenda_latest_docx.unlink(missing_ok=True)
            except Exception:
                pass

        legacy_barcode_png = legacy_agenda_dir / "barcode.png"
        if legacy_barcode_png.exists():
            try:
                if not internal_barcode_png.exists():
                    shutil.copy2(legacy_barcode_png, internal_barcode_png)
                legacy_barcode_png.unlink(missing_ok=True)
            except Exception:
                pass

    # Migrate old K138 output from legacy visible folders into hidden latest.
    if not internal_k138_latest.exists():
        k138_candidates: List[Path] = []
        for d in legacy_k138_dirs:
            k138_candidates.extend([
                d / "K138.pdf",
                d / f"K138_{selected_file.stem}.pdf",
            ])
        if legacy_nested_case_root != case_root:
            k138_candidates.extend([
                legacy_nested_case_root / "K138.pdf",
                legacy_nested_case_root / f"K138_{selected_file.stem}.pdf",
            ])
        for cand in k138_candidates:
            if not cand.exists():
                continue
            try:
                shutil.copy2(cand, internal_k138_latest)
                cand.unlink(missing_ok=True)
            except Exception:
                pass
            break

    if not client_agenda_pdf.exists():
        agenda_pdf_candidates: List[Path] = []
        for d in legacy_agenda_dirs:
            agenda_pdf_candidates.extend([
                d / f"Agenda_{selected_file.stem}.pdf",
                d / "agenda_latest.pdf",
            ])
        if legacy_nested_case_root != case_root:
            agenda_pdf_candidates.append(legacy_nested_case_root / f"Agenda_{selected_file.stem}.pdf")
        for cand in agenda_pdf_candidates:
            if cand.exists():
                try:
                    shutil.copy2(cand, client_agenda_pdf)
                    cand.unlink(missing_ok=True)
                except Exception:
                    pass
                break

    if not client_agenda_docx.exists():
        agenda_docx_candidates: List[Path] = []
        for d in legacy_agenda_dirs:
            agenda_docx_candidates.extend([
                d / f"Agenda_{selected_file.stem}.docx",
                d / "agenda_latest.docx",
            ])
        if legacy_nested_case_root != case_root:
            agenda_docx_candidates.append(legacy_nested_case_root / f"Agenda_{selected_file.stem}.docx")
        for cand in agenda_docx_candidates:
            if cand.exists():
                try:
                    shutil.copy2(cand, client_agenda_docx)
                    cand.unlink(missing_ok=True)
                except Exception:
                    pass
                break

    # Do NOT auto-expose K138 to the client-visible case folder.
    # K138.pdf must appear only after explicit "Generate K138".

    # Migrate legacy client name to the short standard name when present.
    if legacy_client_k138_pdf.exists():
        try:
            if not client_k138_pdf.exists():
                shutil.copy2(legacy_client_k138_pdf, client_k138_pdf)
            legacy_client_k138_pdf.unlink(missing_ok=True)
        except Exception:
            pass

    # Backfill internal latest files from client-visible files when needed.
    if client_k138_pdf.exists() and (not internal_k138_latest.exists()):
        try:
            shutil.copy2(client_k138_pdf, internal_k138_latest)
        except Exception:
            pass
    if client_agenda_pdf.exists() and (not internal_agenda_latest_pdf.exists()):
        try:
            shutil.copy2(client_agenda_pdf, internal_agenda_latest_pdf)
        except Exception:
            pass
    if client_agenda_docx.exists() and (not internal_agenda_latest_docx.exists()):
        try:
            shutil.copy2(client_agenda_docx, internal_agenda_latest_docx)
        except Exception:
            pass

    # Remove now-empty legacy visible directories.
    for old_dir in set(legacy_k138_dirs + legacy_agenda_dirs):
        try:
            old_dir.rmdir()
        except Exception:
            pass
    if legacy_nested_case_root != case_root:
        try:
            legacy_nested_case_root.rmdir()
        except Exception:
            pass

    return {
        "case_root": case_root,
        "internal_case_root": internal_case_root,
        "source_copy": selected_file,
        "values_latest_json": internal_values,
        "k138_dir": internal_k138_dir,
        "k138_output_pdf": client_k138_pdf,
        "client_k138_pdf": client_k138_pdf,
        "k138_latest_pdf": internal_k138_latest,
        "agenda_dir": internal_agenda_dir,
        "agenda_output_pdf": client_agenda_pdf,
        "agenda_output_docx": client_agenda_docx,
        "agenda_latest_pdf": internal_agenda_latest_pdf,
        "agenda_latest_docx": internal_agenda_latest_docx,
        "barcode_png": internal_barcode_png,
    }


def write_values_latest_json(path: Path, payload: Dict[str, object]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)


def read_values_latest_json(path: Path) -> Dict[str, object]:
    if not path.exists():
        return {}
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def update_values_latest_json(path: Path, patch: Dict[str, object]) -> Dict[str, object]:
    data = read_values_latest_json(path)
    data.update(patch or {})
    write_values_latest_json(path, data)
    return data


def _resolve_agenda_values_with_cache(
    case_paths: Dict[str, Path],
    inventory_number: str,
    agent_id: str,
) -> Tuple[str, str]:
    """
    Use cached case values when extracted values are missing or suspicious.
    Priority:
    - inventory: extracted first, cached fallback
    - agent: cached preferred on mismatch (badge from INSIGNE source is more reliable)
    """
    inv = _normalize_inventory_number(inventory_number)
    ag = re.sub(r"\D", "", agent_id or "")
    if ag:
        ag = ag[-5:] if len(ag) > 5 else ag.zfill(5)

    cached = read_values_latest_json(case_paths["values_latest_json"])
    cinv = _normalize_inventory_number(cached.get("inventory_number", ""))
    cag = re.sub(r"\D", "", cached.get("agent_id", "") or "")
    if cag:
        cag = cag[-5:] if len(cag) > 5 else cag.zfill(5)

    if not inv and cinv:
        inv = cinv
    if cag and (not ag or ag != cag):
        ag = cag
    return inv, ag


def ensure_case_source_file(selected_file: Path, case_paths: Dict[str, Path]) -> Tuple[Path, bool]:
    """
    Ensure the selected SAISIE file is copied into the case folder.
    Never moves/deletes the original from working directory.
    Returns (path_in_case_folder_or_original, copied_flag).
    """
    src = selected_file.resolve()
    dst = case_paths["source_copy"].resolve()
    if src == dst:
        return dst, False

    dst.parent.mkdir(parents=True, exist_ok=True)
    try:
        if dst.exists():
            dst.unlink()
    except Exception:
        pass

    try:
        shutil.copy2(str(src), str(dst))
        return dst, True
    except Exception:
        # Fallback: keep using original file when copy fails.
        return src, False


def _timestamp_compact() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def _timestamp_iso() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


CRITICAL_K138_FIELDS = [
    "description_inventory",
    "description_item",
    "notice_to",
    "seizing_officer",
]


def _safe_len(s: str) -> int:
    return len((s or "").strip())


def _normalize_token(s: str) -> str:
    return re.sub(r"\W+", "", (s or "").upper())


def _top_first_match(top: Dict[str, str], needles: List[str]) -> str:
    nd = [_normalize_token(n) for n in needles]
    for k, v in top.items():
        kn = _normalize_token(k)
        if any(n in kn for n in nd):
            cv = clean_value(v)
            if cv:
                return cv
    return ""


# ======================== =
# Form geometry (PDF points) -----
# ======================== =

# Page is US Letter: 612 x 792 points (PyMuPDF coordinate space)

FIELD_BOXES = {
    "BOND ROOM LEDGER #": (330, 0, 612, 30),
    "SIED #": (330, 30, 612, 60),

    "# INSIGNE AGENT SAISISSANT:": (21, 165, 245, 185),
    "# INVENTAIRE:": (245, 165, 445, 185),
    "PAYS:": (445, 165, 610, 185),

    "LIEU INTERCEPTION:": (21, 190, 245, 222),
    "DATE / HEURE INTERCEPTION:": (245, 190, 445, 222),
    "DÃ‰CLARATION:": (445, 190, 610, 222),

    "POIDS / QTÃ‰ MARCH.:": (21, 225, 245, 265),
    "DESCRIPTION DE Lâ€™ITEM Ã€ SAISIR:": (245, 225, 445, 265),

    "Notes": (445, 225, 610, 390),

    "EXPÃ‰DITEUR :": (21, 275, 245, 388),
    "DESTINATAIRE:": (245, 275, 445, 388),

    "INDICES:": (21, 388, 610, 520),

    "BOTTOM_TABLE_REGION": (21, 575, 610, 770),
}

BOTTOM_COLS = ["Date et Heure", "DÃ©placement de", "Acheminement Ã ", "Nom ASF", "Insigne", "TÃ©moin", "PIN"]
BOTTOM_X_BINS = [21, 115, 240, 360, 445, 505, 560, 610]  # 7 cols => 8 edges

TOP_CSV_FIELDS = [
    "BOND ROOM LEDGER #",
    "SIED #",

    "1-AEDS :",
    "Date examen:",
    "# Rapport :",
    "RÃ©sultat :",
    "2-AEADS (MTL) :",
    "Date examen:",
    "# Rapport :",
    "RÃ©sultat :",
    "3-Y15 (OTT) :",
    "Date examen:",
    "# Rapport :",
    "RÃ©sultat :",
    "SAISIE:",
    "CONFISCATION:",
    "K9:",
    "SAISIE Dâ€™ENVERGURE :",
    "# INSIGNE AGENT SAISISSANT:",
    "# INVENTAIRE:",
    "PAYS:",
    "LIEU INTERCEPTION:",
    "DATE / HEURE INTERCEPTION:",
    "DÃ‰CLARATION:",
    "POIDS / QTÃ‰ MARCH.:",
    "DESCRIPTION DE Lâ€™ITEM Ã€ SAISIR:",
    "EXPÃ‰DITEUR :",
    "DESTINATAIRE:",
    "INDICES:",
    "Notes",
]


# ======================== =
# Helpers -----
# ======================== =

def _repair_mojibake_text(s: str) -> str:
    """
    Repair common mojibake patterns like 'DÃ‰CLARATION' -> 'DÉCLARATION'.
    """
    out = str(s or "")
    for _ in range(2):
        if not any(ch in out for ch in ("\u00C3", "\u00C2", "\u00E2")):
            break
        candidate = ""
        for enc in ("latin1", "cp1252"):
            try:
                candidate = out.encode(enc).decode("utf-8")
                break
            except UnicodeError:
                candidate = ""
        if not candidate or candidate == out:
            break
        out = candidate
    return out


def _normalize_common_ocr_french(s: str) -> str:
    """
    Post-OCR normalization for frequent French label/content artifacts.
    """
    out = str(s or "")
    out = re.sub(r"\bDECLARATION\b", "DÉCLARATION", out, flags=re.IGNORECASE)
    out = re.sub(r"\bEXPEDITEUR\b", "EXPÉDITEUR", out, flags=re.IGNORECASE)
    out = re.sub(r"\bTEMOIN\b", "TÉMOIN", out, flags=re.IGNORECASE)
    # Common OCR split for "À l'intérieur"
    out = re.sub(r"\bA\s+V['’]?(?=int[eé]rieur)", "À l'", out, flags=re.IGNORECASE)
    return out


def normalize_output_text(s: str) -> str:
    """Normalize text for CSV/log/UI output without aggressive trimming."""
    out = _repair_mojibake_text(str(s or ""))
    out = _normalize_common_ocr_french(out)
    return out


def clean_value(s: str) -> str:
    s = normalize_output_text(s or "")
    s = s.replace("\u2019", "'").replace("\u2018", "'")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip(" \n\t:_|-")

def in_box(word: Tuple[float, float, float, float, str], box: Tuple[float, float, float, float]) -> bool:
    x0, y0, x1, y1, _ = word
    bx0, by0, bx1, by1 = box
    return (x0 >= bx0) and (x1 <= bx1) and (y0 >= by0) and (y1 <= by1)

def words_from_pdf(pdf_path: Path) -> List[Tuple[float, float, float, float, str]]:
    doc = fitz.open(str(pdf_path))
    page = doc.load_page(0)
    words = page.get_text("words")  # x0,y0,x1,y1,word,block,line,wordno
    doc.close()
    return [(x0, y0, x1, y1, w) for (x0, y0, x1, y1, w, *_rest) in words]


def get_page_text(pdf_path: Path) -> str:
    """Full page 0 text for label-anchored regex extraction (more reliable than position alone)."""
    try:
        doc = fitz.open(str(pdf_path))
        page = doc.load_page(0)
        text = page.get_text("text")
        doc.close()
        return text or ""
    except Exception:
        return ""

def word_key(w: Tuple[float, float, float, float, str]) -> Tuple[float, float, float, float, str]:
    x0, y0, x1, y1, t = w
    return (round(x0, 1), round(y0, 1), round(x1, 1), round(y1, 1), t)

def diff_words(template_words, completed_words) -> List[Tuple[float, float, float, float, str]]:
    tmpl_keys = set(word_key(w) for w in template_words)
    extras = [w for w in completed_words if word_key(w) not in tmpl_keys]
    extras.sort(key=lambda t: (t[1], t[0]))
    return extras

def join_words(extras: List[Tuple[float, float, float, float, str]]) -> str:
    if not extras:
        return ""
    lines: List[List[str]] = []
    cur: List[str] = []
    cur_y = extras[0][1]
    for x0, y0, x1, y1, t in extras:
        if abs(y0 - cur_y) > 5:
            lines.append(cur)
            cur = [t]
            cur_y = y0
        else:
            cur.append(t)
    if cur:
        lines.append(cur)
    # Dedupe: merge consecutive duplicate lines
    seen: Dict[str, bool] = {}
    unique_lines: List[str] = []
    for line in lines:
        line_txt = " ".join(line).strip()
        if line_txt and line_txt not in seen:
            seen[line_txt] = True
            unique_lines.append(line_txt)
    return clean_value("\n".join(unique_lines))

# ---- OCR helpers (for image-based PDFs) ----
def render_clip(pdf_path: Path, clip: Tuple[float, float, float, float], dpi: int = 300) -> "Image.Image":
    if not HAVE_PIL:
        raise RuntimeError("PIL is required for OCR")
    doc = fitz.open(str(pdf_path))
    page = doc.load_page(0)
    rect = fitz.Rect(*clip)
    mat = fitz.Matrix(dpi / 72.0, dpi / 72.0)
    pix = page.get_pixmap(matrix=mat, clip=rect, alpha=False)
    doc.close()
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    return img

def ocr_image(img: "Image.Image") -> str:
    if not HAVE_TESSERACT:
        return ""
    txt = pytesseract.image_to_string(img, lang="fra+eng", config="--psm 6")
    return clean_value(txt)


def get_ocr_text_from_image(image_path: Path) -> str:
    """OCR a standalone image file (png/jpg/tif/bmp) and return cleaned text."""
    if not HAVE_PIL or not HAVE_TESSERACT:
        return ""
    try:
        with Image.open(str(image_path)) as img:
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")
            txt = pytesseract.image_to_string(img, lang="fra+eng", config="--psm 6")
            return clean_value(txt)
    except Exception:
        return ""

def token_diff(template_text: str, filled_text: str) -> str:
    """
    Remove template tokens from filled while preserving meaningful repeats.
    Important: repeated numeric chunks (e.g., "0000" in inventory numbers)
    must be kept in order.
    """
    t = clean_value(template_text)
    f = clean_value(filled_text)
    if not f:
        return ""
    ttoks = [x for x in re.split(r"\s+", t) if x]
    ftoks = [x for x in re.split(r"\s+", f) if x]
    tset = set(x.upper() for x in ttoks)
    keep: List[str] = []
    for tok in ftoks:
        up = tok.upper()
        if (up not in tset) or re.search(r"\d", tok):
            keep.append(tok)
    return clean_value(" ".join(keep))


def token_diff_preserve_lines(template_text: str, filled_text: str) -> str:
    """Like token_diff but preserves line breaks from filled text (for addresses)."""
    t = clean_value(template_text)
    ttoks = set(x.upper() for x in re.split(r"\s+", t) if x)
    lines_out = []
    for line in (filled_text or "").splitlines():
        line = line.strip()
        if not line:
            continue
        ftoks = [x for x in re.split(r"\s+", line) if x]
        keep = []
        for tok in ftoks:
            up = tok.upper()
            if (up not in ttoks) or re.search(r"\d", tok):
                keep.append(tok)
        if keep:
            lines_out.append(" ".join(keep))
    return clean_value("\n".join(lines_out))


def _expand_two_digit_year(yy: int) -> int:
    """Convert 2-digit year to 4-digit year (default modern cases)."""
    yy = int(yy)
    return 2000 + yy if yy <= 79 else 1900 + yy


def _is_valid_calendar_date(year: int, month: int, day: int) -> bool:
    try:
        datetime(year, month, day)
        return True
    except Exception:
        return False


def parse_first_date(s: str) -> Optional[Tuple[int, int, int]]:
    s = clean_value(s or "")
    if not s:
        return None

    numeric_patterns = [
        # yyyy-mm-dd or yyyy/mm/dd, optional time (including datetime with T separator).
        (r"\b(\d{4})[./-](\d{1,2})[./-](\d{1,2})(?:[T ][\d:.]+(?:\s*[APMapm]{2})?)?\b", "ymd"),
        # dd-mm-yyyy or dd/mm/yyyy, optional time.
        (r"\b(\d{1,2})[./-](\d{1,2})[./-](\d{4})(?:[T ][\d:.]+(?:\s*[APMapm]{2})?)?\b", "dmy"),
        # Ambiguous 2-digit year: yy-mm-dd (e.g. 24-02-26 = 2024-02-26) — try ymd2 first when
        # first group looks like a plausible 2-digit year (00-99) AND result is valid.
        (r"\b(\d{2})[./-](\d{1,2})[./-](\d{2})(?:[T ][\d:.]+(?:\s*[APMapm]{2})?)?\b", "ambig"),
    ]
    for pat, order in numeric_patterns:
        m = re.search(pat, s)
        if not m:
            continue
        a, b, c = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if order == "ymd":
            yyyy, mm, dd = a, b, c
            if _is_valid_calendar_date(yyyy, mm, dd):
                return yyyy, mm, dd
        elif order == "dmy":
            dd, mm, yyyy = a, b, c
            if _is_valid_calendar_date(yyyy, mm, dd):
                return yyyy, mm, dd
        elif order == "ambig":
            # For ambiguous xx-yy-zz: try yy-mm-dd interpretation first (most common on SAISIE forms)
            # A valid year candidate is <= 99 (2-digit). We prefer ymd2 (yy-mm-dd) since
            # the forms are dated within the current decade and that's how officers write dates.
            # Heuristic: if b is 1-12 and c is 1-31 and a is 0-99 -> yy-mm-dd
            # Also try dmy2 (dd-mm-yy) as fallback.
            candidates = []
            # Try yy-mm-dd: a=year, b=month, c=day
            yyyy_ymd = _expand_two_digit_year(a)
            if _is_valid_calendar_date(yyyy_ymd, b, c):
                candidates.append((yyyy_ymd, b, c))
            # Try dd-mm-yy: a=day, b=month, c=year
            yyyy_dmy = _expand_two_digit_year(c)
            if _is_valid_calendar_date(yyyy_dmy, b, a):
                candidates.append((yyyy_dmy, b, a))
            if not candidates:
                continue
            # If both are valid, prefer yy-mm-dd (year first) since that's typical on SAISIE
            # UNLESS the day-first candidate has a more recent year (closer to current year)
            if len(candidates) == 1:
                yyyy, mm, dd = candidates[0]
            else:
                # Prefer whichever interpretation gives a year that looks like a valid recent year
                # Both are valid - use yy-mm-dd (first candidate) as primary for SAISIE forms
                yyyy, mm, dd = candidates[0]
            return yyyy, mm, dd

    # Unicode-safe "day month year" matcher (avoids fragile accented char ranges).
    month_matches = re.finditer(
        r"(\d{1,2})\s+([^\W\d_]+)\s+(\d{2,4})(?:\s+\d{1,2}:\d{2}(?::\d{2})?)?",
        s,
        flags=re.UNICODE,
    )
    month_map = {
        "janvier": 1, "january": 1,
        "fevrier": 2, "february": 2,
        "mars": 3, "march": 3,
        "avril": 4, "april": 4,
        "mai": 5, "may": 5,
        "juin": 6, "june": 6,
        "juillet": 7, "july": 7,
        "aout": 8, "august": 8,
        "septembre": 9, "september": 9,
        "octobre": 10, "october": 10,
        "novembre": 11, "november": 11,
        "decembre": 12, "december": 12,
    }
    for m in month_matches:
        dd = int(m.group(1))
        month_raw = normalize_output_text(m.group(2)).lower()
        month_raw = (
            month_raw.replace("é", "e")
            .replace("è", "e")
            .replace("ê", "e")
            .replace("ë", "e")
            .replace("à", "a")
            .replace("â", "a")
            .replace("ä", "a")
            .replace("û", "u")
            .replace("ù", "u")
            .replace("ü", "u")
            .replace("î", "i")
            .replace("ï", "i")
            .replace("ô", "o")
            .replace("ö", "o")
            .replace("ç", "c")
        )
        month_raw = re.sub(r"[^a-z]", "", month_raw)
        mm = month_map.get(month_raw)
        if not mm:
            continue
        year_raw = int(m.group(3))
        yyyy = _expand_two_digit_year(year_raw) if year_raw < 100 else year_raw
        if _is_valid_calendar_date(yyyy, mm, dd):
            return yyyy, mm, dd
    return None

def month_en_fr(mm: int) -> Tuple[str, str]:
    en = ["JANUARY","FEBRUARY","MARCH","APRIL","MAY","JUNE","JULY","AUGUST","SEPTEMBER","OCTOBER","NOVEMBER","DECEMBER"]
    fr = ["JANVIER","FÉVRIER","MARS","AVRIL","MAI","JUIN","JUILLET","AOÛT","SEPTEMBRE","OCTOBRE","NOVEMBRE","DÉCEMBRE"]
    mm = max(1, min(12, mm))
    return en[mm-1], fr[mm-1]

def normalize_multiline_to_pipe(s: str) -> str:
    lines = [clean_value(x) for x in (s or "").splitlines()]
    lines = [x for x in lines if x]
    return " | ".join(lines)

def _dedupe_lines_keep_order(s: str) -> str:
    """Remove duplicate lines while preserving first occurrence order."""
    if not s:
        return ""
    seen = set()
    out = []
    for line in (s or "").splitlines():
        line = clean_value(line)
        if not line:
            continue
        k = line.upper()
        if k in seen:
            continue
        seen.add(k)
        out.append(line)
    return "\n".join(out)

_COUNTRY_CANON = {
    # Canada
    "CANADA": "CANADA",
    "CAN": "CANADA",
    # USA — English and French variants
    "USA": "USA",
    "JSA": "USA",   # OCR artifact
    "ISA": "USA",   # OCR artifact
    "U5A": "USA",   # OCR artifact
    "US": "USA",
    "UNITED STATES": "USA",
    "UNITED STATES OF AMERICA": "USA",
    "ETATS-UNIS": "USA",
    "ÉTATS-UNIS": "USA",
    "ETATS UNIS": "USA",
    "ÉTATS UNIS": "USA",
    "E.U.": "USA",
    # UK — English and French variants
    "UK": "UK",
    "UNITED KINGDOM": "UK",
    "ROYAUME-UNI": "UK",
    "ROYAUME UNI": "UK",
    "GB": "UK",
    "GREAT BRITAIN": "UK",
    "GRANDE-BRETAGNE": "UK",
    "GRANDE BRETAGNE": "UK",
}

_RE_POSTAL_CA = re.compile(r"\b[ABCEGHJ-NPRSTVXY]\d[ABCEGHJ-NPRSTV-Z]\s*\d[ABCEGHJ-NPRSTV-Z]\d\b", re.IGNORECASE)
_RE_ZIP_US = re.compile(r"\b\d{5}(?:-\d{4})?\b")
_RE_POSTCODE_UK = re.compile(r"\b[A-Z]{1,2}\d[A-Z\d]?\s*\d[A-Z]{2}\b", re.IGNORECASE)
_RE_POSTAL_CA_ONLY = re.compile(r"^\s*([ABCEGHJ-NPRSTVXY]\d[ABCEGHJ-NPRSTV-Z])\s*(\d[ABCEGHJ-NPRSTV-Z]\d)\s*$", re.IGNORECASE)
_RE_ZIP_US_ONLY = re.compile(r"^\s*\d{5}(?:-\d{4})?\s*$")
_RE_POSTCODE_UK_ONLY = re.compile(r"^\s*[A-Z]{1,2}\d[A-Z\d]?\s*\d[A-Z]{2}\s*$", re.IGNORECASE)


def _canonical_country_name(line: str) -> str:
    token = clean_value(line).upper()
    token = re.sub(r"[.,]", " ", token)
    token = re.sub(r"\s+", " ", token).strip()
    return _COUNTRY_CANON.get(token, "")


def _clean_notice_artifact_line(line: str) -> str:
    """Remove common OCR/address artifacts before layout normalization."""
    ln = clean_value(line)
    if not ln:
        return ""
    # Common OCR noise prefixes found in destination lines.
    ln = re.sub(r"^(?:ADRESSE\s+ILLISIBLE|ILLISIBLE)\s+", "", ln, flags=re.IGNORECASE)
    ln = re.sub(r"^(?:A/?C|C/O|CO)\s+", "", ln, flags=re.IGNORECASE)
    # Merged sender/recipient line often ends with the true recipient token (e.g., "S. Jones").
    if "," in ln:
        m = re.search(r"([A-Za-z]\.\s*[A-Za-z][A-Za-z'’.\-]+)$", ln)
        if m:
            ln = m.group(1)
    return clean_value(ln)


def _infer_country_from_text(line: str) -> str:
    t = (line or "").upper()
    if _RE_POSTAL_CA.search(t):
        return "CANADA"
    if _RE_ZIP_US.search(t):
        return "USA"
    if _RE_POSTCODE_UK.search(t):
        return "UK"
    if "CANADA" in t:
        return "CANADA"
    if re.search(r"\b(USA|US|UNITED STATES|ETATS-UNIS|ÉTATS-UNIS|ETATS UNIS|ÉTATS UNIS)\b", t):
        return "USA"
    if re.search(r"\b(UK|UNITED KINGDOM|ROYAUME-UNI|ROYAUME UNI|GREAT BRITAIN|GRANDE-BRETAGNE)\b", t):
        return "UK"
    return ""


def _normalize_notice_address_layout(s: str) -> str:
    """
    Normalize address order for K138:
    - keep street/name lines first
    - force city/postal as second-last line when found
    - force country as last line when found or inferred
    """
    lines = [_clean_notice_artifact_line(x) for x in (s or "").splitlines() if clean_value(x)]
    lines = [ln for ln in lines if ln]
    if not lines:
        return ""

    # Remove obvious label noise while preserving address content.
    noise_re = re.compile(r"^(?:DESTINATAIRE|EXP(?:É|E|Ã‰)DITEUR|NOTES?)\s*:?\s*$", re.IGNORECASE)
    lines = [ln for ln in lines if not noise_re.fullmatch(ln)]
    if not lines:
        return ""

    # Drop duplicates before layout normalization.
    dedup = []
    seen = set()
    for ln in lines:
        key = ln.upper()
        if key in seen:
            continue
        seen.add(key)
        dedup.append(ln)
    lines = dedup

    def _normalize_postal_only_token(line: str) -> str:
        """Normalize standalone postal/zip tokens (e.g., L7A2S2 -> L7A 2S2)."""
        t = clean_value(line).upper()
        m = _RE_POSTAL_CA_ONLY.fullmatch(t)
        if m:
            return f"{m.group(1)} {m.group(2)}"
        return t

    def _is_postal_only(line: str) -> bool:
        t = clean_value(line)
        return bool(_RE_POSTAL_CA_ONLY.fullmatch(t) or _RE_ZIP_US_ONLY.fullmatch(t) or _RE_POSTCODE_UK_ONLY.fullmatch(t))

    country = ""
    kept: List[str] = []
    trailing_country_re = re.compile(
        r"\b(?:CANADA|USA|US|UNITED STATES(?: OF AMERICA)?|UK|UNITED KINGDOM)\b\.?\s*$",
        re.IGNORECASE,
    )
    leading_country_re = re.compile(
        r"^\s*(CANADA|USA|US|UNITED STATES(?: OF AMERICA)?|UK|UNITED KINGDOM)\b[,\s]+(.+)$",
        re.IGNORECASE,
    )

    for ln in lines:
        lead = leading_country_re.match(ln)
        if lead:
            tail = clean_value(lead.group(2))
            inferred_tail = _infer_country_from_text(tail)
            if inferred_tail:
                country = inferred_tail
            elif not country:
                country = _canonical_country_name(lead.group(1))
            ln = tail
            if not ln:
                continue

        canon = _canonical_country_name(ln)
        if canon:
            country = canon
            continue

        # If country appears at end of a mixed line, strip it from this line and keep separately.
        stripped = trailing_country_re.sub("", ln).strip(" ,")
        if stripped != ln:
            if not country:
                country = _infer_country_from_text(ln)
            if stripped:
                kept.append(stripped)
            continue

        kept.append(ln)

    # Merge postal-only line into the previous city/province line when split by OCR/layout.
    # Example: "Brampton ON" + "L7A 2S2" -> "Brampton ON L7A 2S2"
    merged: List[str] = []
    for ln in kept:
        # Keep US ZIP lines standalone (expected K138 layout has ZIP on its own line).
        is_ca_postal_only = bool(_RE_POSTAL_CA_ONLY.fullmatch(clean_value(ln)))
        if merged and _is_postal_only(ln) and is_ca_postal_only:
            prev = merged[-1]
            prev_has_postal = bool(_RE_POSTAL_CA.search(prev) or _RE_ZIP_US.search(prev) or _RE_POSTCODE_UK.search(prev))
            # Only merge when previous line looks like city/province (typically no digits).
            if (not prev_has_postal) and (not re.search(r"\d", prev)):
                merged[-1] = f"{prev} {_normalize_postal_only_token(ln)}".strip()
                continue
            merged.append(_normalize_postal_only_token(ln))
            continue
        merged.append(ln)
    kept = merged

    city_postal = ""
    city_idx = -1
    for i in range(len(kept) - 1, -1, -1):
        ln = kept[i]
        if _RE_POSTAL_CA.search(ln) or _RE_ZIP_US.search(ln) or _RE_POSTCODE_UK.search(ln):
            city_postal = ln
            city_idx = i
            break

    if city_idx >= 0:
        kept.pop(city_idx)

    if city_postal and not country:
        country = _infer_country_from_text(city_postal)

    if not country:
        for ln in kept:
            inferred = _infer_country_from_text(ln)
            if inferred:
                country = inferred
                break

    out = [ln for ln in kept if ln]
    if city_postal:
        out.append(city_postal)
    if country:
        out.append(country)

    # Final de-duplication in final order.
    final = []
    seen_final = set()
    for ln in out:
        k = ln.upper()
        if k in seen_final:
            continue
        seen_final.add(k)
        final.append(ln)

    return "\n".join(final)


def _clean_notice_address(s: str) -> str:
    """Normalize destination address block and remove repeated label/line noise."""
    if not s:
        return ""
    v = _strip_repeated_label_blocks(s, r"DESTINATAIRE\s*:?\s*")
    # Hard stop at NOTES section even when OCR merges labels into one line.
    v = re.split(r"\bNOTES?\b\s*:?", v, maxsplit=1, flags=re.IGNORECASE)[0]
    # Remove header/placeholder noise often present in template rows.
    v = re.sub(r"\bDESTINATAIRE\s*:?\s*", "", v, flags=re.IGNORECASE)
    v = re.sub(r"\bINCONNU\b\s*_*\s*", "", v, flags=re.IGNORECASE)
    v = re.sub(r"_+", " ", v)
    v = _dedupe_lines_keep_order(v)
    stop_re = re.compile(
        r"(?:DATE\s*/\s*HEURE|LIEU\s+INTERCEPTION|D(?:Ã‰|E|É)CLARATION|POIDS|INDICES|CHECKLIST|NOTES?|"
        r"EXP(?:Ã‰|É|E)DITEUR|"
        r"ACHEMINEMENT|BOND\s+ROOM|SIED|#\s*INVENTAIRE|#\s*INSIGNE)",
        re.IGNORECASE,
    )
    filtered: List[str] = []
    for ln in v.splitlines():
        line = _clean_notice_artifact_line(ln)
        if not line:
            continue
        if re.match(r"^(?:DESTINATAIRE|EXP(?:Ã‰|É|E)DITEUR)\b", line, re.IGNORECASE):
            continue
        if stop_re.search(line):
            break
        # Drop leaked date/time-only lines from other sections.
        if re.fullmatch(r"\d{4}[-/]\d{2}[-/]\d{2}(?:\s+\d{1,2}:\d{2})?", line):
            continue
        if re.fullmatch(r"\d{1,2}:\d{2}", line):
            continue
        filtered.append(line)
        # Destination block should end at country line.
        if _canonical_country_name(line):
            break
    return _normalize_notice_address_layout("\n".join(filtered))


def _strip_repeated_label_blocks(s: str, label_pattern: str) -> str:
    """Strip repeated 'LABEL: value' from table-style extraction; keep unique content."""
    if not s:
        return ""
    # Split by label pattern; first part is usually the value, rest are duplicates
    parts = re.split(label_pattern, s, flags=re.IGNORECASE)
    # Collect non-empty parts, skip duplicates
    seen = set()
    out_lines = []
    for p in parts:
        p = clean_value(p)
        if not p:
            continue
        for line in p.splitlines():
            line = clean_value(line)
            if not line:
                continue
            key = line.upper().strip()
            if key not in seen:
                seen.add(key)
                out_lines.append(line)
    return "\n".join(out_lines) if out_lines else ""

def safe_get(top: Dict[str, str], k: str) -> str:
    return clean_value(top.get(k, ""))

def _normalize_declared_text(s: str) -> str:
    """Clean OCR artifacts in declaration field (strip label remnants, fix common truncated words)."""
    v = clean_value(s)
    if not v:
        return ""
    # Remove leading label artifacts like "DÃ‰CLARATION:", "DECLARATION:", "ECLARATION:"
    v = re.sub(r"^(?:D|E)?\s*(?:[Ã‰E])?\s*CLARATION\s*:?\s*", "", v, flags=re.IGNORECASE)
    # Also handle variants seen in K138 and OCR outputs.
    v = re.sub(
        r"^(?:DECLARED|D(?:É|E|Ã‰|\?)CLAR(?:É|E|Ã‰|\?))\s*(?:/\s*(?:DECLARED|D(?:É|E|Ã‰|\?)CLAR(?:É|E|Ã‰|\?)))?\s*:?\s*",
        "",
        v,
        flags=re.IGNORECASE,
    )
    v = re.sub(r"^/\s*D(?:É|E|Ã‰|\?)CLAR(?:É|E|Ã‰|\?)\s*:?\s*", "", v, flags=re.IGNORECASE)
    v = clean_value(v)
    # Common OCR truncation: "ucune" -> "Aucune"
    if re.search(r"\baucune\b", v, flags=re.IGNORECASE):
        return "Aucune"
    if re.fullmatch(r"[uU]?\s*cune", v):
        return "Aucune"
    return v

def _extract_declared_from_top(top: Dict[str, str]) -> str:
    """Read declaration value from any variant key, then normalize."""
    candidates = [
        "DÃ‰CLARATION:",
        "DECLARATION:",
        "DÉCLARATION:",
        "DECLARED / DÉCLARÉ:",
        "DECLARED/DÉCLARÉ:",
        "DECLARED:",
        "DÉCLARÉ:",
        "DÃƒâ€°CLARATION:",
        "D\\u00C9CLARATION:",
        "DECLARATION / DÉCLARATION:",
        "DÉCLARATION / DECLARATION:",
    ]
    raw = ""
    for k in candidates:
        raw = safe_get(top, k)
        if raw:
            break
    if not raw:
        raw = _top_first_match(top, ["DECLARED", "DÉCLARÉ", "DECLARATION", "DÉCLARATION"])
    if not raw:
        # Last fallback: fuzzy key search for any field containing 'CLARATION' or 'DECLARED'
        for k, v in top.items():
            kn = normalize_output_text(k or "").upper()
            kn_stripped = re.sub(r"[^A-Z0-9]", "", kn)
            if (
                ("CLARATION" in kn)
                or ("DECLARED" in kn)
                or ("CLARATION" in kn_stripped)
            ) and clean_value(v):
                raw = clean_value(v)
                break
    if not raw:
        # Word edge-case fallback: declaration text can be pasted inside DESCRIPTION field.
        for k, v in top.items():
            kn = normalize_output_text(k or "").upper()
            if ("DESCRIPTION" not in kn) and ("ITEM" not in kn):
                continue
            vv = clean_value(v)
            if not vv:
                continue
            m = re.search(
                r"(?:DECLARED\s*/\s*D(?:É|E|Ã‰|\?)CLAR(?:É|E|Ã‰|\?)|D(?:É|E|Ã‰|\?)CLARATION)\s*:?\s*([^\n\r]+)",
                vv,
                re.IGNORECASE,
            )
            if m:
                raw = clean_value(m.group(1))
                break
    return _normalize_declared_text(raw)

def _extract_seizure_number(top: Dict[str, str]) -> str:
    """Extract a clean seizure number from noisy OCR fields."""
    sied = safe_get(top, "SIED #")
    bond = safe_get(top, "BOND ROOM LEDGER #")

    # 1) Prefer explicit SIED-like pattern (e.g., 3952-25-1234)
    if sied:
        m = re.search(r"\b\d{3,5}-\d{2}-\d{3,5}\b", sied)
        if m:
            return m.group(0)
        m = re.search(r"\bSIED\s*#?\s*([A-Z0-9-]{4,20})\b", sied, re.IGNORECASE)
        if m:
            return m.group(1)

    # 2) Try bond ledger marker
    if bond:
        m = re.search(r"(?:BOND\s+ROOM\s+LEDGER|BOND\s+LEDGER|LEDGER)\s*#?\s*([A-Z0-9-]{2,20})", bond, re.IGNORECASE)
        if m:
            return m.group(1)
        # 3) Fallback: first short numeric token (common ledger style like 371)
        nums = re.findall(r"\b\d{3,6}\b", bond)
        if nums:
            return nums[0]
        nums = re.findall(r"\b\d{2,6}\b", bond)
        if nums:
            return nums[0]

    # 4) Last fallback: short meaningful token from SIED if any
    if sied:
        tokens = re.findall(r"\b[A-Z0-9-]{3,20}\b", sied.upper())
        for tok in tokens:
            if any(ch.isdigit() for ch in tok):
                return tok
    return ""


# ======================== =
# Label-anchored regex fallbacks (Dmitry: more reliable than position alone) -----
# ======================== =

# Terminating labels for address blocks: stop capturing when we hit the next section
ADDRESS_TERMINATING_LABELS = re.compile(
    r"(?:INDICES:|LIEU INTERCEPTION|DATE / HEURE|DÃ‰CLARATION|POIDS|DESCRIPTION DE L'ITEM|EXPÃ‰DITEUR|DESTINATAIRE|# INVENTAIRE|# INSIGNE)",
    re.IGNORECASE
)


def _normalize_inventory_number(raw: str) -> str:
    """
    Normalize inventory number from noisy OCR/text.
    Canonical target:
    - 2 letters + 9 digits + 2 letters (e.g., AB123456789CA)
    Rules:
    - remove spaces/separators
    - uppercase
    - ignore label/junk characters
    Compatible fallback for legacy format:
    - 1 letter + long digit sequence (e.g., W00006042000043012)
    """
    text = normalize_output_text(clean_value(raw or ""))
    if not text:
        return ""

    # Keep value part when full label/value text is present.
    m = re.search(
        r"(?:NO\.?\s*D['’]?\s*INVENTAIRE|NO\s*INVENTAIRE|#\s*INVENTAIRE|INVENTAIRE)\s*:?\s*(.+)",
        text,
        re.IGNORECASE | re.DOTALL,
    )
    if m:
        text = clean_value(m.group(1))

    # Upper + remove separators.
    compact = re.sub(r"[^A-Za-z0-9]+", "", text).upper()
    if not compact:
        return ""

    # Strip common label junk that can leak into token stream.
    compact = re.sub(
        r"(?:NODINVENTAIRE|NOINVENTAIRE|INVENTAIRE|NUMEROINVENTAIRE|NUMERODEINVENTAIRE)",
        "",
        compact,
        flags=re.IGNORECASE,
    )
    compact = re.sub(
        r"(?:DESTINATAIRE|EXPEDITEUR|DECLARATION|DESCRIPTION|SIED|BONDROOMLEDGER|CHECKLIST|INDICES|NOTES?)",
        "",
        compact,
        flags=re.IGNORECASE,
    )

    # 1) Strict canonical format: 2 letters + 9 digits + 2 letters.
    m = re.search(r"[A-Z]{2}\d{9}[A-Z]{2}", compact)
    if m:
        return m.group(0)

    # 2) Frequent legacy format: one leading letter + long digits.
    # OCR often misreads leading 'W' as 'V' for this format; normalize to W.
    m = re.search(r"[A-Z]\d{10,20}", compact)
    if m:
        legacy = m.group(0)
        if (
            legacy.startswith("V")
            and legacy[1:].isdigit()
            and legacy[1:].startswith("0000")
        ):
            legacy = f"W{legacy[1:]}"
        return legacy

    # 3) Tolerant fallback: compact token containing letters+digits.
    for tok in re.findall(r"[A-Z0-9]{8,40}", compact):
        if re.search(r"[A-Z]", tok) and re.search(r"\d", tok):
            return tok

    return ""


def extract_inventory_by_label(text: str) -> str:
    """Extract inventory number near # INVENTAIRE: / NO INVENTAIRE: via regex."""
    if not text:
        return ""
    # Match # INVENTAIRE / NO INVENTAIRE and pull a meaningful token after it.
    for pattern in [
        r"(?:#\s*INVENTAIRE|NO\s*INVENTAIRE)\s*:?\s*(?:\n|\s)*([A-Z0-9][A-Z0-9\-/ ]{5,40})",
        r"(?:INVENTAIRE)\s*:?\s*(?:\n|\s)*([A-Z0-9][A-Z0-9\-/ ]{5,40})",
    ]:
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            v = clean_value(m.group(1))
            norm = _normalize_inventory_number(v)
            if norm:
                return norm
            tok = re.search(r"[A-Z0-9][A-Z0-9\-]{7,30}", v, re.IGNORECASE)
            if tok:
                return tok.group(0)
            return v
    return ""

def extract_officer_by_label(text: str) -> str:
    """Extract 5-digit officer number near # INSIGNE AGENT SAISISSANT."""
    if not text:
        return ""
    # Common DOCX/OCR order: badge digits appear before the label.
    m = re.search(r"(\d{4,6})\s*(?:#\s*INSIGNE\s+AGENT\s+SAISISSANT|INSIGNE\s+AGENT\s+SAISISSANT)", text, re.IGNORECASE)
    if m:
        num = m.group(1).strip()
        return num[-5:] if len(num) > 5 else num.zfill(5) if len(num) < 5 else num
    m = re.search(r"(?:#\s*INSIGNE\s+AGENT\s+SAISISSANT|INSIGNE\s+AGENT\s+SAISISSANT)\s*:?\s*(\d{4,6})", text, re.IGNORECASE)
    if m:
        num = m.group(1).strip()
        return num[-5:] if len(num) > 5 else num.zfill(5) if len(num) < 5 else num
    m = re.search(r"(?:#\s*INSIGNE|SAISISSANT)\s*:?\s*(\d{5})", text, re.IGNORECASE)
    if m:
        return m.group(1)
    return ""


def extract_interception_date_by_label(text: str) -> Optional[Tuple[int, int, int]]:
    """Extract DATE / HEURE INTERCEPTION value and parse a date from that block only."""
    if not text:
        return None
    patterns = [
        r"DATE\s*/\s*HEURE\s*INTERCEPTION\s*:?\s*(.+?)(?=\n\s*(?:D(?:É|E)CLARATION|POIDS|DESCRIPTION|DESTINATAIRE|EXP(?:É|E)DITEUR|INDICES|NOTES?)\b|\Z)",
        r"DATE\s+HEURE\s+INTERCEPTION\s*:?\s*(.+?)(?=\n\s*(?:D(?:É|E)CLARATION|POIDS|DESCRIPTION|DESTINATAIRE|EXP(?:É|E)DITEUR|INDICES|NOTES?)\b|\Z)",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE | re.DOTALL)
        if not m:
            continue
        block = clean_value(m.group(1))
        dt = parse_first_date(block)
        if dt:
            return dt
    return None

def extract_declaration_by_label(text: str) -> str:
    """Extract DÉCLARATION value from page text and normalize it."""
    if not text:
        return ""
    patterns = [
        r"D(?:É|E|Ã‰|\?)CLARATION\s*:?\s*(.+?)(?=\n\s*(?:POIDS|DESCRIPTION|DESTINATAIRE|EXP(?:É|E|Ã‰|\?)DITEUR|INDICES|NOTES?|LIEU\s+INTERCEPTION|DATE\s*/\s*HEURE|#\s*INVENTAIRE|#\s*INSIGNE)\b|\Z)",
        r"D(?:É|E|Ã‰|\?)CLARATION\s*:?\s*([^\n\r]+)",
        r"DECLARED\s*/\s*D(?:É|E|Ã‰|\?)CLAR(?:É|E|Ã‰|\?)\s*:?\s*(.+?)(?=\n\s*(?:POIDS|DESCRIPTION|DESTINATAIRE|EXP(?:É|E|Ã‰|\?)DITEUR|INDICES|NOTES?|LIEU\s+INTERCEPTION|DATE\s*/\s*HEURE|#\s*INVENTAIRE|#\s*INSIGNE)\b|\Z)",
        r"(?:DECLARED|D(?:É|E|Ã‰|\?)CLAR(?:É|E|Ã‰|\?))\s*:?\s*([^\n\r]+)",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE | re.DOTALL)
        if not m:
            continue
        val = _normalize_declared_text(m.group(1))
        if val:
            return val
    return ""

def extract_address_by_label(text: str, label: str, until_labels: Optional[List[str]] = None) -> str:
    """Extract multi-line address after label until next section label. Handles line-wrapping + OCR chunking."""
    if not text:
        return ""
    until = until_labels or ["INDICES:", "EXPÃ‰DITEUR", "DESTINATAIRE", "LIEU INTERCEPTION", "DATE ", "Notes", "NOTES"]
    esc_label = re.escape(label.rstrip(" :"))
    esc_until = "|".join(re.escape(u) for u in until)
    pattern = r"(?:{})\s*:?\s*(.*?)(?=\n\s*(?:{})|\Z)".format(esc_label, esc_until)
    m = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
    if m:
        block = m.group(1)
        lines = [clean_value(ln) for ln in block.splitlines() if clean_value(ln)]
        return _normalize_notice_address_layout("\n".join(lines)) if lines else ""
    return ""


# ======================== =
# Extraction: template vs completed (text-diff first, OCR fallback) -----
# ======================== =

# Order fields by box area (smallest first) so each word is assigned to at most one box (no duplicates).
def _field_boxes_sorted() -> List[Tuple[str, Tuple[float, float, float, float]]]:
    order = []
    for name, box in FIELD_BOXES.items():
        if name == "BOTTOM_TABLE_REGION":
            continue
        x0, y0, x1, y1 = box
        order.append((name, box))
    order.sort(key=lambda x: (x[1][2] - x[1][0]) * (x[1][3] - x[1][1]))
    return order

def _assign_words_to_boxes(extras: List[Tuple[float, float, float, float, str]]) -> Dict[str, List[Tuple[float, float, float, float, str]]]:
    """Assign each word to exactly one box (smallest containing box first) to avoid duplicates."""
    by_field: Dict[str, List[Tuple[float, float, float, float, str]]] = {}
    for name, _ in _field_boxes_sorted():
        by_field[name] = []
    used: set = set()  # word_key(w) tuples already assigned
    for w in extras:
        key = word_key(w)
        if key in used:
            continue
        for name, box in _field_boxes_sorted():
            if in_box(w, box):
                by_field[name].append(w)
                used.add(key)
                break
    return by_field

def _refine_top_with_label_regex(top: Dict[str, str], completed_pdf: Path) -> None:
    """Apply label-anchored regex fallbacks when position-based value is empty or weak."""
    page_text = get_page_text(completed_pdf)
    # For image-only PDFs, get_page_text is empty; build from OCR results in top
    if not page_text:
        parts = []
        for k, v in top.items():
            if v and k != "BOTTOM_TABLE_REGION":
                parts.append(f"{k} {v}")
        page_text = "\n".join(parts)
    if not page_text:
        return
    # Inventory: regex fallback if empty or too short
    inv = top.get("# INVENTAIRE:", "").strip()
    if not inv or len(inv) < 3:
        ref = extract_inventory_by_label(page_text)
        if ref:
            top["# INVENTAIRE:"] = ref
    # Officer 5-digit: regex fallback if empty
    officer = top.get("# INSIGNE AGENT SAISISSANT:", "").strip()
    if not officer or not re.search(r"\d{5}", officer):
        ref = extract_officer_by_label(page_text)
        if ref:
            top["# INSIGNE AGENT SAISISSANT:"] = ref
    # Declaration: fallback if empty/weak.
    declared = _extract_declared_from_top(top)
    if not declared or len(declared) < 2:
        ref = extract_declaration_by_label(page_text)
        if ref:
            # Keep a canonical declaration key populated for downstream mapping.
            top["DÃ‰CLARATION:"] = ref
    # Addresses: regex fallback for multi-line until terminating label
    dest_val = top.get("DESTINATAIRE:", "").strip()
    if not dest_val or (len(dest_val) < 10 and "\n" not in dest_val):
        ref = extract_address_by_label(page_text, "DESTINATAIRE", until_labels=["INDICES:", "EXPÃ‰DITEUR", "LIEU INTERCEPTION", "Notes", "NOTES"])
        if ref:
            top["DESTINATAIRE:"] = ref
    exp_val = top.get("EXPÃ‰DITEUR :", "").strip()
    if not exp_val or (len(exp_val) < 10 and "\n" not in exp_val):
        ref = extract_address_by_label(page_text, "EXPÃ‰DITEUR", until_labels=["INDICES:", "DESTINATAIRE", "LIEU INTERCEPTION", "Notes", "NOTES"])
        if ref:
            top["EXPÃ‰DITEUR :"] = ref
    # Description item (heroine etc.): regex fallback when empty â€“ match both apostrophes
    item_keys = ("DESCRIPTION DE L'ITEM Ã€ SAISIR:", "DESCRIPTION DE L\u2019ITEM Ã€ SAISIR:")
    if not any((top.get(k) or "").strip() for k in item_keys):
        ref = _regex_after_label(
            page_text,
            r"DESCRIPTION\s+DE\s+L['\u2019]ITEM\s+(?:Ã€|À|A|\?)\s+SAISIR\s*:?\s*",
            ["EXPÃ‰DITEUR", "EXPEDITEUR", "DESTINATAIRE", "INDICES", "Notes"],
        )
        if ref:
            for k in item_keys:
                top[k] = ref
    return


def extract_field_values_textdiff(template_pdf: Path, completed_pdf: Path) -> Tuple[Dict[str, str], List[Dict[str, str]]]:
    tmpl_w = words_from_pdf(template_pdf)
    comp_w = words_from_pdf(completed_pdf)
    extras = diff_words(tmpl_w, comp_w)
    by_field = _assign_words_to_boxes(extras)

    top: Dict[str, str] = {}
    for field in FIELD_BOXES:
        if field == "BOTTOM_TABLE_REGION":
            continue
        words = sorted(by_field.get(field, []), key=lambda t: (t[1], t[0]))
        top[field] = join_words(words)

    _refine_top_with_label_regex(top, completed_pdf)

    bt_box = FIELD_BOXES["BOTTOM_TABLE_REGION"]
    bt_words = [w for w in extras if in_box(w, bt_box)]

    rows_by_key: Dict[int, List[Tuple[float, float, float, float, str]]] = {}
    for w in bt_words:
        _, y0, *_ = w
        rk = int(round(y0 / 12.0))
        rows_by_key.setdefault(rk, []).append(w)

    bottom_rows: List[Dict[str, str]] = []
    for rk in sorted(rows_by_key.keys()):
        words = sorted(rows_by_key[rk], key=lambda t: (t[1], t[0]))
        line_txt = " ".join(t for *_, t in words).upper()
        if ("ACHEMINEMENT" in line_txt) or ("DÃ‰PLACEMENT" in line_txt) or ("TÃ‰MOIN" in line_txt):
            continue

        col_vals = [""] * 7
        for x0, y0, x1, y1, t in words:
            col = None
            for i in range(7):
                if BOTTOM_X_BINS[i] <= x0 < BOTTOM_X_BINS[i + 1]:
                    col = i
                    break
            if col is None:
                continue
            col_vals[col] = (col_vals[col] + " " + t).strip()

        if not re.search(r"\d{4}[-/]\d{2}[-/]\d{2}", col_vals[0]):
            continue

        bottom_rows.append({BOTTOM_COLS[i]: clean_value(col_vals[i]) for i in range(7)})

    return top, bottom_rows

# Address fields: preserve line breaks (multi-line as in template)
_ADDRESS_FIELDS = {"DESTINATAIRE:", "EXPÃ‰DITEUR :"}


def extract_field_values_ocr(template_pdf: Path, filled_pdf: Path) -> Tuple[Dict[str, str], List[Dict[str, str]]]:
    """OCR each field box; subtract template text from filled; preserve line breaks for addresses."""
    if not HAVE_PIL or not HAVE_TESSERACT:
        return {}, []
    top: Dict[str, str] = {}
    for field, box in FIELD_BOXES.items():
        if field == "BOTTOM_TABLE_REGION":
            continue
        try:
            img_t = render_clip(template_pdf, box, dpi=300)
            img_f = render_clip(filled_pdf, box, dpi=300)
            ttxt = ocr_image(img_t)
            ftxt = ocr_image(img_f)
            # Preserve line breaks for address fields so output matches template layout
            if field in _ADDRESS_FIELDS:
                top[field] = token_diff_preserve_lines(ttxt, ftxt)
            else:
                top[field] = token_diff(ttxt, ftxt)
        except Exception:
            top[field] = ""
    _refine_top_with_label_regex(top, filled_pdf)
    bottom_rows: List[Dict[str, str]] = []
    return top, bottom_rows

def _blank_top() -> Dict[str, str]:
    top: Dict[str, str] = {}
    for key in FIELD_BOXES:
        if key != "BOTTOM_TABLE_REGION":
            top[key] = ""
    return top

def _build_top_from_text(text: str) -> Dict[str, str]:
    """Build top fields from raw OCR/text content using tolerant label regexes."""
    top = _blank_top()
    if not text:
        return top
    top["# INVENTAIRE:"] = extract_inventory_by_label(text)
    top["# INSIGNE AGENT SAISISSANT:"] = extract_officer_by_label(text)
    top["DESTINATAIRE:"] = extract_address_by_label(
        text,
        "DESTINATAIRE",
        until_labels=["INDICES:", "EXPEDITEUR", "LIEU", "DESTINATAIRE", "NOTES", "Notes"],
    )
    top["DESTINATAIRE:"] = _clean_notice_address(top["DESTINATAIRE:"])
    top["EXP\u00C9DITEUR :"] = _regex_after_label(
        text,
        r"EXP(?:E|Ã‰|\?)DITEUR\??\s*:?\s*",
        ["INDICES", "DESTINATAIRE", "LIEU", "EXPEDITEUR", "NOTES", "Notes"],
    )
    top["SIED #"] = _regex_after_label(text, r"SIED\s*#?\s*:?\s*", ["BOND", "1-AEDS", "LIEU", "INDICES"])
    m = re.search(r"SAISIE\s+.\s*FAIRE\s*[:#]?\s*(\d{2,6})", text, re.IGNORECASE)
    if m:
        top["BOND ROOM LEDGER #"] = m.group(1)
    else:
        top["BOND ROOM LEDGER #"] = _regex_after_label(text, r"(?:BOND\s+ROOM\s+LEDGER|BOND\s+LEDGER)\s*#?\s*:?\s*", ["SIED", "1-AEDS", "LIEU", "INDICES"])
    top["POIDS / QT\u00C9 MARCH.:"] = _regex_after_label(
        text,
        r"POIDS\s*/\s*QT(?:Ã‰|E|\?)\s*MARCH\.?\s*:?\s*",
        ["DESCRIPTION", "EXPEDITEUR", "INDICES"],
    )
    dt = parse_first_date(text)
    if dt:
        y, mo, d = dt
        top["DATE / HEURE INTERCEPTION:"] = f"{y:04d}-{mo:02d}-{d:02d}"
    top["LIEU INTERCEPTION:"] = _regex_after_label(
        text,
        r"LIEU\s+INTERCEPTION\s*:?\s*",
        ["DATE", "DECLARATION", "LIEU INTERCEPTION"],
    )
    top["D\u00C9CLARATION:"] = _regex_after_label(
        text,
        r"D(?:Ã‰|E|\?)CLARATION\s*:?\s*",
        ["POIDS", "DESCRIPTION"],
    )
    if not clean_value(top["D\u00C9CLARATION:"]):
        top["D\u00C9CLARATION:"] = _regex_after_label(
            text,
            r"DECLARED\s*/\s*D(?:Ã‰|É|E|\?)CLAR(?:É|E|\?)\s*:?\s*",
            ["POIDS", "DESCRIPTION", "EXPEDITEUR", "DESTINATAIRE", "INDICES", "NOTES"],
        )
    item_desc = _regex_after_label(
        text,
        r"DESCRIPTION\s+DE\s+L(?:['\u2019]|\?)?ITEM\s+(?:Ã€|À|A|\?)\s+SAISIR\s*:?\s*",
        ["EXPEDITEUR", "DESTINATAIRE", "INDICES", "Notes", "DESCRIPTION DE L"],
    )
    item_desc = _dedupe_lines_keep_order(item_desc)
    top["DESCRIPTION DE L'ITEM \u00C0 SAISIR:"] = item_desc
    top["DESCRIPTION DE L\u2019ITEM \u00C0 SAISIR:"] = item_desc
    return top

def _top_is_effectively_blank(top: Dict[str, str]) -> bool:
    if not top:
        return True
    nonempty = sum(1 for v in top.values() if clean_value(v))
    if nonempty >= 3:
        return False
    key_candidates = [
        "# INVENTAIRE:",
        "# INSIGNE AGENT SAISISSANT:",
        "DESTINATAIRE:",
        "DATE / HEURE INTERCEPTION:",
        "DESCRIPTION DE L'ITEM \u00C0 SAISIR:",
        "DESCRIPTION DE L\u2019ITEM \u00C0 SAISIR:",
    ]
    return not any(clean_value(top.get(k, "")) for k in key_candidates)

def _top_quality_score(top: Dict[str, str]) -> int:
    """Heuristic score for extracted top fields; higher means better quality."""
    if not top:
        return 0
    keys = [
        "DESTINATAIRE:",
        "EXPÃ‰DITEUR :",
        "DESCRIPTION DE L'ITEM Ã€ SAISIR:",
        "DESCRIPTION DE L\u2019ITEM Ã€ SAISIR:",
        "DATE / HEURE INTERCEPTION:",
        "DÃ‰CLARATION:",
        "SIED #",
        "BOND ROOM LEDGER #",
    ]
    score = 0
    for k in keys:
        v = clean_value(top.get(k, ""))
        if not v:
            continue
        score += min(len(v), 80)
        # Bonus when value starts with an alphanumeric (less likely clipped)
        if re.match(r"^[A-Za-z0-9]", v):
            score += 10
        # Bonus for address-like multi-line content
        if k == "DESTINATAIRE:" and ("\n" in v or "|" in v):
            score += 15
    return score

def get_ocr_page_text_pdf(pdf_path: Path, dpi: int = 300) -> str:
    """OCR full first page of PDF; fallback path when field-level OCR is empty."""
    if not HAVE_PIL or not HAVE_TESSERACT:
        return ""
    try:
        doc = fitz.open(str(pdf_path))
        page = doc.load_page(0)
        rect = page.rect
        doc.close()
        img = render_clip(pdf_path, (0, 0, rect.width, rect.height), dpi=dpi)
        return ocr_image(img)
    except Exception:
        return ""

def get_text_from_docx(docx_path: Path) -> str:
    """Extract full text from .docx (paragraphs + tables). Dedupes consecutive duplicate lines."""
    if not HAVE_DOCX:
        return ""
    try:
        doc = DocxDocument(str(docx_path))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        # Dedupe consecutive duplicate lines (tables often repeat structure)
        deduped = []
        for i, line in enumerate(parts):
            if not deduped or line != deduped[-1]:
                deduped.append(line)
        return "\n".join(deduped) or ""
    except Exception:
        pass
    # Fallback: pull all w:t text nodes directly from document.xml (handles some textbox-only docs)
    try:
        with zipfile.ZipFile(str(docx_path), "r") as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        chunks = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml, flags=re.IGNORECASE | re.DOTALL)
        if not chunks:
            return ""
        text = "\n".join(clean_value(html.unescape(c)) for c in chunks if clean_value(html.unescape(c)))
        return text
    except Exception:
        return ""

def get_ocr_text_from_docx_images(docx_path: Path) -> str:
    """OCR embedded images in DOCX (for scanned/image-only documents)."""
    if not HAVE_DOCX or not HAVE_PIL or not HAVE_TESSERACT:
        return ""
    try:
        doc = DocxDocument(str(docx_path))
        out: List[str] = []
        seen_hashes: set = set()
        for rel in doc.part._rels.values():
            if "image" not in rel.reltype:
                continue
            blob = rel.target_part.blob
            if not blob:
                continue
            h = hash(blob[:128])
            if h in seen_hashes:
                continue
            seen_hashes.add(h)
            try:
                img = Image.open(io.BytesIO(blob))
                txt = ocr_image(img)
                if txt:
                    out.append(txt)
            except Exception:
                continue
        return clean_value("\n".join(out))
    except Exception:
        return ""


def extract_field_values_from_docx(docx_path: Path) -> Tuple[Dict[str, str], List[Dict[str, str]], str]:
    """Extract from Word .docx using regex; OCR images when document has no selectable text."""
    text = get_text_from_docx(docx_path)
    if not text:
        text = get_ocr_text_from_docx_images(docx_path)
        if not text:
            return _blank_top(), [], "docx"
        return _build_top_from_text(text), [], "docx-ocr"
    top = _build_top_from_text(text)
    if _top_is_effectively_blank(top):
        ocr_text = get_ocr_text_from_docx_images(docx_path)
        if ocr_text:
            return _build_top_from_text(ocr_text), [], "docx-ocr"
    return top, [], "docx"


def extract_field_values_from_image(image_path: Path) -> Tuple[Dict[str, str], List[Dict[str, str]], str]:
    """Extract from standalone image files via full-image OCR + label regex."""
    text = get_ocr_text_from_image(image_path)
    if not text:
        return _blank_top(), [], "image-ocr"
    return _build_top_from_text(text), [], "image-ocr"


def _regex_after_label(text: str, label_pat: str, until: List[str]) -> str:
    """Capture text after label until one of the until patterns."""
    m = re.search(label_pat + r"(.*?)(?=" + "|".join(re.escape(u) for u in until) + r"|\Z)", text, re.IGNORECASE | re.DOTALL)
    return clean_value(m.group(1)) if m and m.lastindex else ""


def should_use_ocr(pdf_path: Path, word_count: int) -> bool:
    """
    Determine if PDF should use OCR instead of text-diff.
    Returns True if OCR is recommended.
    
    Detection methods:
    - Environment variable FORCE_OCR=1 (demo/testing override)
    - Very low word count (< 20)
    - Image-based PDF detection (images present + few text blocks)
    - Text quality check (high single-char ratio = OCR artifacts)
    """
    # DEMO/TESTING: Force OCR via environment variable (no code changes needed)
    if os.environ.get("FORCE_OCR") == "1":
        return True
    
    # Very few words = definitely OCR
    if word_count < 20:
        return True
    
    # If many words, check if it's actually image-based
    try:
        doc = fitz.open(str(pdf_path))
        page = doc[0]
        
        # Check image content
        images = page.get_images()
        text_dict = page.get_text("dict")
        text_blocks = [b for b in text_dict.get("blocks", []) if "lines" in b]
        
        # Heuristic: If has images AND few text blocks, likely scanned/image PDF
        if len(images) > 3 and len(text_blocks) < 10:
            doc.close()
            return True
        
        # Check text quality: if many single-char "words", likely OCR artifact
        words = page.get_text("words")
        if len(words) > 0:
            single_char_ratio = sum(1 for w in words if len(w[4]) == 1) / len(words)
            if single_char_ratio > 0.4:  # 40% single chars = likely OCR'd
                doc.close()
                return True
        
        doc.close()
    except Exception:
        pass
    
    # Default: use text-diff if word count is high enough
    return word_count < 80


def extract_field_values(template_pdf: Optional[Path], completed_pdf: Path) -> Tuple[Dict[str, str], List[Dict[str, str]], str]:
    """Use text-diff/OCR for PDFs, regex for DOCX, and OCR+regex for image files."""
    suffix = (completed_pdf.suffix or "").lower()
    if suffix == ".docx" and HAVE_DOCX:
        return extract_field_values_from_docx(completed_pdf)
    if suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}:
        return extract_field_values_from_image(completed_pdf)
    if suffix != ".pdf":
        return _blank_top(), [], f"unsupported:{suffix or 'none'}"
    if template_pdf is None:
        if HAVE_PIL and HAVE_TESSERACT:
            txt = get_ocr_page_text_pdf(completed_pdf)
            if txt:
                return _build_top_from_text(txt), [], "ocr-no-template"
        return _blank_top(), [], "text-diff-no-template"
    
    # Get word count
    fill_words = words_from_pdf(completed_pdf) if suffix == ".pdf" else []
    word_count = len(fill_words)
    
    # Use better detection (includes FORCE_OCR env var for demo)
    use_ocr = should_use_ocr(completed_pdf, word_count)
    
    # Try OCR if detection says so AND OCR libraries are available
    if use_ocr:
        if HAVE_PIL and HAVE_TESSERACT:
            top, bottom = extract_field_values_ocr(template_pdf, completed_pdf)
            page_text = get_ocr_page_text_pdf(completed_pdf)
            if page_text:
                top_page = _build_top_from_text(page_text)
                score_box = _top_quality_score(top)
                score_page = _top_quality_score(top_page)
                if _top_is_effectively_blank(top) or (score_page > score_box + 15):
                    return top_page, bottom, "ocr-page-fallback"
            return top, bottom, "ocr"
        else:
            # OCR needed but not available - warn and fall back to text-diff
            # (In production, you might want to raise an error here)
            import warnings
            if not HAVE_TESSERACT:
                warnings.warn(f"OCR mode needed but Tesseract not available. Install: pip install pytesseract (and Tesseract OCR). Falling back to text-diff.")
            top, bottom = extract_field_values_textdiff(template_pdf, completed_pdf)
            return top, bottom, "text-diff-fallback"
    
    # Default to text-diff
    top, bottom = extract_field_values_textdiff(template_pdf, completed_pdf)
    return top, bottom, "text-diff"


# ======================== =
# CSV writers -----
# ======================== =

def _open_csv_append_or_create(path: Path):
    """
    Create CSV with UTF-8 BOM for Excel compatibility, append later as UTF-8.
    """
    if path.exists():
        return open(path, "a", newline="", encoding="utf-8")
    return open(path, "w", newline="", encoding="utf-8-sig")


def write_saisie_csv(out_csv: Path, top: Dict[str, str], bottom_rows: List[Dict[str, str]]) -> None:
    with open(out_csv, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.writer(f)
        w.writerow([normalize_output_text("SAISIE Ã€ FAIRE"), ""])
        for fn in TOP_CSV_FIELDS:
            w.writerow([normalize_output_text(fn), normalize_output_text(top.get(fn, ""))])
        w.writerow(["", ""])
        w.writerow([normalize_output_text("ACHEMINEMENT DES PIÃˆCES Ã€ CONVICTION"), ""])
        for i, r in enumerate(bottom_rows, start=1):
            if i > 1:
                w.writerow(["", ""])
            for col in BOTTOM_COLS:
                w.writerow([normalize_output_text(col), normalize_output_text(r.get(col, ""))])

def detect_form_type(item_description: str) -> str:
    """
    Detect form type from item description.
    Returns: "Cannabis-Stupefiant", "Knives-Arms", or "Stupefiant-Others"
    """
    desc_lower = (item_description or "").lower()
    
    # Check for cannabis-related keywords
    cannabis_keywords = ["cannabis", "marijuana", "marihuana", "weed", "thc", "cbd", "hash", "hashish"]
    if any(keyword in desc_lower for keyword in cannabis_keywords):
        return "Cannabis-Stupefiant"
    
    # Check for knives/arms-related keywords
    knives_keywords = ["knife", "couteau", "centrifugal", "blade", "arm", "weapon", "arme"]
    if any(keyword in desc_lower for keyword in knives_keywords):
        return "Knives-Arms"
    
    # Default
    return "Stupefiant-Others"


def _find_notice_start(s: str) -> int:
    """First character index of any legal notice in s, or -1. Used to strip notice from item text."""
    if not s:
        return -1
    s_lower = s.lower()
    markers = [
        "(please note:",
        "please note:",
        "(attention :",
        "attention :",
        "centrifugal knife",
        "couteau centrifuge",
    ]
    first = len(s)
    for m in markers:
        i = s_lower.find(m)
        if i != -1:
            if i > 0 and s[i - 1] == "(" and not m.startswith("("):
                i = i - 1
            first = min(first, i)
    return first if first < len(s) else -1


def _extract_item_by_label(text: str) -> str:
    """Extract DESCRIPTION DE L'ITEM Ã€ SAISIR value from full-page text."""
    if not text:
        return ""
    # Flexible for apostrophe/accents/OCR substitutions.
    pattern = (
        r"DESCRIPTION\s+DE\s+L(?:['\u2019]|\?)?ITEM\s+(?:Ã€|À|A|\?)\s+SAISIR\s*:?\s*"
        r"(.*?)(?=\s*(?:EXP(?:Ã‰|É|\ufffd|E|\?)DITEUR|DESTINATAIRE|INDICES|NOTES?|POIDS|DECLARATION|DÃ‰CLARATION|LIEU\s+INTERCEPTION|DATE\s*/\s*HEURE)\b\s*:?\s*|\Z)"
    )
    m = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
    if not m:
        return ""
    block = m.group(1)
    lines = [clean_value(ln) for ln in block.splitlines() if clean_value(ln)]
    return clean_value("\n".join(lines))


def _strip_address_leak_from_item(s: str) -> str:
    """
    Remove non-item section leakage that can appear after item text
    (e.g., EXPEDITEUR/DESTINATAIRE blocks).
    """
    if not s:
        return ""
    out = str(s)
    # Detect EXPÉDITEUR at start in any encoding (mojibake, replacement char, etc.)
    alpha_start = re.sub(r"[^A-Za-z]", "", out[:20]).upper()
    if alpha_start.startswith("EXPDITEUR"):
        return ""
    m = re.search(
        r"\b(?:EXP(?:Ã‰|É|\ufffd|E|\?)DITEUR|DESTINATAIRE|INDICES|LIEU\s+INTERCEPTION|DATE\s*/\s*HEURE)\b\s*:?",
        out,
        re.IGNORECASE,
    )
    if m:
        out = out[:m.start()]
    lines = [clean_value(ln) for ln in out.splitlines() if clean_value(ln)]
    # Common residue when cutting "INCONNU DESTINATAIRE: ..."
    while lines and re.fullmatch(r"(?:INCONNU|UNKNOWN|N/?A)", lines[0], re.IGNORECASE):
        lines.pop(0)
    while lines and re.fullmatch(r"(?:INCONNU|UNKNOWN|N/?A)", lines[-1], re.IGNORECASE):
        lines.pop()
    return clean_value("\n".join(lines))


def _page_lines(text: str) -> List[str]:
    return [clean_value(ln) for ln in (text or "").splitlines() if clean_value(ln)]


def _is_unknown_notice_address_block(s: str) -> bool:
    """
    True when DESTINATAIRE content is only placeholders (e.g., INCONNU/UNKNOWN/N-A)
    and contains no real address text.
    """
    raw = clean_value(s or "")
    if not raw:
        return False

    normalized = _strip_repeated_label_blocks(raw, r"DESTINATAIRE\s*:?\s*")
    lines = [clean_value(ln) for ln in normalized.splitlines() if clean_value(ln)]
    if not lines:
        return True

    unknown_tokens = {"INCONNU", "UNKNOWN", "NA", "N/A", "ND", "NONDISPONIBLE"}
    for ln in lines:
        ln = re.sub(r"^\s*DESTINATAIRE\s*:?\s*", "", ln, flags=re.IGNORECASE)
        ln = clean_value(ln).strip(" _-:/")
        if not ln:
            continue
        token = re.sub(r"[^A-Za-z0-9/]+", "", ln).upper()
        if token in unknown_tokens:
            continue
        if re.fullmatch(r"(?:INCONNU|UNKNOWN|N/?A)(?:[ _-]+(?:INCONNU|UNKNOWN|N/?A))*", ln, flags=re.IGNORECASE):
            continue
        return False
    return True


def _is_weak_notice_address(s: str) -> bool:
    """True when destination block is too short/incomplete for K138 notice."""
    lines = [clean_value(ln) for ln in (s or "").splitlines() if clean_value(ln)]
    if len(lines) < 3:
        return True
    if not any(re.search(r"\d", ln) for ln in lines):
        return True
    has_postal = any(_RE_POSTAL_CA.search(ln) or _RE_ZIP_US.search(ln) or _RE_POSTCODE_UK.search(ln) for ln in lines)
    has_country = any(_canonical_country_name(ln) for ln in lines)
    return not (has_postal or has_country)


def _is_suspicious_notice_address(s: str) -> bool:
    """
    Detect likely-clipped destination blocks (missing first characters, bad country token, etc.).
    """
    lines = [clean_value(ln) for ln in (s or "").splitlines() if clean_value(ln)]
    if not lines:
        return True

    # Lines starting with lowercase letters are usually OCR clipping artifacts.
    if any(re.match(r"^[a-z]", ln) for ln in lines if re.search(r"[A-Za-z]", ln)):
        return True

    # Common first-letter clipping on city/country.
    if any(re.match(r"^(?:ort\b|nited\b|anada\b|xpediteur\b|estinataire\b)", ln, re.IGNORECASE) for ln in lines):
        return True

    # Country line variants close to "USA" produced by OCR.
    last = lines[-1].upper().replace(" ", "")
    if last in {"JSA", "ISA", "U5A"}:
        return True

    # US ZIP should be 5 digits (or ZIP+4), not 4.
    if any(re.fullmatch(r"\d{4}", ln) for ln in lines):
        return True
    if any(re.search(r"\bEXP(?:Ã‰|É|E)DITEUR\b", ln, re.IGNORECASE) for ln in lines):
        return True

    return False


def _is_weak_inventory_token(s: str) -> bool:
    """True when value does not look like an inventory id (letters+digits, min length)."""
    return not bool(_normalize_inventory_number(s))


def _is_weak_item_text(s: str) -> bool:
    """True when extracted item text is empty or clearly a leaked label/address block."""
    v = clean_value(s or "")
    if not v:
        return True
    if re.fullmatch(r"(?:INCONNU|UNKNOWN|N/?A|_+)", v, re.IGNORECASE):
        return True
    if re.match(r"^(?:INCONNU|UNKNOWN|N/?A)\b", v, re.IGNORECASE):
        return True
    if re.match(r"^(?:EXP(?:Ã‰|E|\?)DITEUR|DESTINATAIRE|INDICES|NOTES?|D(?:Ã‰|E)CLARATION)\b", v, re.IGNORECASE):
        return True
    if re.search(r"\b(?:DESTINATAIRE|EXP(?:Ã‰|É|\ufffd|E|\?)DITEUR|INDICES)\b", v, re.IGNORECASE):
        return True
    # Detect EXPÉDITEUR in any encoding (mojibake, replacement char) by stripping non-alpha
    alpha_start = re.sub(r"[^A-Za-z]", "", v[:20]).upper()
    if alpha_start.startswith("EXPDITEUR"):
        return True
    return False


def _extract_inventory_from_text_loose(text: str) -> str:
    """
    Heuristic fallback for inventory token when label extraction fails.
    Looks for first alnum token containing both letters and digits.
    """
    if not text:
        return ""
    bad = {
        "CIBLAGE", "CLOTHES", "CANADA", "INCONNU", "BELANGER",
        "SAISIE", "CONFISCATION", "CHECKLIST", "DESTINATAIRE",
        "EXPEDITEUR", "INDICES",
    }
    # Prefer lines around explicit inventory labels.
    for ln in _page_lines(text):
        if re.search(r"\bINVENTAIRE\b", ln, re.IGNORECASE):
            norm = _normalize_inventory_number(ln)
            if norm:
                return norm

    # Prefer explicit legacy inventory style (W + long digits), allowing OCR V/W confusion.
    for m in re.finditer(r"\b[WV]\s*(?:\d[\s-]*){10,24}\b", text, re.IGNORECASE):
        norm = _normalize_inventory_number(m.group(0))
        if norm:
            return norm

    # Spaced inventory style fallback, e.g., "W 0000 6042 0000 43012".
    for m in re.finditer(r"\b[A-Za-z]\s*(?:\d[\s-]*){6,30}\b", text):
        norm = _normalize_inventory_number(m.group(0))
        if norm:
            return norm

    for tok in re.findall(r"\b[A-Za-z0-9-]{8,40}\b", text):
        if not (re.search(r"[A-Za-z]", tok) and re.search(r"\d", tok)):
            continue
        if tok.upper() in bad:
            continue
        norm = _normalize_inventory_number(tok)
        if norm:
            return norm
    return ""


def _extract_item_from_text_loose(text: str) -> str:
    """
    Heuristic fallback for item: often appears immediately after weight line.
    """
    lines = _page_lines(text)
    if not lines:
        return ""
    skip = re.compile(
        r"(?:UNDER LAYER|CLOTHES|DECLARATION|EXP(?:Ã‰|E|\?)DITEUR|DESTINATAIRE|INDICES|CHECKLIST|CANADA|USA|UK|INCONNU)",
        re.IGNORECASE,
    )
    for i, line in enumerate(lines):
        if re.search(r"\b\d+(?:[.,]\d+)?\s*(?:KG|G|GRAM|GRAMME|LBS|LB)\b", line, re.IGNORECASE):
            for j in range(i + 1, min(i + 4, len(lines))):
                cand = lines[j]
                if ":" in cand:
                    continue
                if skip.search(cand):
                    continue
                if 2 <= len(cand) <= 60:
                    return cand
    return ""


def _extract_destination_from_text_loose(text: str, seed: str = "") -> str:
    """
    Heuristic fallback for destination address from page text.
    """
    lines = _page_lines(text)
    if not lines:
        return ""

    def _clean_block(block_lines: List[str]) -> str:
        stop_re = re.compile(
            r"(?:DATE\s*/\s*HEURE|LIEU\s+INTERCEPTION|D(?:Ã‰|E|É)CLARATION|POIDS|INDICES|CHECKLIST|NOTES?|"
            r"EXP(?:Ã‰|É|E)DITEUR|"
            r"ACHEMINEMENT|BOND\s+ROOM|SIED|#\s*INVENTAIRE|#\s*INSIGNE)",
            re.IGNORECASE,
        )
        out: List[str] = []
        for ln in block_lines:
            c = _clean_notice_artifact_line(ln)
            if not c:
                continue
            if re.match(r"^(?:DESTINATAIRE|EXP(?:Ã‰|É|E)DITEUR)\b", c, re.IGNORECASE):
                continue
            if stop_re.search(c):
                break
            # Remove label/placeholder noise.
            if re.fullmatch(r"DESTINATAIRE\s*:?", c, re.IGNORECASE):
                continue
            if re.fullmatch(r"INCONNU|UNKNOWN|_+", c, re.IGNORECASE):
                continue
            if re.fullmatch(r"\d{4}[-/]\d{2}[-/]\d{2}(?:\s+\d{1,2}:\d{2})?", c):
                continue
            if re.fullmatch(r"\d{1,2}:\d{2}", c):
                continue
            out.append(c)
            # Address usually ends when we reach country line.
            if c.upper() in {"CANADA", "USA", "UK"}:
                break
        return _normalize_notice_address_layout("\n".join(out))

    # If we already have a partial destination line, expand around that anchor.
    anchor = ""
    for ln in _page_lines(seed):
        if len(ln) >= 6:
            anchor = ln
            break
    if anchor:
        for i, ln in enumerate(lines):
            if anchor.lower() in ln.lower() or ln.lower() in anchor.lower():
                start = max(0, i - 2)
                end = min(len(lines), i + 6)
                block_txt = _clean_block(lines[start:end])
                if re.search(r"\d", block_txt):
                    return block_txt

    # Canadian postal format fallback (e.g., L7A 2S2).
    for i, ln in enumerate(lines):
        if re.search(r"\b[A-Z]\d[A-Z]\s*\d[A-Z]\d\b", ln.upper()):
            start = max(0, i - 3)
            end = min(len(lines), i + 4)
            block_txt = _clean_block(lines[start:end])
            if re.search(r"\d", block_txt):
                return block_txt

    return ""


def _top_blob(top: Dict[str, str]) -> str:
    parts: List[str] = []
    for k, v in (top or {}).items():
        if k:
            parts.append(str(k))
        if v:
            parts.append(str(v))
    return "\n".join(parts)


def _fallback_inventory_from_top(top: Dict[str, str]) -> str:
    """
    Conservative fallback: only use values attached to explicit inventory keys.
    Avoid scanning all fields, which can pick tracking numbers from other blocks.
    """
    if not top:
        return ""
    inventory_lines: List[str] = []
    for k, v in top.items():
        kn = re.sub(r"[^A-Za-z0-9]+", "", str(k or "")).upper()
        if "INVENTAIRE" in kn or "INVENTORY" in kn:
            cv = clean_value(v or "")
            if cv:
                inventory_lines.append(cv)
    for line in inventory_lines:
        by_label = extract_inventory_by_label(line)
        if by_label:
            return by_label
        norm = _normalize_inventory_number(line)
        if norm:
            return norm
    return ""


def _fallback_officer_from_top(top: Dict[str, str]) -> str:
    blob = _top_blob(top)
    if not blob:
        return ""
    # Prefer number near officer label
    m = re.search(r"(?:INSIGNE|AGENT\s+SAISISSANT|OFFICER)\s*[:#]?\s*(\d{4,6})", blob, re.IGNORECASE)
    if m:
        d = m.group(1)
        return d[-5:] if len(d) > 5 else d.zfill(5) if len(d) < 5 else d
    # Fallback: first 5-digit number
    m = re.search(r"\b(\d{5})\b", blob)
    if m:
        return m.group(1)
    return ""


_BOX_FALLBACK_PAD: Dict[str, Tuple[float, float, float, float]] = {
    # x0, y0, x1, y1 padding
    "# INVENTAIRE:": (-2, 0, 2, 3),
    "DESTINATAIRE:": (-2, 0, 2, 2),
    "DESCRIPTION DE Lâ€™ITEM Ã€ SAISIR:": (-2, 0, 2, 2),
    "DÃ‰CLARATION:": (-2, 0, 2, 2),
    "# INSIGNE AGENT SAISISSANT:": (-2, 0, 2, 2),
}


def _word_overlaps_box(word: Tuple[float, float, float, float, str], box: Tuple[float, float, float, float]) -> bool:
    """True when the word center falls inside the box rectangle."""
    x0, y0, x1, y1, _ = word
    bx0, by0, bx1, by1 = box
    cx = (x0 + x1) / 2.0
    cy = (y0 + y1) / 2.0
    return (bx0 <= cx <= bx1) and (by0 <= cy <= by1)


def _extract_box_text_from_pdf(completed_pdf: Path, field_name: str) -> str:
    """
    Read text directly from a field box in the completed PDF (without template diff).
    Useful as fallback when text-diff misses values.
    """
    box = FIELD_BOXES.get(field_name)
    if not box:
        return ""
    px0, py0, px1, py1 = _BOX_FALLBACK_PAD.get(field_name, (0, 0, 0, 0))
    box = (box[0] + px0, box[1] + py0, box[2] + px1, box[3] + py1)
    try:
        words = words_from_pdf(completed_pdf)
        in_field = [w for w in words if _word_overlaps_box(w, box)]
        in_field = sorted(in_field, key=lambda t: (t[1], t[0]))
        return join_words(in_field)
    except Exception:
        return ""


def build_k138_values_from_saisie(top: Dict[str, str], completed_pdf: Path, form_type: str = None, notice_text: str = "") -> Dict[str, str]:
    # Map SAISIE fields â†’ k138_values.csv fields (based on your sample file format)
    d: Dict[str, str] = {}
    from datetime import date
    page_text = get_page_text(completed_pdf)

    # notice_to: strict extraction from DESTINATAIRE only.
    # Never borrow from full-page heuristics (can leak date/other section text).
    raw_dest = safe_get(top, "DESTINATAIRE:")
    dest_is_unknown = _is_unknown_notice_address_block(raw_dest)
    dest_block = "" if dest_is_unknown else _clean_notice_address(raw_dest)

    # Fallback limited to DESTINATAIRE box only (same section, no cross-section scanning).
    if (not dest_is_unknown) and _is_weak_notice_address(dest_block):
        box_dest = _extract_box_text_from_pdf(completed_pdf, "DESTINATAIRE:")
        if box_dest:
            if _is_unknown_notice_address_block(box_dest):
                dest_is_unknown = True
                dest_block = ""
            else:
                dest_block = _clean_notice_address(box_dest)

    if dest_is_unknown:
        dest_block = ""
    notice_to_pipe = normalize_multiline_to_pipe(dest_block)
    # Keep K138 generation unblocked when DESTINATAIRE is blank/unknown on SAISIE.
    if not clean_value(notice_to_pipe):
        notice_to_pipe = "INCONNU"
    d["notice_to"] = notice_to_pipe

    # notice_date: date the letter was generated (Avis de saisie) â€“ today, not seizure date
    today = date.today()
    d["notice_date"] = f"{today.year:04d}-{today.month:02d}-{today.day:02d}"

    # seizure_date_line + split year (left/right) â€“ from DATE/HEURE INTERCEPTION
    # Prefer label-anchored extraction from full page text when available.
    dt = None
    if page_text:
        dt = extract_interception_date_by_label(page_text)
    if not dt:
        dt = parse_first_date(safe_get(top, "DATE / HEURE INTERCEPTION:"))
    if not dt and page_text:
        dt = parse_first_date(page_text)
    if dt:
        yyyy, mm, dd = dt
        en, fr = month_en_fr(mm)
        d["seizure_date_line"] = f"{dd} {en} / {dd} {fr}"
        d["seizure_year_left"] = str(yyyy)[:2]
        d["seizure_year_right"] = str(yyyy)[2:]
    else:
        d["seizure_date_line"] = ""
        d["seizure_year_left"] = ""
        d["seizure_year_right"] = ""

    # seizure_location: ALWAYS fixed value for K138 form (Dmitry)
    d["seizure_location"] = K138_SEIZURE_LOCATION_FIXED
    # lieu_interception: extracted value, saved for hidden folder analysis (not used on form)
    lieu = safe_get(top, "LIEU INTERCEPTION:")
    # Strip repeated "Ciblage LIEU INTERCEPTION:" from table-style extraction
    lieu_clean = _strip_repeated_label_blocks(lieu, r"Ciblage\s+LIEU\s+INTERCEPTION\s*:?\s*")
    d["lieu_interception"] = lieu_clean if lieu_clean else lieu

    # description_inventory: no spaces (e.g. AB123456789CA)
    raw_inv = safe_get(top, "# INVENTAIRE:")
    inv_from_explicit_source = bool(clean_value(raw_inv))
    if not raw_inv:
        raw_inv = _extract_box_text_from_pdf(completed_pdf, "# INVENTAIRE:")
        inv_from_explicit_source = bool(clean_value(raw_inv))
    if not raw_inv:
        raw_inv = _fallback_inventory_from_top(top)
        inv_from_explicit_source = bool(clean_value(raw_inv))

    # description_declared: keep this as a single line for K138 rendering.
    declared_single = _extract_declared_from_top(top)
    if not declared_single:
        box_declared = _extract_box_text_from_pdf(completed_pdf, "DÃ‰CLARATION:")
        if box_declared:
            declared_single = _normalize_declared_text(box_declared)
    if not declared_single and page_text:
        declared_single = extract_declaration_by_label(page_text)
    if not declared_single and page_text:
        parsed_top = _build_top_from_text(page_text)
        declared_single = _extract_declared_from_top(parsed_top)
    if (
        not declared_single
        and (completed_pdf.suffix or "").lower() == ".pdf"
        and HAVE_PIL
        and HAVE_TESSERACT
    ):
        page_text_ocr = get_ocr_page_text_pdf(completed_pdf)
        if page_text_ocr:
            declared_single = extract_declaration_by_label(page_text_ocr)
            if not declared_single:
                parsed_top_ocr = _build_top_from_text(page_text_ocr)
                declared_single = _extract_declared_from_top(parsed_top_ocr)
    declared_single = re.sub(r"\s+", " ", (declared_single or "").replace("|", " ")).strip()
    d["description_declared"] = declared_single

    if raw_inv:
        by_label = extract_inventory_by_label(raw_inv)
        raw_inv = by_label if by_label else raw_inv
    if _is_weak_inventory_token(raw_inv) and page_text:
        raw_inv = extract_inventory_by_label(page_text)
    if _is_weak_inventory_token(raw_inv) and page_text:
        raw_inv = _extract_inventory_from_text_label_only(page_text)
    if (
        _is_weak_inventory_token(raw_inv)
        and inv_from_explicit_source
        and (completed_pdf.suffix or "").lower() == ".pdf"
        and HAVE_PIL
        and HAVE_TESSERACT
    ):
        page_text_ocr = get_ocr_page_text_pdf(completed_pdf)
        if page_text_ocr:
            raw_inv = extract_inventory_by_label(page_text_ocr) or _extract_inventory_from_text_label_only(page_text_ocr)
    if _is_weak_inventory_token(raw_inv):
        # Do not loose-scan the whole document for inventory:
        # this can pull postal tracking values when the inventory field is blank.
        raw_inv = _normalize_inventory_number(raw_inv) if inv_from_explicit_source else ""
    d["description_inventory"] = _normalize_inventory_number(raw_inv) if raw_inv else ""

    # description_item: only the actual item text (no legal notice here â€“ notice is added once in fill_k138 via legal_notice)
    # First try direct key variants, then tolerant key match.
    # Scan all keys for the DESCRIPTION DE L*ITEM*SAISIR field using alpha-normalized
    # key matching (avoids picking up wrong straight-apostrophe key with EXPEDITEUR value).
    item_desc = ""
    for _ik, _iv in top.items():
        _ik_alpha = re.sub(r"[^A-Za-z]", "", _ik).upper()
        if _ik_alpha in ("DESCRIPTIONDELITEMASAISIR", "DESCRIPTIONDELITEMSAISIR"):
            _cv = clean_value(_iv)
            if _cv and not re.sub(r"[^A-Za-z]", "", _cv[:20]).upper().startswith("EXPDITEUR"):
                item_desc = _cv
                break
    if not item_desc:
        item_desc = (
            safe_get(top, "DESCRIPTION DE L\u2019ITEM \u00c0 SAISIR:")
            or safe_get(top, "DESCRIPTION DE L'ITEM \u00c0 SAISIR:")
            or safe_get(top, "DESCRIPTION DE L\u2019ITEM \u00c3\u20ac SAISIR:")
            or safe_get(top, "DESCRIPTION DE L'ITEM \u00c3\u20ac SAISIR:")
        )
    # Discard if value is EXPEDITEUR content (misextracted due to PDF layout artifacts)
    if re.sub(r"[^A-Za-z]", "", (item_desc or "")[:20]).upper().startswith("EXPDITEUR"):
        item_desc = ""
    if not item_desc:
        item_desc = _top_first_match(top, ["DESCRIPTION", "ITEM", "SAISIR"])
    # Strip repeated "DESCRIPTION DE L'ITEM Ã€ SAISIR:" blocks from table-style DOCX extraction
    item_desc = _strip_repeated_label_blocks(
        item_desc,
        r"DESCRIPTION\s+DE\s+L['\u2019]?ITEM\s+(?:Ã€|À|A|\?)\s+SAISIR\s*:?\s*",
    )
    if not item_desc:
        item_desc = (
            safe_get(top, "DESCRIPTION DE L\u2019ITEM À SAISIR:")
            or safe_get(top, "DESCRIPTION DE L'ITEM À SAISIR:")
            or safe_get(top, "DESCRIPTION DE L\u2019ITEM Ã€ SAISIR:")
            or safe_get(top, "DESCRIPTION DE L'ITEM Ã€ SAISIR:")
        )
    if not item_desc:
        item_desc = _extract_item_by_label(page_text)
    if not item_desc and page_text:
        parsed_top = _build_top_from_text(page_text)
        item_desc = (
            safe_get(parsed_top, "DESCRIPTION DE L\u2019ITEM À SAISIR:")
            or safe_get(parsed_top, "DESCRIPTION DE L'ITEM À SAISIR:")
            or safe_get(parsed_top, "DESCRIPTION DE L\u2019ITEM Ã€ SAISIR:")
            or safe_get(parsed_top, "DESCRIPTION DE L'ITEM Ã€ SAISIR:")
            or _top_first_match(parsed_top, ["DESCRIPTION", "ITEM", "SAISIR"])
        )
    if _is_weak_item_text(item_desc):
        item_desc = ""
    if not item_desc:
        box_item = _extract_box_text_from_pdf(completed_pdf, "DESCRIPTION DE L\u2019ITEM Ã€ SAISIR:")
        if box_item:
            item_desc = _strip_repeated_label_blocks(
                box_item,
                r"DESCRIPTION\s+DE\s+L['\u2019]?ITEM\s+(?:Ã€|À|A|\?)\s+SAISIR\s*:?\s*"
            )
    if _is_weak_item_text(item_desc):
        item_desc = ""
    if not item_desc and page_text:
        item_desc = _extract_item_from_text_loose(page_text)
    # Strip any legal notice block from extracted text (so it never appears inside description_item)
    first = _find_notice_start(item_desc)
    if first != -1:
        item_desc = item_desc[:first].rstrip()
    # Guardrail: remove leaked sender/recipient/address block from item.
    item_desc = _strip_address_leak_from_item(item_desc)
    if _is_weak_item_text(item_desc):
        # Prefer declaration as fallback rather than leaking sender/recipient labels.
        item_desc = declared_single or ""
    # Store full description for Narrative; K138 uses a shorter version.
    d["description_item_full"] = item_desc
    # SHORT: truncate at 150 chars on a word boundary for K138 form field.
    _short_limit = 150
    if len(item_desc) > _short_limit:
        truncated = item_desc[:_short_limit]
        last_space = truncated.rfind(" ")
        item_desc = (truncated[:last_space] if last_space > 80 else truncated).rstrip(" ,;") + "…"
    d["description_item"] = item_desc

    # legal_notice: passed separately so fill_k138 can append it exactly once (no duplication)
    d["legal_notice"] = (notice_text or "").strip() if notice_text else ""

    # Form type: use provided form_type if given, otherwise detect from item description
    if form_type:
        d["form_type"] = form_type
    else:
        d["form_type"] = detect_form_type(item_desc)

    # description_seizure_number: extract a clean number from SIED/BOND fields
    d["description_seizure_number"] = ""

    # seizing_officer: map to K138; normalize to 5 digits (Dmitry: officer 5-digit from # INSIGNE AGENT SAISISSANT)
    raw = safe_get(top, "# INSIGNE AGENT SAISISSANT:")
    if not raw:
        box_officer = _extract_box_text_from_pdf(completed_pdf, "# INSIGNE AGENT SAISISSANT:")
        if box_officer:
            m = re.search(r"\b(\d{4,6})\b", box_officer)
            raw = m.group(1) if m else box_officer
    if not raw:
        raw = _top_first_match(top, ["INSIGNE", "SAISISSANT"])
    if not raw:
        raw = _fallback_officer_from_top(top)
    if not raw and page_text:
        raw = extract_officer_by_label(page_text)
    if not raw and page_text:
        m = re.search(r"\b(\d{5})\b", page_text)
        if m:
            raw = m.group(1)
    digits = re.sub(r"\D", "", raw)
    if len(digits) > 5:
        d["seizing_officer"] = digits[-5:]
    elif len(digits) > 0:
        d["seizing_officer"] = digits.zfill(5)
    else:
        d["seizing_officer"] = raw

    return d

def write_k138_values_csv(out_csv: Path, values: Dict[str, str]) -> None:
    # exactly like your sample: header row "field,value"
    with open(out_csv, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.writer(f)
        w.writerow(["field", "value"])
        for k in [
            "notice_to",
            "notice_date",
            "seizure_date_line",
            "seizure_year_left",
            "seizure_year_right",
            "seizure_location",
            "lieu_interception",
            "description_inventory",
            "description_declared",
            "description_item",
            "description_seizure_number",
            "legal_notice",
            "seizing_officer",
            "form_type",
        ]:
            w.writerow([k, normalize_output_text(values.get(k, ""))])


def compute_field_confidence(top: Dict[str, str], values: Dict[str, str]) -> Dict[str, int]:
    """Return field-level confidence scores (0-100)."""
    scores: Dict[str, int] = {}

    notice_to = values.get("notice_to", "")
    scores["confidence_notice_to"] = 95 if ("|" in notice_to and _safe_len(notice_to) >= 20) else (70 if _safe_len(notice_to) >= 12 else 25)

    inv = re.sub(r"\s+", "", values.get("description_inventory", "")).upper()
    scores["confidence_inventory"] = 95 if re.fullmatch(r"[A-Z0-9\-]{8,30}", inv or "") else (65 if _safe_len(inv) >= 5 else 20)

    officer = re.sub(r"\D", "", values.get("seizing_officer", ""))
    scores["confidence_officer"] = 95 if len(officer) == 5 else (70 if len(officer) >= 4 else 20)

    item = values.get("description_item", "")
    scores["confidence_item"] = 90 if _safe_len(item) >= 10 else (70 if _safe_len(item) >= 3 else 20)

    seizure_no = values.get("description_seizure_number", "")
    # Seizure number is intentionally blank for current workflow.
    if not clean_value(seizure_no):
        scores["confidence_seizure_number"] = 90
    else:
        scores["confidence_seizure_number"] = 90 if re.search(r"\d", seizure_no) else 65

    declared = values.get("description_declared", "")
    scores["confidence_declared"] = 85 if _safe_len(declared) >= 2 else 35

    return scores


def validate_k138_values(values: Dict[str, str]) -> Tuple[List[str], List[str]]:
    """Validation rules before fill. Returns (errors, warnings)."""
    errors: List[str] = []
    warnings: List[str] = []

    for field in CRITICAL_K138_FIELDS:
        if not clean_value(values.get(field, "")):
            errors.append(f"Missing critical field: {field}")

    inv = re.sub(r"\s+", "", values.get("description_inventory", "")).upper()
    if inv and not re.fullmatch(r"[A-Z0-9\-]{6,35}", inv):
        warnings.append("Inventory format looks unusual")

    officer_digits = re.sub(r"\D", "", values.get("seizing_officer", ""))
    if officer_digits and len(officer_digits) != 5:
        warnings.append("Seizing officer should usually be 5 digits")

    if _safe_len(values.get("notice_to", "")) < 10:
        warnings.append("Address block (notice_to) is very short")

    seiz = values.get("description_seizure_number", "") or ""
    if clean_value(seiz) and (not re.search(r"\d", seiz)):
        warnings.append("Seizure number looks unusual")

    return errors, warnings


def _build_case_key(values: Dict[str, str]) -> str:
    inv = re.sub(r"\W+", "", (values.get("description_inventory", "") or "").upper())
    seiz = re.sub(r"\W+", "", (values.get("description_seizure_number", "") or "").upper())
    addr = re.sub(r"\W+", "", (values.get("notice_to", "") or "").upper())[:24]
    base = f"{inv}|{seiz}|{addr}"
    return hashlib.sha1(base.encode("utf-8")).hexdigest()[:16]


def append_review_queue_csv(
    out_csv: Path,
    latest: Dict[str, str],
    errors: List[str],
    warnings: List[str],
) -> None:
    """Append records requiring manual review."""
    needs_review = latest.get("confidence_status") == "review" or bool(errors) or bool(warnings)
    if not needs_review:
        return
    write_header = not out_csv.exists()
    with _open_csv_append_or_create(out_csv) as f:
        w = csv.writer(f)
        if write_header:
            w.writerow(["record_id", "extracted_at", "source_file", "reason", "details"])
        reasons = []
        if latest.get("confidence_status") == "review":
            reasons.append("low_confidence")
        if errors:
            reasons.append("validation_error")
        if warnings:
            reasons.append("validation_warning")
        details = " | ".join(errors + warnings)
        w.writerow([
            normalize_output_text(latest.get("record_id", "")),
            normalize_output_text(latest.get("extracted_at", "")),
            normalize_output_text(latest.get("source_file", "")),
            ",".join(reasons),
            normalize_output_text(details),
        ])


def update_case_tracking_csv(
    case_index_csv: Path,
    duplicate_csv: Path,
    latest: Dict[str, str],
) -> bool:
    """Track case linking/dedupes. Returns True when duplicate case key is detected."""
    case_key = latest.get("case_key", "")
    if not case_key:
        return False

    existing: Dict[str, Dict[str, str]] = {}
    if case_index_csv.exists():
        try:
            with open(case_index_csv, "r", newline="", encoding="utf-8-sig") as f:
                r = csv.DictReader(f)
                for row in r:
                    if row.get("case_key"):
                        existing[row["case_key"]] = row
        except Exception:
            existing = {}

    is_dup = case_key in existing
    now = _timestamp_iso()
    row = existing.get(case_key, {})
    first_seen = row.get("first_seen", now)
    seen_count = int(row.get("seen_count", "0") or "0") + 1

    existing[case_key] = {
        "case_key": case_key,
        "first_seen": first_seen,
        "last_seen": now,
        "seen_count": str(seen_count),
        "last_record_id": latest.get("record_id", ""),
        "last_source_file": latest.get("source_file", ""),
        "inventory": latest.get("description_inventory", ""),
        "seizure_number": latest.get("description_seizure_number", ""),
    }

    with open(case_index_csv, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.DictWriter(
            f,
            fieldnames=[
                "case_key",
                "first_seen",
                "last_seen",
                "seen_count",
                "last_record_id",
                "last_source_file",
                "inventory",
                "seizure_number",
            ],
        )
        w.writeheader()
        for _, v in sorted(existing.items(), key=lambda kv: kv[0]):
            w.writerow(v)

    if is_dup:
        write_header = not duplicate_csv.exists()
        with _open_csv_append_or_create(duplicate_csv) as f:
            w = csv.writer(f)
            if write_header:
                w.writerow(["detected_at", "case_key", "record_id", "source_file", "inventory", "seizure_number"])
            w.writerow([
                normalize_output_text(now),
                normalize_output_text(case_key),
                normalize_output_text(latest.get("record_id", "")),
                normalize_output_text(latest.get("source_file", "")),
                normalize_output_text(latest.get("description_inventory", "")),
                normalize_output_text(latest.get("description_seizure_number", "")),
            ])
    return is_dup


def build_latest_civ_values(
    top: Dict[str, str],
    values: Dict[str, str],
    source_file: Path,
    extract_mode: str,
) -> Dict[str, str]:
    """Build merged and confidence-scored values for latest_civ.csv."""
    quality = _top_quality_score(top)
    field_scores = compute_field_confidence(top, values)
    min_field_score = min(field_scores.values()) if field_scores else 0
    confident = "yes" if (quality >= 80 and min_field_score >= 60) else "review"
    latest = {
        "record_id": f"{source_file.stem}_{_timestamp_compact()}",
        "extracted_at": _timestamp_iso(),
        "source_file": str(source_file),
        "extract_mode": extract_mode,
        "quality_score": str(quality),
        "min_field_confidence": str(min_field_score),
        "confidence_status": confident,
        "notice_to": values.get("notice_to", ""),
        "notice_date": values.get("notice_date", ""),
        "seizure_date_line": values.get("seizure_date_line", ""),
        "seizure_year_left": values.get("seizure_year_left", ""),
        "seizure_year_right": values.get("seizure_year_right", ""),
        "seizure_location": values.get("seizure_location", ""),
        "lieu_interception": values.get("lieu_interception", ""),
        "description_inventory": values.get("description_inventory", ""),
        "description_declared": values.get("description_declared", ""),
        "description_item": values.get("description_item", ""),
        "description_seizure_number": values.get("description_seizure_number", ""),
        "seizing_officer": values.get("seizing_officer", ""),
        "form_type": values.get("form_type", ""),
        "raw_destinataire": _top_first_match(top, ["DESTINATAIRE"]),
        "raw_expediteur": _top_first_match(top, ["EXPEDITEUR", "EXP?DITEUR"]),
        "raw_inventory": _top_first_match(top, ["INVENTAIRE"]),
        "raw_officer": _top_first_match(top, ["INSIGNE", "SAISISSANT"]),
        "raw_declared": _top_first_match(top, ["DECLARATION", "D?CLARATION"]),
        "raw_item": _top_first_match(top, ["DESCRIPTION", "ITEM", "SAISIR"]),
        "case_key": _build_case_key(values),
    }
    for k, v in field_scores.items():
        latest[k] = str(v)
    return latest

def write_latest_civ_csv(out_csv: Path, latest: Dict[str, str]) -> None:
    """Overwrite latest_civ.csv with the most recent merged/confident record."""
    with open(out_csv, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.writer(f)
        w.writerow(["field", "value"])
        for k, v in latest.items():
            w.writerow([normalize_output_text(k), normalize_output_text(v)])


def append_all_values_csv(
    out_csv: Path,
    top: Dict[str, str],
    values: Dict[str, str],
    source_file: Path,
    extract_mode: str,
) -> None:
    """Append every extracted value to all_values.csv for history/audit."""
    header = ["record_id", "extracted_at", "source_file", "extract_mode", "scope", "field", "value"]
    record_id = f"{source_file.stem}_{_timestamp_compact()}"
    ts = _timestamp_iso()
    write_header = not out_csv.exists()
    with _open_csv_append_or_create(out_csv) as f:
        w = csv.writer(f)
        if write_header:
            w.writerow(header)
        for k, v in top.items():
            w.writerow([
                normalize_output_text(record_id),
                normalize_output_text(ts),
                normalize_output_text(str(source_file)),
                normalize_output_text(extract_mode),
                "raw_top",
                normalize_output_text(k),
                normalize_output_text(v),
            ])
        for k, v in values.items():
            w.writerow([
                normalize_output_text(record_id),
                normalize_output_text(ts),
                normalize_output_text(str(source_file)),
                normalize_output_text(extract_mode),
                "k138",
                normalize_output_text(k),
                normalize_output_text(v),
            ])


# ======================== =
# Optional K138 fill hook -----
# ======================== =

# ======================== =
# Config file handling
# ======================== =

def load_config() -> configparser.ConfigParser:
    """Load configuration from radiance_copilot.cfg"""
    config = configparser.ConfigParser()
    if CONFIG_FILE.exists():
        config.read(CONFIG_FILE)
    return config

def save_config(config: configparser.ConfigParser):
    """Save configuration to radiance_copilot.cfg"""
    with open(CONFIG_FILE, 'w') as f:
        config.write(f)

def get_config_path(section: str, key: str, default: str = "") -> Optional[Path]:
    """Get a path from config file"""
    config = load_config()
    if section in config and key in config[section]:
        path_str = config[section][key]
        if path_str:
            return Path(path_str)
    return Path(default) if default else None

def set_config_path(section: str, key: str, path: Optional[Path]):
    """Set a path in config file"""
    config = load_config()
    if section not in config:
        config.add_section(section)
    config[section][key] = str(path) if path else ""
    save_config(config)


def get_config_text(section: str, key: str, default: str = "") -> str:
    """Get a plain text setting from config file."""
    config = load_config()
    if section in config and key in config[section]:
        return str(config[section].get(key, default) or default)
    return default


def set_config_text(section: str, key: str, value: str) -> None:
    """Set a plain text setting in config file."""
    config = load_config()
    if section not in config:
        config.add_section(section)
    config[section][key] = str(value or "")
    save_config(config)

# ======================== =
# Notice text file reading
# ======================== =

def load_notice_text(templates_folder: Path, form_type: str) -> str:
    """Load notice text from .txt file in templates folder."""
    if not templates_folder or not templates_folder.exists():
        return ""
    
    # Map form type to filename
    filename_map = {
        "Cannabis-Stupefiant": "k138_note_cannabis.txt",
        "Knives-Arms": "k138_note_arms.txt",
        "Stupefiant-Others": "k138_note_other.txt"
    }
    
    filename = filename_map.get(form_type, "k138_note_other.txt")
    file_path = templates_folder / filename
    
    if file_path.exists():
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except Exception as e:
            print(f"Warning: Could not read notice file {file_path}: {e}")
            return ""
    return ""

def find_saisie_template(templates_folder: Path) -> Optional[Path]:
    """Find SAISIE template PDF, excluding Agenda/K138 templates.

    Searches configured folder, its parent, and the app-relative 'templates/' folder.
    """
    def _search_folder(folder: Path) -> Optional[Path]:
        candidates: List[Tuple[int, str, Path]] = []
        for pdf in folder.glob("*.pdf"):
            name = _fold_ascii_lower(pdf.name)
            if "template" not in name:
                continue
            if ("k138" in name) or ("agenda" in name):
                continue
            if "saisie" not in name:
                continue
            score = 0
            if "a faire" in name:
                score += 5
            if "francompact" in name:
                score += 3
            if "dummy" not in name:
                score += 2
            if name.startswith("saisie"):
                score += 1
            candidates.append((score, name, pdf))
        if candidates:
            candidates.sort(key=lambda t: (t[0], t[1]), reverse=True)
            return candidates[0][2]
        # Conservative fallback: explicit "saisie" only, never Agenda/K138.
        for pdf in folder.glob("*.pdf"):
            name = _fold_ascii_lower(pdf.name)
            if ("saisie" in name) and ("agenda" not in name) and ("k138" not in name):
                return pdf
        return None

    for folder in _template_search_folders(templates_folder):
        found = _search_folder(folder)
        if found:
            return found
    return None


SAISIE_AFFAIRE_TEXT_FIELD_TO_BOX = {
    "bond_room_ledger": "BOND ROOM LEDGER #",
    "agent_badge": "# INSIGNE AGENT SAISISSANT:",
    "inventory_number": "# INVENTAIRE:",
    "country": "PAYS:",
    "interception_location": "LIEU INTERCEPTION:",
    "interception_datetime": "DATE / HEURE INTERCEPTION:",
    "declaration": "DÃ‰CLARATION:",
    "weight_qty": "POIDS / QTÃ‰ MARCH.:",
    "item_description": "DESCRIPTION DE Lâ€™ITEM Ã€ SAISIR:",
    "sender": "EXPÃ‰DITEUR :",
    "recipient": "DESTINATAIRE:",
    "indices": "INDICES:",
    "notes": "Notes",
}

SAISIE_AFFAIRE_MULTILINE_KEYS = {"item_description", "sender", "recipient", "indices", "notes"}


def _saisie_widget_norm_name(raw_name: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", _fold_ascii_lower(_repair_mojibake_text(clean_value(raw_name))))


def _saisie_guess_text_key(norm_name: str) -> Optional[str]:
    if not norm_name:
        return None
    token_map: Dict[str, List[str]] = {
        "bond_room_ledger": ["bondroomledger", "ledger", "bondroom"],
        "agent_badge": ["insigneagentsaisissant", "agentsaisissant", "insigneagent", "badge"],
        "inventory_number": ["noinventaire", "inventaire"],
        "country": ["pays", "country"],
        "interception_location": ["lieuinterception", "interceptionlocation"],
        "interception_datetime": ["dateheureinterception", "dateinterception", "interceptiondatetime"],
        "declaration": ["declaration", "declared"],
        "weight_qty": ["poidsqtemarch", "poids", "weightqty"],
        "item_description": ["descriptiondelitemasaisir", "descriptionitemasaisir", "itemdescription", "descriptionitem"],
        "sender": ["expediteur", "sender"],
        "recipient": ["destinataire", "recipient"],
        "indices": ["indices"],
        "notes": ["notes", "note"],
    }
    for key, tokens in token_map.items():
        if any(tok in norm_name for tok in tokens):
            return key
    return None


def _saisie_guess_checkbox_key(norm_name: str) -> Optional[str]:
    if not norm_name:
        return None
    if "envergure" in norm_name and "saisie" in norm_name:
        return "saisie_denvergure"
    if "confiscation" in norm_name:
        return "confiscation"
    if "k9" in norm_name:
        return "k9"
    if "saisie" in norm_name:
        return "saisie"
    return None


def _saisie_best_text_key_for_rect(
    wrect: "fitz.Rect",
    text_regions: Dict[str, "fitz.Rect"],
) -> Optional[str]:
    """Map widget rectangle to nearest SAISIE text key by overlap ratio."""
    if not text_regions:
        return None
    if (wrect.width <= 40) and (wrect.height <= 20):
        return None
    best_key: Optional[str] = None
    best_overlap = 0.0
    for candidate, region in text_regions.items():
        ov = _rect_overlap_area_fitz(wrect, region)
        if ov > best_overlap:
            best_overlap = ov
            best_key = candidate
    if not best_key or best_overlap <= 0:
        return None
    region = text_regions[best_key]
    w_area = max(1.0, float((wrect.x1 - wrect.x0) * (wrect.y1 - wrect.y0)))
    r_area = max(1.0, float((region.x1 - region.x0) * (region.y1 - region.y0)))
    overlap_ratio_widget = best_overlap / w_area
    overlap_ratio_region = best_overlap / r_area
    if (overlap_ratio_widget >= 0.45) or (overlap_ratio_region >= 0.20):
        return best_key
    return None


def extract_saisie_affaire_manual_fields_from_pdf(pdf_path: Path) -> Dict[str, str]:
    """
    Best-effort read of Saisie D'affaire text fields directly from fillable widgets.
    Used as a fallback when cache/manual JSON is missing.
    """
    out: Dict[str, str] = {}
    if not pdf_path or (pdf_path.suffix or "").lower() != ".pdf" or not pdf_path.exists():
        return out

    text_regions: Dict[str, fitz.Rect] = {}
    for key, box_label in SAISIE_AFFAIRE_TEXT_FIELD_TO_BOX.items():
        rect = FIELD_BOXES.get(box_label)
        if rect:
            text_regions[key] = fitz.Rect(*rect)

    try:
        doc = fitz.open(str(pdf_path))
    except Exception:
        return out
    try:
        if len(doc) == 0:
            return out
        page = doc[0]
        widgets = list(page.widgets() or [])
        text_type = int(getattr(fitz, "PDF_WIDGET_TYPE_TEXT", 7))
        for w in widgets:
            try:
                wtype = int(getattr(w, "field_type", 0) or 0)
            except Exception:
                continue
            if wtype != text_type:
                continue
            norm_name = _saisie_widget_norm_name(str(getattr(w, "field_name", "") or ""))
            key_by_name = _saisie_guess_text_key(norm_name)
            key_by_rect = _saisie_best_text_key_for_rect(fitz.Rect(w.rect), text_regions)
            key = key_by_rect or key_by_name
            if not key:
                continue
            val = clean_value(str(getattr(w, "field_value", "") or ""))
            if not val:
                continue
            # Keep the longest non-empty value if multiple widgets map to same key.
            if len(val) >= len(clean_value(out.get(key, ""))):
                out[key] = val
    finally:
        try:
            doc.close()
        except Exception:
            pass

    badge = re.sub(r"\D", "", out.get("agent_badge", ""))
    if badge:
        out["agent_badge"] = badge[-5:] if len(badge) > 5 else badge.zfill(5)
    inv = _normalize_inventory_number(out.get("inventory_number", ""))
    if inv:
        out["inventory_number"] = inv
    return out


def apply_saisie_affaire_manual_to_k138_values(
    values_dict: Dict[str, str],
    manual_fields: Dict[str, str],
) -> None:
    """Apply structured Saisie D'affaire fields onto K138 base values."""
    if not isinstance(values_dict, dict) or not isinstance(manual_fields, dict):
        return

    inv = _normalize_inventory_number(clean_value(manual_fields.get("inventory_number", "")))
    if inv:
        values_dict["description_inventory"] = inv

    decl = clean_value(manual_fields.get("declaration", ""))
    if decl:
        values_dict["description_declared"] = decl

    item = clean_value(manual_fields.get("item_description", ""))
    if item:
        values_dict["description_item"] = item

    badge = re.sub(r"\D", "", clean_value(manual_fields.get("agent_badge", "")))
    if badge:
        values_dict["seizing_officer"] = badge[-5:] if len(badge) > 5 else badge.zfill(5)

    recipient = clean_value(manual_fields.get("recipient", ""))
    if recipient:
        values_dict["notice_to"] = normalize_multiline_to_pipe(recipient.replace(",", "\n"))

    dt = parse_first_date(clean_value(manual_fields.get("interception_datetime", "")))
    if dt:
        yyyy, mm, dd = dt
        en, fr = month_en_fr(mm)
        values_dict["seizure_date_line"] = f"{dd} {en} / {dd} {fr}"
        values_dict["seizure_year_left"] = str(yyyy)[:2]
        values_dict["seizure_year_right"] = str(yyyy)[2:]

    # Hard guardrail: never allow sender/recipient leakage in item field.
    if _is_weak_item_text(values_dict.get("description_item", "")):
        fallback_item = decl or clean_value(values_dict.get("description_declared", ""))
        values_dict["description_item"] = fallback_item


def fill_saisie_affaire_pdf(
    template_path: Path,
    output_path: Path,
    field_values: Dict[str, str],
    check_values: Dict[str, bool],
) -> Tuple[int, int, int]:
    """
    Fill SAISIE template with Saisie D'affaire values while keeping the PDF editable.

    Strategy:
    - Copy template to output.
    - Populate existing widgets by name (and geometry fallback).
    - Create missing text widgets over known field boxes.
    Returns: (updated_text_widgets, created_text_widgets, updated_check_widgets)
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output_path)

    updated_text = 0
    created_text = 0
    updated_checks = 0
    tmp_path = output_path.with_name(f"{output_path.stem}._saving_tmp{output_path.suffix}")
    if tmp_path.exists():
        try:
            tmp_path.unlink()
        except Exception:
            pass

    text_regions: Dict[str, fitz.Rect] = {}
    for key, box_label in SAISIE_AFFAIRE_TEXT_FIELD_TO_BOX.items():
        rect = FIELD_BOXES.get(box_label)
        if rect:
            text_regions[key] = fitz.Rect(*rect)

    doc = fitz.open(str(output_path))
    try:
        if len(doc) == 0:
            raise RuntimeError("SAISIE template has no pages.")
        page = doc[0]
        widgets = list(page.widgets() or [])
        text_type = int(getattr(fitz, "PDF_WIDGET_TYPE_TEXT", 7))
        checkbox_type = int(getattr(fitz, "PDF_WIDGET_TYPE_CHECKBOX", 2))
        radio_type = int(getattr(fitz, "PDF_WIDGET_TYPE_RADIOBUTTON", 5))
        button_type = int(getattr(fitz, "PDF_WIDGET_TYPE_BUTTON", 1))

        written_text_keys: set[str] = set()
        covered_text_keys: set[str] = set()
        written_checkbox_keys: set[str] = set()
        unmatched_top_checkbox_widgets: List["fitz.Widget"] = []
        checkbox_order = ["saisie", "confiscation", "k9", "saisie_denvergure"]
        top_checkbox_y_max = 260.0
        named_top_checkbox_keys_present: set[str] = set()
        has_any_top_checkbox_widgets = False

        def _write_checkbox(widget, checked: bool) -> bool:
            try:
                flags = int(getattr(widget, "field_flags", 0) or 0)
                if flags & 1:
                    widget.field_flags = flags & ~1
            except Exception:
                pass
            candidates = []
            on_state = None
            try:
                on_state = widget.on_state()
            except Exception:
                on_state = None
            if checked:
                candidates = [on_state, "Yes", "On", "1", "True", "X"]
            else:
                candidates = ["Off", "0", "False", "", " "]
            for candidate in candidates:
                if candidate is None:
                    continue
                try:
                    widget.field_value = str(candidate)
                    widget.update()
                    return True
                except Exception:
                    continue
            return False

        def _best_text_key_for_widget_rect(wrect: "fitz.Rect") -> Optional[str]:
            if not text_regions:
                return None
            # Ignore tiny top widgets; these are often checkbox helper fields.
            if (wrect.width <= 40) and (wrect.height <= 20):
                return None
            best_key: Optional[str] = None
            best_overlap = 0.0
            for candidate, region in text_regions.items():
                ov = _rect_overlap_area_fitz(wrect, region)
                if ov > best_overlap:
                    best_overlap = ov
                    best_key = candidate
            if not best_key or best_overlap <= 0:
                return None
            region = text_regions[best_key]
            w_area = max(1.0, float((wrect.x1 - wrect.x0) * (wrect.y1 - wrect.y0)))
            r_area = max(1.0, float((region.x1 - region.x0) * (region.y1 - region.y0)))
            overlap_ratio_widget = best_overlap / w_area
            overlap_ratio_region = best_overlap / r_area
            if (overlap_ratio_widget >= 0.45) or (overlap_ratio_region >= 0.20):
                return best_key
            return None

        def _has_existing_text_widget_overlap(target_rect: "fitz.Rect") -> bool:
            region_area = max(1.0, float((target_rect.x1 - target_rect.x0) * (target_rect.y1 - target_rect.y0)))
            for ww in widgets:
                try:
                    ww_type = int(getattr(ww, "field_type", 0) or 0)
                except Exception:
                    continue
                if ww_type != text_type:
                    continue
                wrect = fitz.Rect(ww.rect)
                ov = _rect_overlap_area_fitz(wrect, target_rect)
                if ov <= 0:
                    continue
                w_area = max(1.0, float((wrect.x1 - wrect.x0) * (wrect.y1 - wrect.y0)))
                # Conservative overlap thresholds to avoid creating duplicate fallback widgets.
                if (ov / region_area) >= 0.08 or (ov / w_area) >= 0.30:
                    return True
            return False

        for w in widgets:
            try:
                wtype = int(getattr(w, "field_type", 0) or 0)
            except Exception:
                continue
            if wtype not in (checkbox_type, radio_type, button_type):
                continue
            norm_name = _saisie_widget_norm_name(str(getattr(w, "field_name", "") or ""))
            ckey = _saisie_guess_checkbox_key(norm_name)
            rw = fitz.Rect(w.rect)
            if rw.y0 <= top_checkbox_y_max:
                has_any_top_checkbox_widgets = True
            if ckey and (rw.y0 <= top_checkbox_y_max):
                named_top_checkbox_keys_present.add(ckey)

        for idx, w in enumerate(widgets):
            norm_name = _saisie_widget_norm_name(str(getattr(w, "field_name", "") or ""))
            try:
                wtype = int(getattr(w, "field_type", 0) or 0)
            except Exception:
                continue

            if wtype == text_type:
                wrect = fitz.Rect(w.rect)
                key_by_name = _saisie_guess_text_key(norm_name)
                key_by_rect = _best_text_key_for_widget_rect(wrect)
                # Prefer geometry when it disagrees with name.
                # Some templates reuse one field name across multiple boxes
                # (e.g. DECLARATION / NOTES), so name-only mapping is unreliable.
                key = key_by_name
                if key_by_rect and (not key or key_by_rect != key):
                    key = key_by_rect
                if key:
                    covered_text_keys.add(key)
                if not key:
                    ckey_text = _saisie_guess_checkbox_key(norm_name)
                    if (
                        ckey_text
                        and (wrect.y0 <= top_checkbox_y_max)
                        and (not has_any_top_checkbox_widgets)
                        and (ckey_text not in named_top_checkbox_keys_present)
                        and (ckey_text not in written_checkbox_keys)
                    ):
                        checked = bool(check_values.get(ckey_text, False))
                        try:
                            flags = int(getattr(w, "field_flags", 0) or 0)
                            if flags & 1:
                                w.field_flags = flags & ~1
                        except Exception:
                            pass
                        try:
                            w.field_value = "X" if checked else ""
                            w.update()
                            updated_checks += 1
                            written_checkbox_keys.add(ckey_text)
                        except Exception:
                            pass
                    else:
                        # Clear common placeholder artifacts in recipient fields.
                        cur = _fold_ascii_lower(clean_value(str(getattr(w, "field_value", "") or "")))
                        is_placeholder = ("inconnu" in cur) or (cur in {"8", "inconnu8"})
                        rec_region = text_regions.get("recipient")
                        overlaps_recipient = bool(rec_region and (_rect_overlap_area_fitz(fitz.Rect(w.rect), rec_region) > 0))
                        if is_placeholder and (overlaps_recipient or ("destinataire" in norm_name)):
                            try:
                                w.field_value = " "
                                w.update()
                            except Exception:
                                pass
                    continue
                val = clean_value(field_values.get(key, ""))
                try:
                    flags = int(getattr(w, "field_flags", 0) or 0)
                    if flags & 1:
                        w.field_flags = flags & ~1
                except Exception:
                    pass
                # Some templates reuse the same field name for multiple widgets.
                # If declaration/notes share a field name, values mirror each other.
                # Disambiguate per widget before writing so each area keeps its own value.
                try:
                    desired_name = f"SAISIE_{key}_w{idx}"
                    if clean_value(str(getattr(w, "field_name", "") or "")) != desired_name:
                        w.field_name = desired_name
                        w.update()
                except Exception:
                    pass
                try:
                    if key in SAISIE_AFFAIRE_MULTILINE_KEYS:
                        ml_flag = int(getattr(fitz, "PDF_FIELD_IS_MULTILINE", 4096))
                        cur_flags = int(getattr(w, "field_flags", 0) or 0)
                        w.field_flags = cur_flags | ml_flag
                except Exception:
                    pass
                try:
                    w.field_value = val if val else " "
                    w.update()
                    updated_text += 1
                    written_text_keys.add(key)
                except Exception:
                    continue
                continue

            if wtype in (checkbox_type, radio_type, button_type):
                ckey = _saisie_guess_checkbox_key(norm_name)
                if not ckey:
                    # Keep only top-band small boxes for fallback mapping to avoid checklist section.
                    rw = fitz.Rect(w.rect)
                    if rw.y0 <= top_checkbox_y_max and (rw.width <= 26) and (rw.height <= 26):
                        unmatched_top_checkbox_widgets.append(w)
                    continue
                if w.rect.y0 > top_checkbox_y_max:
                    # Never touch checklist section checkboxes.
                    continue
                checked = bool(check_values.get(ckey, False))
                if _write_checkbox(w, checked):
                    updated_checks += 1
                    written_checkbox_keys.add(ckey)
                continue

        # Geometry/order fallback for unnamed TOP checkboxes only.
        if unmatched_top_checkbox_widgets:
            remaining_keys = [
                k
                for k in checkbox_order
                if (k in check_values)
                and (k not in named_top_checkbox_keys_present)
                and (k not in written_checkbox_keys)
            ]
            unresolved = sorted(unmatched_top_checkbox_widgets, key=lambda ww: (ww.rect.y0, ww.rect.x0))
            for widget, key in zip(unresolved, remaining_keys):
                if _write_checkbox(widget, bool(check_values.get(key, False))):
                    updated_checks += 1
                    written_checkbox_keys.add(key)

        # If template lacks some text widgets, create them so output remains editable.
        for key, rect in text_regions.items():
            if key in covered_text_keys:
                continue
            if _has_existing_text_widget_overlap(rect):
                continue
            val = clean_value(field_values.get(key, ""))
            try:
                widget = fitz.Widget()
                widget.field_type = fitz.PDF_WIDGET_TYPE_TEXT
                widget.field_name = f"SAISIE_{key}"
                widget.field_value = val
                widget.rect = fitz.Rect(rect)
                widget.fill_color = (1, 1, 1)
                widget.text_color = (0, 0, 0)
                widget.text_fontsize = 8
                widget.border_width = 0.5
                widget.border_color = (0.65, 0.65, 0.65)
                if key in SAISIE_AFFAIRE_MULTILINE_KEYS:
                    widget.field_flags = int(getattr(fitz, "PDF_FIELD_IS_MULTILINE", 4096))
                page.add_widget(widget)
                created_text += 1
            except Exception:
                continue

        doc.save(str(tmp_path), incremental=False, encryption=fitz.PDF_ENCRYPT_KEEP)
    finally:
        doc.close()

    os.replace(str(tmp_path), str(output_path))
    if tmp_path.exists():
        try:
            tmp_path.unlink()
        except Exception:
            pass

    return updated_text, created_text, updated_checks

def find_k138_template(templates_folder: Path, form_type: str) -> Optional[Path]:
    """Find K138 template PDF in templates folder matching the form type."""
    if not templates_folder or not templates_folder.exists():
        return None

    def _matches_form(name_lower: str, ft: str) -> bool:
        if "k138" not in name_lower:
            return False
        if ft == "Cannabis-Stupefiant":
            return "cannabis" in name_lower
        if ft == "Knives-Arms":
            return ("knife" in name_lower) or ("knives" in name_lower) or ("arms" in name_lower)
        # Stupefiant-Others: exclude cannabis/knives/arms.
        if ("cannabis" in name_lower) or ("knife" in name_lower) or ("knives" in name_lower) or ("arms" in name_lower):
            return False
        return ("stupefiant-others" in name_lower) or ("stupefiant" in name_lower) or ("others" in name_lower)

    def _widget_count(pdf_path: Path) -> int:
        """Count fillable widgets on first pages; higher count => better fillable template."""
        try:
            doc = fitz.open(str(pdf_path))
            try:
                total = 0
                for i in range(min(4, len(doc))):
                    total += len(list(doc[i].widgets() or []))
                return total
            finally:
                doc.close()
        except Exception:
            return 0

    candidates: List[Tuple[int, str, Path]] = []
    for pdf in templates_folder.glob("*.pdf"):
        name_lower = _fold_ascii_lower(pdf.name)
        if not _matches_form(name_lower, form_type):
            continue

        score = 0
        if "template" in name_lower:
            score += 12
        if "dummy" not in name_lower:
            score += 8
        if "stupefiant-others-template" in name_lower:
            score += 10
        if name_lower.startswith("2-"):
            score += 2

        # Prefer templates that already contain real fillable fields.
        widgets = _widget_count(pdf)
        score += min(widgets, 60)

        candidates.append((score, name_lower, pdf))

    if not candidates:
        # Fallback: if a specific form template is missing, use best available K138 template.
        fallback_candidates: List[Tuple[int, str, Path]] = []
        for pdf in templates_folder.glob("*.pdf"):
            name_lower = _fold_ascii_lower(pdf.name)
            if "k138" not in name_lower:
                continue

            score = 0
            if "template" in name_lower:
                score += 10
            if "dummy" not in name_lower:
                score += 6
            widgets = _widget_count(pdf)
            score += min(widgets, 60)

            fallback_candidates.append((score, name_lower, pdf))

        if not fallback_candidates:
            return None
        fallback_candidates.sort(key=lambda t: (t[0], t[1]), reverse=True)
        return fallback_candidates[0][2]

    candidates.sort(key=lambda t: (t[0], t[1]), reverse=True)
    return candidates[0][2]


def _template_search_folders(templates_folder: Optional[Path]) -> List[Path]:
    """Return ordered list of folders to search for template files.

    Priority:
      1. Configured templates_folder
      2. Parent of configured folder (handles accidental case-subfolder selection)
      3. 'templates/' next to the running script / frozen exe
    """
    seen: set = set()
    result: List[Path] = []

    def _add(p: Optional[Path]) -> None:
        if p and p.exists() and p.is_dir():
            key = str(p.resolve())
            if key not in seen:
                seen.add(key)
                result.append(p.resolve())

    _add(templates_folder)
    if templates_folder:
        _add(templates_folder.parent)

    # App-relative 'templates/' folder (works both as script and frozen exe)
    if getattr(sys, "frozen", False):
        app_dir = Path(sys.executable).parent
    else:
        app_dir = Path(__file__).parent
    _add(app_dir / "templates")

    return result


def find_agenda_template(templates_folder: Path) -> Optional[Path]:
    """Find official Agenda template (PDF preferred, DOCX fallback).

    Searches configured folder, its parent, and the app-relative 'templates/' folder.
    """
    def _search_folder(folder: Path) -> Optional[Path]:
        pdfs = [p for p in folder.glob("*.pdf") if "agenda" in p.name.lower()]
        if pdfs:
            for p in pdfs:
                if "template" in p.name.lower():
                    return p
            return pdfs[0]
        docxs = [p for p in folder.glob("*.docx") if "agenda" in p.name.lower()]
        if docxs:
            for p in docxs:
                if "template" in p.name.lower():
                    return p
            return docxs[0]
        return None

    for folder in _template_search_folders(templates_folder):
        found = _search_folder(folder)
        if found:
            return found
    return None


def _fold_ascii_lower(s: str) -> str:
    out = normalize_output_text(s or "").lower()
    repl = {
        "é": "e",
        "è": "e",
        "ê": "e",
        "ë": "e",
        "à": "a",
        "â": "a",
        "ä": "a",
        "î": "i",
        "ï": "i",
        "ô": "o",
        "ö": "o",
        "ù": "u",
        "û": "u",
        "ü": "u",
        "ç": "c",
        "’": "'",
    }
    for k, v in repl.items():
        out = out.replace(k, v)
    return out


def find_saisie_interet_template(templates_folder: Path) -> Optional[Path]:
    """Find SAISIE D'INTERET Excel template.

    Searches configured folder, its parent, and the app-relative 'templates/' folder.
    """
    def _search_folder(folder: Path) -> Optional[Path]:
        candidates: List[Tuple[int, str, Path]] = []
        for ext in ("*.xlsx", "*.xlsm", "*.xltx", "*.xltm"):
            for xl in folder.glob(ext):
                name = _fold_ascii_lower(xl.name)
                if ("saisie" in name) and ("interet" in name):
                    score = 0
                    if "template" in name:
                        score += 3
                    if "dummy" not in name:
                        score += 1
                    if xl.suffix.lower() == ".xlsx":
                        score += 1
                    candidates.append((score, name, xl))
        if candidates:
            candidates.sort(key=lambda t: (t[0], t[1]), reverse=True)
            return candidates[0][2]
        # Fallback: any Excel file with "interet".
        for ext in ("*.xlsx", "*.xlsm", "*.xltx", "*.xltm"):
            for xl in folder.glob(ext):
                if "interet" in _fold_ascii_lower(xl.name):
                    return xl
        return None

    for folder in _template_search_folders(templates_folder):
        found = _search_folder(folder)
        if found:
            return found
    return None


def _split_name_and_address(raw: str) -> Tuple[str, str]:
    lines: List[str] = []
    seen: set = set()
    for ln in re.split(r"[\n|]+", raw or ""):
        v = clean_value(ln)
        if not v:
            continue
        if re.fullmatch(r"(?:DESTINATAIRE|EXP(?:Ã‰|É|E)DITEUR)\s*:?", v, re.IGNORECASE):
            continue
        key = v.upper()
        if key in seen:
            continue
        seen.add(key)
        lines.append(v)
    if not lines:
        return "", ""
    if len(lines) == 1:
        return lines[0], ""
    return lines[0], ", ".join(lines[1:])


def _interet_yes_no(raw: str) -> str:
    t = clean_value(raw or "").upper()
    if not t:
        return ""
    if re.search(r"(?:\bX\b|✓|✔|\bOUI\b|\bYES\b|\bTRUE\b|\b1\b)", t):
        return "Oui"
    if re.search(r"(?:\bNON\b|\bNO\b|\bFALSE\b|\b0\b)", t):
        return "Non"
    return ""


def _interet_type_from_form(form_type: str, item_desc: str) -> str:
    t = clean_value(form_type or "").lower()
    if "cannabis" in t:
        return "Cannabis"
    if ("knife" in t) or ("arm" in t):
        return "Arme blanche"
    if "stupefiant" in t:
        return "Stupefiant"
    item = clean_value(item_desc or "")
    return item[:80] if item else ""


def build_saisie_interet_rows(
    top: Dict[str, str],
    values: Dict[str, str],
    case_folder_name: str = "",
) -> Dict[int, object]:
    """
    Map extracted SAISIE/K138 values to SAISIE D'INTERET row numbers (column C).
    """
    rows: Dict[int, object] = {}

    dt = parse_first_date(safe_get(top, "DATE / HEURE INTERCEPTION:"))
    if dt:
        y, m, d = dt
        rows[7] = f"{y:04d}-{m:02d}-{d:02d}"
    elif clean_value(values.get("notice_date", "")):
        rows[7] = clean_value(values.get("notice_date", ""))

    agent = re.sub(r"\D", "", values.get("seizing_officer", "") or "")
    if agent:
        rows[8] = agent[-5:] if len(agent) > 5 else agent.zfill(5)

    inv = _normalize_inventory_number(values.get("description_inventory", "") or "")
    if inv:
        rows[10] = inv

    item = clean_value(values.get("description_item", "") or "")
    item_type = _interet_type_from_form(values.get("form_type", ""), item)
    if item_type:
        rows[12] = item_type
    qty = safe_get(top, "POIDS / QTÃ‰ MARCH.:") or _top_first_match(top, ["POIDS", "QT"])
    if qty:
        rows[13] = qty
    if item:
        rows[15] = item

    sender_raw = safe_get(top, "EXPÃ‰DITEUR :") or _top_first_match(top, ["EXPEDITEUR", "EXP?DITEUR"])
    s_name, s_addr = _split_name_and_address(sender_raw)
    if s_name:
        rows[17] = s_name
    if s_addr:
        rows[18] = s_addr

    dest_raw = clean_value(values.get("notice_to", "") or "").replace("|", "\n")
    if not dest_raw:
        dest_raw = safe_get(top, "DESTINATAIRE:") or _top_first_match(top, ["DESTINATAIRE"])
    d_name, d_addr = _split_name_and_address(dest_raw)
    if d_name:
        rows[21] = d_name
    if d_addr:
        rows[22] = d_addr

    declared = clean_value(values.get("description_declared", "") or "")
    if declared:
        rows[25] = declared

    report_no = _top_first_match(top, ["RAPPORT"])
    if report_no:
        rows[33] = report_no
    k9_raw = _top_first_match(top, ["K9"])
    k9_yes_no = _interet_yes_no(k9_raw)
    if k9_yes_no:
        rows[34] = k9_yes_no

    if clean_value(case_folder_name):
        rows[37] = clean_value(case_folder_name)
    rows[38] = "Generated by Radiance Co-Pilot."
    return rows


def _k138_text_all_pages(k138_pdf: Path) -> str:
    try:
        doc = fitz.open(str(k138_pdf))
        chunks = []
        for i in range(len(doc)):
            chunks.append(doc[i].get_text("text") or "")
        doc.close()
        return normalize_output_text("\n".join(chunks))
    except Exception:
        return ""


def extract_agenda_values_from_k138(k138_pdf: Path) -> Tuple[str, str]:
    """
    Extract inventory_number and agent_id directly from generated K138 PDF text.
    """
    text = _k138_text_all_pages(k138_pdf)
    if not text:
        return "", ""

    inventory = ""
    for pat in [
        r"INVENTORY\s+NO\s*/\s*NO\.?\s*D['â€™]INVENTAIRE\s*:\s*([A-Z0-9\- ]{6,60})",
        r"NO\.?\s*D['â€™]INVENTAIRE\s*:\s*([A-Z0-9\- ]{6,60})",
        r"INVENTAIRE\s*:\s*([A-Z0-9\- ]{6,60})",
    ]:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            inventory = _normalize_inventory_number(m.group(1) or "")
            if inventory:
                break
    if not inventory:
        for ln in (text or "").splitlines():
            if re.search(r"\bINVENTAIRE\b", ln, re.IGNORECASE):
                inventory = _normalize_inventory_number(ln)
                if inventory:
                    break
    if not inventory:
        inventory = _extract_inventory_from_text_loose(text)

    agent_id = ""
    for pat in [
        r"SEIZING\s+OFFICER(?:\s*/\s*AGENT\s+SAISISSANT)?\s*:\s*(\d{4,6})",
        r"AGENT\s+SAISISSANT(?:\s*/\s*SEIZING\s+OFFICER)?\s*:\s*(\d{4,6})",
        r"(?:NO\.?\s*D['â€™]?\s*INSIGNE|INSIGNE|BADGE(?:\s*NO\.?)?)\s*[:#]?\s*(\d{4,6})",
        r"(\d{4,6})\s*(?:NO\.?\s*D['â€™]?\s*INSIGNE|INSIGNE|BADGE(?:\s*NO\.?)?)",
    ]:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            agent_id = m.group(1)
            break

    # Normalize agent to 5 digits when possible.
    d = re.sub(r"\D", "", agent_id or "")
    if d:
        agent_id = d[-5:] if len(d) > 5 else d.zfill(5)

    return inventory, agent_id


def _agenda_text_all_pages(agenda_path: Path) -> str:
    suffix = (agenda_path.suffix or "").lower()
    if suffix == ".pdf":
        try:
            doc = fitz.open(str(agenda_path))
            chunks = []
            for i in range(len(doc)):
                chunks.append(doc[i].get_text("text") or "")
            doc.close()
            return normalize_output_text("\n".join(chunks))
        except Exception:
            return ""
    if suffix == ".docx" and HAVE_DOCX:
        try:
            doc = DocxDocument(str(agenda_path))
            lines: List[str] = []
            for p in doc.paragraphs:
                txt = clean_value(p.text)
                if txt:
                    lines.append(txt)
            for table in doc.tables:
                for row in table.rows:
                    row_txt = " ".join(clean_value(c.text) for c in row.cells if clean_value(c.text))
                    if row_txt:
                        lines.append(row_txt)
            return normalize_output_text("\n".join(lines))
        except Exception:
            return ""
    return ""


def _extract_inventory_from_text_label_only(text: str) -> str:
    """
    Strictly extract inventory candidate only from inventory-labeled text blocks.
    Used as a last-resort recovery path before blocking agenda generation.
    """
    if not text:
        return ""
    patterns = [
        r"(?:#\s*INVENTAIRE|NO\.?\s*D['’]?\s*INVENTAIRE|NO\s*INVENTAIRE|INVENTAIRE)\s*:?\s*([A-Z0-9][A-Z0-9\-/ ]{6,60})",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if not m:
            continue
        raw = clean_value(m.group(1))
        norm = _normalize_inventory_number(raw)
        if norm:
            return norm
        compact = re.sub(r"[^A-Za-z0-9]+", "", raw).upper()
        # Keep a strong fallback token when OCR dropped expected letters.
        if len(compact) >= 10 and re.search(r"\d", compact):
            return compact
    return ""


def _compact_alnum_token(raw: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "", (raw or "")).upper()


def _looks_like_inventory_number_strict(raw: str) -> bool:
    """
    Strict inventory-shape detector used to prevent CE/CID mixups.
    Keep this stricter than _normalize_inventory_number() fallback logic.
    """
    compact = _compact_alnum_token(clean_value(raw or ""))
    if not compact:
        return False
    # Canonical: 2 letters + 9 digits + 2 letters
    if re.fullmatch(r"[A-Z]{2}\d{9}[A-Z]{2}", compact):
        return True
    # Legacy style often used in these forms: 1 letter + long digit run
    if re.fullmatch(r"[A-Z]\d{10,24}", compact):
        return True
    return False


def _extract_sied_candidate(raw: str) -> str:
    """
    Extract a CE/CID-like token from noisy text and return a normalized value.
    Returns empty string when no reliable token is found.
    """
    v = clean_value(raw)
    if not v:
        return ""
    v = re.sub(r"(?:CE\s*/\s*CID|CECID|SIED)\s*[:#]?\s*", " ", v, flags=re.IGNORECASE)
    v = re.sub(r"\s*-\s*", "-", v)
    v = re.sub(r"\s+", " ", v).strip(" _-:")
    if not v:
        return ""

    # Reject prose/noise snippets that often appear near the SIED label.
    if re.search(r"\b(?:ETIQUET|TICKET|PIECE|PCS|LABELS?)\b", v, re.IGNORECASE):
        return ""

    upper = v.upper()
    def _is_known_form_code(tok: str) -> bool:
        t = (tok or "").upper()
        if t in {"K132", "K138", "K1138"}:
            return True
        # Common document codes that are not CE/CID values.
        if re.fullmatch(r"K\d{3}", t):
            return True
        return False

    # Prefer strongly structured numeric CE/CID values first.
    pattern_priority = [
        r"\b\d{3,6}-\d{2}-\d{3,6}\b",
        r"\b\d{3,6}(?:-\d{2,6}){1,3}\b",
        r"\b[A-Z]{1,5}\d{4,}(?:-\d{1,6}){0,3}\b",
    ]
    for pat in pattern_priority:
        candidates = []
        for m in re.finditer(pat, upper):
            tok = m.group(0).strip()
            if _is_known_form_code(tok):
                continue
            candidates.append(tok)
        if candidates:
            candidates.sort(key=len, reverse=True)
            return candidates[0]

    compact = re.sub(r"\s+", "", upper)
    if re.fullmatch(r"[A-Z0-9]+(?:-[A-Z0-9]+){0,3}", compact):
        has_letters = bool(re.search(r"[A-Z]", compact))
        has_hyphen = "-" in compact
        if (
            (sum(ch.isdigit() for ch in compact) >= 3)
            and (has_letters or has_hyphen)
            and (not _is_known_form_code(compact))
        ):
            return compact
    return ""


def _normalize_sied_value(raw: str, allow_raw: bool = False) -> str:
    """Normalize CE/CID(SIED) to a clean token. Optionally keep raw value when no token is found."""
    v = clean_value(raw)
    if not v:
        return ""
    token = _extract_sied_candidate(v)
    if token:
        return token
    return v if allow_raw else ""


def _is_valid_sied_value(raw: str) -> bool:
    """Validate CE/CID(SIED) candidate and reject checkbox/noise tokens."""
    v = clean_value(raw)
    if not v:
        return False
    v = re.sub(r"\s+", " ", v).strip()

    # Typical checkbox/noise artifacts from OCR/PDF extraction.
    if re.fullmatch(r"[xXoO✓✔\-\._ ]{1,20}", v):
        return False
    if re.fullmatch(r"\d+\s*[xX]", v):
        return False
    if re.fullmatch(r"[xX]\s*\d+", v):
        return False

    token = _extract_sied_candidate(v)
    if not token:
        return False
    compact = re.sub(r"\s+", "", token).upper()
    # Never accept values that look like inventory IDs.
    if _looks_like_inventory_number_strict(compact):
        return False
    # Ignore common template/placeholder artifacts that should not auto-populate.
    if compact in {"929", "000", "0000", "N/A", "NA"}:
        return False
    # Pure 1-3 digit tokens are too weak and often false positives.
    if re.fullmatch(r"\d{1,3}", compact):
        return False
    return True


def _extract_sied_from_agenda_pdf_field(agenda_path: Path) -> str:
    """Read CE/CID directly from the SIED value rectangle on Agenda PDF page 1."""
    if (agenda_path.suffix or "").lower() != ".pdf":
        return ""
    try:
        doc = fitz.open(str(agenda_path))
        page = doc[0]
        sied_label = _find_first_label_rect(page, ["SIED:", "SIED :", "SIED"])
        if sied_label:
            row_words: List[Tuple[float, float, float, float, str]] = []
            for w in page.get_text("words"):
                x0, y0, x1, y1, t, *_rest = w
                cy = (y0 + y1) / 2.0
                if (
                    y0 <= sied_label.y1 + 8
                    and y1 >= sied_label.y0 - 8
                    and x0 >= sied_label.x1 - 4
                    and x0 <= min(page.rect.width, sied_label.x1 + 340)
                    and sied_label.y0 - 12 <= cy <= sied_label.y1 + 12
                ):
                    row_words.append((x0, y0, x1, y1, t))
            raw_row = join_words(sorted(row_words, key=lambda t: (t[1], t[0])))
            sied_row = _extract_sied_candidate(raw_row)
            if sied_row:
                doc.close()
                return sied_row

        _agent_rect, _barcode_rect, _inv_pt, sied_rect = _agenda_rects_for_page(page)
        rect = fitz.Rect(
            max(0, sied_rect.x0 - 4),
            max(0, sied_rect.y0 - 3),
            min(page.rect.width, sied_rect.x1 + 4),
            min(page.rect.height, sied_rect.y1 + 3),
        )
        picked: List[Tuple[float, float, float, float, str]] = []
        for w in page.get_text("words"):
            x0, y0, x1, y1, t, *_rest = w
            cx = (x0 + x1) / 2.0
            cy = (y0 + y1) / 2.0
            if rect.x0 <= cx <= rect.x1 and rect.y0 <= cy <= rect.y1:
                picked.append((x0, y0, x1, y1, t))
        doc.close()
        raw = join_words(sorted(picked, key=lambda t: (t[1], t[0])))
        return _extract_sied_candidate(raw)
    except Exception:
        return ""


def extract_sied_from_agenda(agenda_path: Path) -> str:
    """
    Extract CE/CID or SIED value from an Agenda file.
    Returns cleaned token (letters/digits/-/ /).
    """
    field_sied = _extract_sied_from_agenda_pdf_field(agenda_path)
    if field_sied and _is_valid_sied_value(field_sied):
        return field_sied

    text = _agenda_text_all_pages(agenda_path)
    if not text:
        return ""

    patterns = [
        r"(?:CE\s*/\s*CID|CECID|SIED)\s*[:#]?\s*([A-Z0-9][A-Z0-9\- ]{1,40})",
        r"(?:CE\s*/\s*CID|CECID|SIED)\s*[:#]?\s*_*([A-Z0-9][A-Z0-9\- ]{1,40})",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if not m:
            continue
        val = _extract_sied_candidate(m.group(1))
        if val and _is_valid_sied_value(val):
            return val

    lines = [clean_value(x) for x in text.splitlines() if clean_value(x)]
    for i, ln in enumerate(lines):
        if re.search(r"\b(?:CE\s*/\s*CID|CECID|SIED)\b", ln, re.IGNORECASE):
            # value can be on same line or next line
            same = re.sub(r".*?(?:CE\s*/\s*CID|CECID|SIED)\s*[:#]?\s*", "", ln, flags=re.IGNORECASE).strip("_ -")
            same_val = _extract_sied_candidate(same)
            if same_val and _is_valid_sied_value(same_val):
                return same_val
            if i + 1 < len(lines):
                nxt_val = _extract_sied_candidate(lines[i + 1])
                if nxt_val and _is_valid_sied_value(nxt_val):
                    return nxt_val

    # Fallback: any strong SIED-like token in agenda text.
    for pat in [
        r"\b\d{3,5}-\d{2}-\d{3,5}\b",
        r"\b[A-Z]{1,5}\d{4,}(?:-\d{1,6})?\b",
    ]:
        for m in re.finditer(pat, text, re.IGNORECASE):
            cand = _extract_sied_candidate(m.group(0))
            if cand and _is_valid_sied_value(cand):
                return cand
    return ""


def _agenda_widget_value_map_pdf(agenda_path: Path) -> Dict[str, str]:
    """
    Read key Agenda values from named PDF form widgets.
    Returns keys: inventory_number, agent_id, sied_number.
    """
    out = {
        "inventory_number": "",
        "agent_id": "",
        "sied_number": "",
    }
    if (agenda_path.suffix or "").lower() != ".pdf":
        return out
    try:
        doc = fitz.open(str(agenda_path))
        page = doc[0]
        for w in list(page.widgets() or []):
            name = re.sub(r"[^A-Za-z0-9]+", "", clean_value(str(w.field_name or ""))).upper()
            val = clean_value(str(w.field_value or ""))
            if not val:
                continue
            if name == "BARCODEETNODINVENTAIRE":
                inv = _normalize_inventory_number(val)
                if inv:
                    out["inventory_number"] = inv
            elif name == "AGENT":
                d = re.sub(r"\D", "", val)
                if d:
                    out["agent_id"] = d[-5:] if len(d) > 5 else d.zfill(5)
            elif name == "SIED":
                sied = _normalize_sied_value(val, allow_raw=True)
                if sied and _is_valid_sied_value(sied):
                    out["sied_number"] = sied
        doc.close()
    except Exception:
        return out
    return out


def extract_agenda_core_values(agenda_path: Path) -> Dict[str, str]:
    """
    Extract inventory / agent / CE-CID from Agenda (manual edits included).
    """
    out = {
        "inventory_number": "",
        "agent_id": "",
        "sied_number": "",
    }
    from_widgets = _agenda_widget_value_map_pdf(agenda_path)
    out.update({k: v for k, v in from_widgets.items() if v})

    text = _agenda_text_all_pages(agenda_path)
    if (not out["inventory_number"]) and text:
        out["inventory_number"] = (
            _extract_inventory_from_text_label_only(text)
            or _extract_inventory_from_text_loose(text)
            or ""
        )
        out["inventory_number"] = _normalize_inventory_number(out["inventory_number"])
    if (not out["agent_id"]) and text:
        m_agent = re.search(r"\bAGENT\s*:?\s*(\d{4,6})\b", text, re.IGNORECASE)
        if m_agent:
            d = re.sub(r"\D", "", m_agent.group(1))
            out["agent_id"] = d[-5:] if len(d) > 5 else d.zfill(5)
    if not out["sied_number"]:
        sied = _normalize_sied_value(extract_sied_from_agenda(agenda_path), allow_raw=True)
        if sied and _is_valid_sied_value(sied):
            out["sied_number"] = sied

    # Guardrail: CE/CID must not be identical to inventory (exact match only).
    # Substring checks were too aggressive — a CE/CID like 9823-43-2 would get
    # falsely cleared because its compact form "9823432" appears inside "W09823432".
    if out["sied_number"] and out["inventory_number"]:
        sied_tok = _compact_alnum_token(out["sied_number"])
        inv_tok = _compact_alnum_token(out["inventory_number"])
        if sied_tok == inv_tok:
            out["sied_number"] = ""
    return out


def _read_sied_raw_from_agenda_pdf(agenda_path: Path) -> str:
    """
    Read whatever the user typed in the SIED field of the Agenda PDF — no format
    validation, no pattern matching.  Returns the raw cleaned string or "".
    Falls back to the positional SIED-rect text if no named widget exists.
    """
    if (agenda_path.suffix or "").lower() != ".pdf" or not agenda_path.exists():
        return ""
    try:
        doc = fitz.open(str(agenda_path))
        page = doc[0]
        # 1) Named "SIED" widget (most reliable)
        for w in list(page.widgets() or []):
            nm = re.sub(r"[^A-Za-z0-9]+", "", clean_value(str(w.field_name or ""))).upper()
            if nm == "SIED":
                val = clean_value(str(w.field_value or ""))
                doc.close()
                return val
        # 2) Positional fallback: read text from the SIED rectangle area
        try:
            _agent_rect, _barcode_rect, _inv_pt, sied_rect = _agenda_rects_for_page(page)
            rect = fitz.Rect(
                max(0, sied_rect.x0 - 4),
                max(0, sied_rect.y0 - 3),
                min(page.rect.width, sied_rect.x1 + 60),
                min(page.rect.height, sied_rect.y1 + 20),
            )
            words = [
                w[4] for w in page.get_text("words")
                if fitz.Rect(w[:4]).intersects(rect)
            ]
            doc.close()
            return clean_value(" ".join(words))
        except Exception:
            pass
        doc.close()
    except Exception:
        pass
    return ""


def sync_agenda_files(case_paths: Dict[str, Path]) -> Optional[Path]:
    """
    Keep client/internal Agenda files synchronized.
    If both exist, newest file wins and is copied over the older one.
    Returns best current Agenda path (prefer client path when present).
    """
    pairs = [
        ("agenda_output_pdf", "agenda_latest_pdf"),
        ("agenda_output_docx", "agenda_latest_docx"),
    ]
    best_path: Optional[Path] = None
    for client_key, internal_key in pairs:
        client = case_paths.get(client_key)
        internal = case_paths.get(internal_key)
        if (not client) or (not internal):
            continue
        c_exists = client.exists()
        i_exists = internal.exists()
        if c_exists and i_exists:
            try:
                c_m = client.stat().st_mtime
                i_m = internal.stat().st_mtime
                newer = client if c_m >= i_m else internal
                older = internal if newer is client else client
                if newer.resolve() != older.resolve():
                    older.parent.mkdir(parents=True, exist_ok=True)
                    shutil.copy2(newer, older)
            except Exception:
                pass
            best_path = client
        elif c_exists:
            try:
                internal.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(client, internal)
            except Exception:
                pass
            best_path = client
        elif i_exists:
            try:
                client.parent.mkdir(parents=True, exist_ok=True)
                shutil.copy2(internal, client)
            except Exception:
                pass
            best_path = client if client.exists() else internal
    return best_path


def resolve_latest_k138_pdf(working_dir: Path, case_paths: Dict[str, Path]) -> Optional[Path]:
    """Prefer case storage k138 files, fallback to client K138 filename(s) in working dir."""
    latest = case_paths["k138_latest_pdf"]
    if latest.exists():
        return latest
    client = case_paths.get("client_k138_pdf")
    if client and client.exists():
        return client
    out = case_paths.get("k138_output_pdf")
    if out and out.exists():
        return out
    plain = working_dir / "K138.pdf"
    if plain.exists():
        return plain
    candidates = sorted(
        list(case_paths["case_root"].glob("K138_*.pdf"))
        + list(case_paths["case_root"].glob("K138.pdf"))
        + list(case_paths["k138_dir"].glob("K138_*.pdf"))
        + list(case_paths["k138_dir"].glob("K138.pdf"))
        + list(working_dir.glob("K138_*.pdf")),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    return candidates[0] if candidates else None


def generate_barcode(inventory_number: str, out_png: Path) -> Path:
    """
    Generate Code128 barcode image from inventory_number only.
    """
    if not HAVE_BARCODE:
        raise RuntimeError("Barcode library not installed. Please install 'python-barcode' and 'Pillow'.")
    inv = _normalize_inventory_number(inventory_number)
    if not inv:
        raise ValueError("Inventory number is empty; cannot generate barcode.")
    out_png.parent.mkdir(parents=True, exist_ok=True)
    barcode_obj = Code128(inv, writer=ImageWriter())
    saved = Path(
        barcode_obj.save(
            str(out_png.with_suffix("")),
            options={
                "write_text": False,
                "module_width": 0.34,
                "module_height": 4.8,
                "quiet_zone": 2.0,
            },
        )
    )
    if saved.resolve() != out_png.resolve():
        shutil.copy2(saved, out_png)
        try:
            saved.unlink(missing_ok=True)
        except Exception:
            pass
    return out_png


def _find_first_label_rect(page: "fitz.Page", terms: List[str]) -> Optional["fitz.Rect"]:
    for term in terms:
        try:
            rects = page.search_for(term)
            if rects:
                return rects[0]
        except Exception:
            continue
    return None


def _agenda_rects_for_page(page: "fitz.Page") -> Tuple["fitz.Rect", "fitz.Rect", Tuple[float, float], "fitz.Rect"]:
    """
    Detect AGENT and barcode rectangles by labels; fall back to conservative defaults.
    Returns: (agent_value_rect, barcode_rect, inventory_text_point, sied_value_rect)
    """
    pr = page.rect
    agent_label = _find_first_label_rect(page, ["AGENT:", "AGENT :", "AGENT"])
    if agent_label:
        agent_rect = fitz.Rect(
            agent_label.x1 + 18,
            max(0, agent_label.y0 - 8),
            min(pr.width - 20, agent_label.x1 + 195),
            min(pr.height, agent_label.y1 + 2),
        )
    else:
        agent_rect = fitz.Rect(162, 110, 326, 132)

    barcode_label = _find_first_label_rect(
        page,
        [
            "APPOSER CODE BARRE",
            "APPOSER CODE-BARRE",
            "CODE BARRE",
            "BARCODE",
        ],
    )
    if barcode_label:
        left = max(20, barcode_label.x0 - 40)
        # Place barcode above the "APPOSER CODE BARRE" label line.
        bar_h = 20
        gap_above_label = -4
        top = max(0, barcode_label.y0 - (bar_h + gap_above_label))
        width = min(225, pr.width - left - 20)
        barcode_rect = fitz.Rect(left, top, left + width, top + bar_h)
        # Keep inventory text clearly separated below barcode.
        inv_y = min(pr.height - 8, barcode_rect.y1 + 32)
    else:
        left = pr.width * 0.49
        top = pr.height * 0.575
        barcode_rect = fitz.Rect(left, top, min(pr.width - 20, left + 225), min(pr.height - 20, top + 20))
        inv_y = min(pr.height - 8, barcode_rect.y1 + 32)

    sied_label = _find_first_label_rect(page, ["SIED:", "SIED :", "SIED"])
    if sied_label:
        sied_rect = fitz.Rect(
            min(pr.width - 20, sied_label.x1 + 16),
            max(0, sied_label.y0 - 8),
            min(pr.width - 20, sied_label.x1 + 292),
            min(pr.height, sied_label.y1 + 8),
        )
    else:
        sied_rect = fitz.Rect(277, 147, 517, 168)

    inv_pt = (barcode_rect.x0, inv_y)
    return agent_rect, barcode_rect, inv_pt, sied_rect


def _hard_clear_rect(page: "fitz.Page", rect: "fitz.Rect") -> None:
    """
    Remove existing content in rect so stale hidden text is not re-extracted later.
    Falls back to white overlay when redaction is unavailable.
    """
    try:
        page.add_redact_annot(rect, fill=(1, 1, 1))
        try:
            page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)
        except Exception:
            page.apply_redactions()
    except Exception:
        page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1), overlay=True)


AGENDA_SIED_Y_OFFSET_DEFAULT = -2
AGENDA_SIED_Y_OFFSET_UPDATE = -6  # Keep update placement a little higher than initial fill.


def _rect_overlap_area_fitz(a: "fitz.Rect", b: "fitz.Rect") -> float:
    ix0 = max(a.x0, b.x0)
    iy0 = max(a.y0, b.y0)
    ix1 = min(a.x1, b.x1)
    iy1 = min(a.y1, b.y1)
    if ix1 <= ix0 or iy1 <= iy0:
        return 0.0
    return (ix1 - ix0) * (iy1 - iy0)


def _rect_center_distance_fitz(a: "fitz.Rect", b: "fitz.Rect") -> float:
    acx = (a.x0 + a.x1) / 2.0
    acy = (a.y0 + a.y1) / 2.0
    bcx = (b.x0 + b.x1) / 2.0
    bcy = (b.y0 + b.y1) / 2.0
    return ((acx - bcx) ** 2 + (acy - bcy) ** 2) ** 0.5


def _agenda_widget_regions(
    page: "fitz.Page",
    agent_rect: "fitz.Rect",
    barcode_rect: "fitz.Rect",
    inv_pt: Tuple[float, float],
    sied_rect: "fitz.Rect",
) -> Dict[str, "fitz.Rect"]:
    inv_text_y = inv_pt[1]
    inventory_rect = fitz.Rect(
        max(0, barcode_rect.x0 - 10),
        max(0, inv_text_y - 14),
        min(page.rect.width, barcode_rect.x1 + 10),
        min(page.rect.height, inv_text_y + 10),
    )
    return {
        "agent": fitz.Rect(agent_rect),
        "sied": fitz.Rect(sied_rect),
        "inventory": inventory_rect,
    }


def _is_checklist_like_agenda_widget(norm_name: str) -> bool:
    """Detect checklist/checkbox text widgets that must never receive inventory/agent values."""
    if not norm_name:
        return False
    tokens = [
        "IMPORT",
        "EXPORT",
        "COCHEZ",
        "CANNABIS",
        "PHOTOGRAPH",
        "SACDEPREUVE",
        "TABLEAU",
        "CSFRM",
        "K9",
        "ETIQUETTE",
        "REMETTRE",
        "DESTINATAIRE",
        "EXPEDITEUR",
        "NUMERODEPOSTEREC",
        "DATE",
        "ILLISIBLE",
        "NONENVOYE",
        "REGMAIL",
        "TEXTFIELD",
    ]
    return any(t in norm_name for t in tokens)


def _agenda_pdf_has_inventory_leak(agenda_path: Path, inventory_number: str) -> bool:
    """True when checklist widgets contain inventory-like text due old mapping bug."""
    if (agenda_path.suffix or "").lower() != ".pdf":
        return False
    inv_norm = _normalize_inventory_number(inventory_number or "")
    inv_compact = _compact_alnum_token(inv_norm)
    try:
        doc = fitz.open(str(agenda_path))
        page = doc[0]
        for w in list(page.widgets() or []):
            name = re.sub(r"[^A-Za-z0-9]+", "", clean_value(str(w.field_name or ""))).upper()
            if not _is_checklist_like_agenda_widget(name):
                continue
            cur = clean_value(str(w.field_value or ""))
            if not cur:
                continue
            cur_compact = _compact_alnum_token(cur)
            if inv_compact and cur_compact and (cur_compact == inv_compact):
                doc.close()
                return True
            if _looks_like_inventory_number_strict(cur):
                doc.close()
                return True
        doc.close()
    except Exception:
        return False
    return False


def _fill_agenda_pdf_widgets(
    page: "fitz.Page",
    agent_rect: "fitz.Rect",
    barcode_rect: "fitz.Rect",
    inv_pt: Tuple[float, float],
    sied_rect: "fitz.Rect",
    agent_id: Optional[str] = None,
    inventory_number: Optional[str] = None,
    sied_number: Optional[str] = None,
) -> set[str]:
    """
    Populate existing form widgets by geometry so Agenda remains fillable.
    Returns which logical fields were written: {'agent','sied','inventory'}.
    """
    widgets = list(page.widgets() or [])
    if not widgets:
        return set()

    values: Dict[str, str] = {}
    if agent_id is not None:
        values["agent"] = clean_value(agent_id)
    if inventory_number is not None:
        values["inventory"] = _normalize_inventory_number(inventory_number)
    if sied_number is not None:
        values["sied"] = _normalize_sied_value(sied_number, allow_raw=True)
    if not values:
        return set()

    regions = _agenda_widget_regions(page, agent_rect, barcode_rect, inv_pt, sied_rect)
    name_map = {
        "AGENT": "agent",
        "SIED": "sied",
        "BARCODEETNODINVENTAIRE": "inventory",
    }
    protected_field_names = set(name_map.keys())
    text_type = int(getattr(fitz, "PDF_WIDGET_TYPE_TEXT", 7))
    candidate_values = {clean_value(v).upper() for v in values.values() if clean_value(v)}
    written: set[str] = set()
    for widget in widgets:
        try:
            if int(getattr(widget, "field_type", 0) or 0) != text_type:
                continue
        except Exception:
            continue
        wrect = widget.rect
        raw_name = clean_value(str(widget.field_name or ""))
        norm_name = re.sub(r"[^A-Za-z0-9]+", "", raw_name).upper()
        best_key = name_map.get(norm_name)
        if not best_key:
            # Strict fallback for legacy templates:
            # require strong geometric overlap (distance-only matching caused false writes).
            best_overlap_ratio = 0.0
            best_candidate = None
            w_area = max(1.0, float((wrect.x1 - wrect.x0) * (wrect.y1 - wrect.y0)))
            for key, region in regions.items():
                overlap = _rect_overlap_area_fitz(wrect, region)
                if overlap <= 0:
                    continue
                ratio = overlap / w_area
                if ratio > best_overlap_ratio:
                    best_overlap_ratio = ratio
                    best_candidate = key
            if best_candidate and best_overlap_ratio >= 0.85:
                best_key = best_candidate
            else:
                best_key = None
        if (not best_key) or (best_key not in values):
            # Clean up stale bad writes from earlier versions (inventory leaked into checklist boxes).
            if norm_name not in protected_field_names and _is_checklist_like_agenda_widget(norm_name):
                cur = clean_value(str(widget.field_value or ""))
                cur_up = cur.upper()
                should_clear = bool(
                    cur
                    and (
                        (cur_up in candidate_values)
                        or _looks_like_inventory_number_strict(cur)
                    )
                )
                if should_clear:
                    try:
                        flags = int(getattr(widget, "field_flags", 0) or 0)
                        if flags & 1:
                            widget.field_flags = flags & ~1
                        # PyMuPDF keeps previous value when set to empty string;
                        # single space reliably clears visual content.
                        widget.field_value = " "
                        widget.update()
                    except Exception:
                        pass
            continue

        try:
            flags = int(getattr(widget, "field_flags", 0) or 0)
            if flags & 1:
                widget.field_flags = flags & ~1
            val_to_write = values[best_key]
            if not clean_value(val_to_write):
                # Empty string can be ignored by some widgets; force a visual blank.
                val_to_write = " "
            widget.field_value = val_to_write
            widget.update()
            written.add(best_key)
        except Exception:
            continue
    return written


def _write_agenda_pdf_fields(
    agenda_pdf: Path,
    agent_id: str,
    inventory_number: str,
    sied_number: str,
    barcode_png: Path,
    clear_first: bool,
) -> None:
    """Write agent, barcode, inventory and SIED values into agenda PDF and save in-place.

    The save strategy avoids the file-lock / partial-write problem on Windows:
    1. Always save to a sibling .tmp file first (doc is fully closed before any rename).
    2. Only after doc.close() do we atomically replace the original with os.replace().
    3. saveIncr() is never used because it keeps the file handle open and can leave
       the PDF in an inconsistent state when the process is killed mid-write.
    """
    tmp_path: Optional[Path] = agenda_pdf.with_name(f"{agenda_pdf.stem}._saving_tmp{agenda_pdf.suffix}")
    # Remove any stale temp from a previous crashed run.
    if tmp_path.exists():
        try:
            tmp_path.unlink()
        except Exception:
            pass

    doc = fitz.open(str(agenda_pdf))
    try:
        page = doc[0]
        agent_rect, barcode_rect, inv_pt, sied_rect = _agenda_rects_for_page(page)

        if clear_first:
            _hard_clear_rect(page, agent_rect)
            _hard_clear_rect(
                page,
                fitz.Rect(barcode_rect.x0, barcode_rect.y0, barcode_rect.x1, min(page.rect.height, barcode_rect.y1 + 28)),
            )
            _hard_clear_rect(page, sied_rect)

        sied_to_write = _normalize_sied_value(sied_number, allow_raw=True)
        written_widgets = _fill_agenda_pdf_widgets(
            page,
            agent_rect,
            barcode_rect,
            inv_pt,
            sied_rect,
            agent_id=agent_id,
            inventory_number=inventory_number,
            sied_number=sied_to_write,
        )

        if clean_value(agent_id) and ("agent" not in written_widgets):
            page.insert_text((agent_rect.x0 + 10, agent_rect.y1 - 8), agent_id, fontsize=11, fontname="helv", color=(0, 0, 0))
        if sied_to_write and ("sied" not in written_widgets):
            page.insert_text(
                (sied_rect.x0, sied_rect.y1 + AGENDA_SIED_Y_OFFSET_DEFAULT),
                sied_to_write,
                fontsize=11,
                fontname="helv",
                color=(0, 0, 0),
            )
        page.insert_image(barcode_rect, filename=str(barcode_png), keep_proportion=False, overlay=True)
        if inventory_number:
            inv_text = _normalize_inventory_number(inventory_number)
            if inv_text and ("inventory" not in written_widgets):
                txt_w = fitz.get_text_length(inv_text, fontname="helv", fontsize=10)
                inv_x = barcode_rect.x0 + max(0.0, (barcode_rect.width - txt_w) / 2.0)
                if inv_x + txt_w > barcode_rect.x1:
                    inv_x = max(barcode_rect.x0, barcode_rect.x1 - txt_w)
                page.insert_text((inv_x, inv_pt[1]), inv_text, fontsize=10, fontname="helv", color=(0, 0, 0))

        # Save to temp file while doc is still in scope (needed for the data).
        doc.save(str(tmp_path), incremental=False, encryption=fitz.PDF_ENCRYPT_KEEP)
    finally:
        # Always close the original doc BEFORE touching the filesystem paths.
        doc.close()

    # Now that the file handle is fully released, atomically replace the original.
    os.replace(str(tmp_path), str(agenda_pdf))
    # Cleanup: tmp should be gone after replace, but guard anyway.
    if tmp_path.exists():
        try:
            tmp_path.unlink()
        except Exception:
            pass


def fill_agenda_pdf(
    template_path: Path,
    output_path: Path,
    agent_id: str,
    inventory_number: str,
    sied_number: str,
    barcode_png: Path,
) -> Path:
    """Create new agenda_latest.pdf from template and fill AGENT + barcode area only."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output_path)
    _write_agenda_pdf_fields(output_path, agent_id, inventory_number, sied_number, barcode_png, clear_first=True)
    return output_path


def update_agenda_pdf(
    agenda_path: Path,
    agent_id: str,
    inventory_number: str,
    sied_number: str,
    barcode_png: Path,
) -> Path:
    """Update existing agenda_latest.pdf by refreshing only AGENT and barcode area."""
    if not agenda_path.exists():
        raise FileNotFoundError(f"Agenda not found: {agenda_path}")
    _write_agenda_pdf_fields(agenda_path, agent_id, inventory_number, sied_number, barcode_png, clear_first=True)
    return agenda_path


_DOCX_BARCODE_MARKER = "[RADIANCE_BARCODE]"


def _remove_docx_marker_paragraphs(doc: "DocxDocument") -> None:
    for p in list(doc.paragraphs):
        if _DOCX_BARCODE_MARKER in (p.text or ""):
            p._element.getparent().remove(p._element)


def _set_agent_in_docx(doc: "DocxDocument", agent_id: str) -> None:
    for p in doc.paragraphs:
        txt = p.text or ""
        if re.search(r"\bAGENT\b", txt, flags=re.IGNORECASE):
            if ":" in txt:
                p.text = re.sub(r"(AGENT\s*:?\s*).*", rf"\1{agent_id}", txt, flags=re.IGNORECASE)
            else:
                p.text = f"{txt} {agent_id}".strip()
            return
    doc.add_paragraph(f"AGENT: {agent_id}")


def _set_sied_in_docx(doc: "DocxDocument", sied_number: str) -> None:
    val = clean_value(sied_number)
    for p in doc.paragraphs:
        txt = p.text or ""
        if re.search(r"\bSIED\b", txt, flags=re.IGNORECASE):
            if ":" in txt:
                p.text = re.sub(r"(SIED\s*:?\s*).*", rf"\1{val}", txt, flags=re.IGNORECASE)
            else:
                p.text = f"SIED: {val}"
            return
    doc.add_paragraph(f"SIED: {val}")


def fill_agenda_docx(
    template_path: Path,
    output_path: Path,
    agent_id: str,
    inventory_number: str,
    sied_number: str,
    barcode_png: Path,
) -> Path:
    if not HAVE_DOCX:
        raise RuntimeError("python-docx is required for DOCX agenda templates.")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output_path)
    doc = DocxDocument(str(output_path))
    _set_agent_in_docx(doc, agent_id)
    if clean_value(sied_number):
        _set_sied_in_docx(doc, sied_number)
    _remove_docx_marker_paragraphs(doc)
    p = doc.add_paragraph(f"{_DOCX_BARCODE_MARKER} {inventory_number}")
    run = p.add_run("\n")
    run.add_picture(str(barcode_png), width=Inches(2.0))
    doc.save(str(output_path))
    return output_path


def update_agenda_docx(
    agenda_path: Path,
    agent_id: str,
    inventory_number: str,
    sied_number: str,
    barcode_png: Path,
) -> Path:
    if not HAVE_DOCX:
        raise RuntimeError("python-docx is required for DOCX agenda templates.")
    if not agenda_path.exists():
        raise FileNotFoundError(f"Agenda not found: {agenda_path}")
    doc = DocxDocument(str(agenda_path))
    _set_agent_in_docx(doc, agent_id)
    if clean_value(sied_number):
        _set_sied_in_docx(doc, sied_number)
    _remove_docx_marker_paragraphs(doc)
    p = doc.add_paragraph(f"{_DOCX_BARCODE_MARKER} {inventory_number}")
    run = p.add_run("\n")
    run.add_picture(str(barcode_png), width=Inches(2.0))
    doc.save(str(agenda_path))
    return agenda_path


def fill_agenda(
    template_path: Path,
    case_paths: Dict[str, Path],
    agent_id: str,
    inventory_number: str,
    sied_number: str,
    barcode_png: Path,
) -> Path:
    """Dispatch agenda fill by template extension."""
    suffix = template_path.suffix.lower()
    if suffix == ".pdf":
        return fill_agenda_pdf(template_path, case_paths["agenda_latest_pdf"], agent_id, inventory_number, sied_number, barcode_png)
    if suffix == ".docx":
        return fill_agenda_docx(template_path, case_paths["agenda_latest_docx"], agent_id, inventory_number, sied_number, barcode_png)
    raise RuntimeError(f"Unsupported agenda template: {template_path.name}")


def update_agenda(
    agenda_path: Path,
    agent_id: str,
    inventory_number: str,
    sied_number: str,
    barcode_png: Path,
) -> Path:
    """Dispatch agenda update by file extension."""
    suffix = agenda_path.suffix.lower()
    if suffix == ".pdf":
        return update_agenda_pdf(agenda_path, agent_id, inventory_number, sied_number, barcode_png)
    if suffix == ".docx":
        return update_agenda_docx(agenda_path, agent_id, inventory_number, sied_number, barcode_png)
    raise RuntimeError(f"Unsupported agenda file: {agenda_path.name}")


def update_agenda_pdf_sied_only(agenda_path: Path, sied_number: str) -> Path:
    """Update only SIED value in an existing agenda PDF (in place)."""
    if not agenda_path.exists():
        raise FileNotFoundError(f"Agenda not found: {agenda_path}")
    doc = fitz.open(str(agenda_path))
    page = doc[0]
    agent_rect, barcode_rect, inv_pt, sied_rect = _agenda_rects_for_page(page)

    _hard_clear_rect(page, sied_rect)
    val = _normalize_sied_value(sied_number, allow_raw=True)
    written_widgets = _fill_agenda_pdf_widgets(
        page,
        agent_rect,
        barcode_rect,
        inv_pt,
        sied_rect,
        sied_number=val,
    )
    if val and ("sied" not in written_widgets):
        page.insert_text(
            (sied_rect.x0, sied_rect.y1 + AGENDA_SIED_Y_OFFSET_UPDATE),
            val,
            fontsize=11,
            fontname="helv",
            color=(0, 0, 0),
        )

    tmp_path: Optional[Path] = None
    try:
        doc.saveIncr()
    except Exception:
        tmp_path = agenda_path.with_name(f"{agenda_path.stem}.tmp{agenda_path.suffix}")
        if tmp_path.exists():
            tmp_path.unlink()
        doc.save(str(tmp_path), incremental=False, encryption=fitz.PDF_ENCRYPT_KEEP)
        os.replace(str(tmp_path), str(agenda_path))
    finally:
        doc.close()
        if tmp_path and tmp_path.exists():
            try:
                tmp_path.unlink()
            except Exception:
                pass
    return agenda_path


def update_agenda_docx_sied_only(agenda_path: Path, sied_number: str) -> Path:
    """Update only SIED value in an existing agenda DOCX (in place)."""
    if not HAVE_DOCX:
        raise RuntimeError("python-docx is required for DOCX agenda templates.")
    if not agenda_path.exists():
        raise FileNotFoundError(f"Agenda not found: {agenda_path}")
    doc = DocxDocument(str(agenda_path))
    _set_sied_in_docx(doc, sied_number)
    doc.save(str(agenda_path))
    return agenda_path


def update_agenda_sied_only(agenda_path: Path, sied_number: str) -> Path:
    """Dispatch SIED-only agenda update by file extension."""
    suffix = agenda_path.suffix.lower()
    if suffix == ".pdf":
        return update_agenda_pdf_sied_only(agenda_path, sied_number)
    if suffix == ".docx":
        return update_agenda_docx_sied_only(agenda_path, sied_number)
    raise RuntimeError(f"Unsupported agenda file: {agenda_path.name}")

def try_run_k138_filler(
    k138_template: Path,
    k138_values_csv: Path,
    output_path: Path,
    log_fn,
) -> bool:
    """
    Call fill_k138_notice.py to fill the K138 form.
    Notice text is already included in the CSV via build_k138_values_from_saisie.
    Output is saved to output_path (must be in working case folder, NOT templates).
    CWD is switched to temp folder during fill to avoid leaking relative-path writes.
    """
    try:
        from fill_k138_notice import fill_k138

        log_fn(f"Filling K138 form: {k138_template.name}")

        # Use absolute paths so chdir doesn't affect them
        template_path = str(k138_template.resolve())
        csv_path = str(k138_values_csv.resolve())
        out_path = str(output_path.resolve())

        cwd_before = os.getcwd()
        try:
            os.chdir(str(k138_values_csv.parent))  # temp folder
            fill_k138(
                template_path=template_path,
                output_path=out_path,
                use_csv=True,
                csv_path=csv_path,
            )
        finally:
            os.chdir(cwd_before)
        return True
    except ModuleNotFoundError as e:
        missing = (getattr(e, "name", "") or "").strip()
        if missing == "fill_k138_notice":
            log_fn(
                "K138 helper module missing: fill_k138_notice.py "
                f"(expected: {helper_module_hint('fill_k138_notice')})"
            )
        else:
            log_fn(f"K138 filler missing dependency: {missing or e}")
        return False
    except Exception as e:
        log_fn(f"K138 filler failed: {e}")
        import traceback
        log_fn(traceback.format_exc())
        return False


# ======================== =
# Tkinter app -----
# ======================== =

@dataclass
class AppState:
    templates_folder: Optional[Path] = None  # Templates folder (contains templates and .txt files)
    saisie_pdf_file: Optional[Path] = None  # Individual SAISIE input file to process
    working_dir: Optional[Path] = None
    case_folder_name: str = ""
    case_folder_locked: bool = False
    badge_number: str = ""
    form_type: str = "Stupefiant-Others"  # Default form type
    clerk_agenda_file: Optional[Path] = None  # Clerk-selected agenda PDF (Clerk workflow)

    last_top: Optional[Dict[str, str]] = None
    last_bottom: Optional[List[Dict[str, str]]] = None
    last_saisie_csv: Optional[Path] = None
    last_k138_values_csv: Optional[Path] = None
    last_k138_output: Optional[Path] = None
    last_inventory_number: str = ""
    last_agent_id: str = ""
    last_sied_number: str = ""
    last_k138_values: Optional[Dict[str, str]] = None
    # Set True only after extraction has run for the CURRENT saisie_pdf_file.
    # Prevents values_latest.json from a prior session bleeding into a freshly loaded file.
    extraction_ran: bool = False
    saisie_affaire_generated: bool = False

class AppBase:
    def __init__(self, root: tk.Tk, profile_role: str = "BSO", profile_badge: str = ""):
        self.root = root
        self.state = AppState()
        self.profile_role  = profile_role  or "BSO"
        self.profile_badge = profile_badge or ""

        # Pre-fill badge number from splash selection for BSO
        if self.profile_role == "BSO" and self.profile_badge:
            self.state.badge_number = self.profile_badge

        self.root.title(f"Radiance Copilot — {self.profile_role}")
        app_ico_path = resolve_asset_path("photos/Radiance-copilot-icon.ico", "Radiance-copilot-icon.ico")
        app_png_path = resolve_asset_path("photos/Radiance-copilot-icon.png", "Radiance-copilot-icon.png")
        try:
            if app_ico_path:
                self.root.iconbitmap(app_ico_path)
        except Exception:
            pass
        try:
            if app_png_path:
                self.root.iconphoto(True, tk.PhotoImage(file=app_png_path))
        except Exception:
            pass
        self.root.configure(bg="#F5F7FA")
        # Allow users to resize window in both directions.
        self.root.resizable(True, True)
        self.root.minsize(780, 600)
        self._banner_images: List[tk.PhotoImage] = []
        self._missing_helper_warnings_shown: set[str] = set()
        self._init_styles()

        # Scrollable outer frame so the UI is accessible at any window size
        outer = tk.Frame(root, bg="#F5F7FA")
        outer.grid(row=0, column=0, sticky="nsew")
        root.columnconfigure(0, weight=1)
        root.rowconfigure(0, weight=1)
        outer.columnconfigure(0, weight=1)
        outer.rowconfigure(0, weight=1)

        canvas = tk.Canvas(outer, bg="#F5F7FA", highlightthickness=0)
        vsb = ttk.Scrollbar(outer, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        canvas.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        outer.columnconfigure(0, weight=1)
        outer.rowconfigure(0, weight=1)

        frm = ttk.Frame(canvas, padding=12)
        frm_id = canvas.create_window((0, 0), window=frm, anchor="nw")

        def _on_frm_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
        def _on_canvas_configure(event):
            canvas.itemconfig(frm_id, width=event.width)
        frm.bind("<Configure>", _on_frm_configure)
        canvas.bind("<Configure>", _on_canvas_configure)

        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        self.root.protocol("WM_DELETE_WINDOW", self._on_app_close)

        frm.columnconfigure(0, weight=1)

        # --- Header banner (Radiance-only) ---
        banner_bg = getattr(self, "_accent_color", "#1F4E79")
        banner = tk.Frame(frm, bg=banner_bg, height=148, bd=0, highlightthickness=0)
        banner.grid(row=0, column=0, sticky="ew")
        banner.grid_columnconfigure(0, weight=1)   # title/logo — takes all spare width
        banner.grid_columnconfigure(1, weight=0)   # buttons — fixed width, never overlaps title
        middle_wrap = tk.Frame(banner, bg=banner_bg, bd=0, highlightthickness=0)
        middle_wrap.grid(row=0, column=0, rowspan=2, padx=(16, 8), pady=(10, 10), sticky="nsew")
        middle_wrap.grid_columnconfigure(1, weight=1)
        center_logo_wrap = tk.Frame(middle_wrap, bg=banner_bg, bd=0, highlightthickness=0)
        center_logo_wrap.grid(row=0, column=0, rowspan=2, padx=(0, 10), sticky="w")
        title_wrap = tk.Frame(middle_wrap, bg=banner_bg, bd=0, highlightthickness=0)
        title_wrap.grid(row=0, column=1, rowspan=2, sticky="w")

        def _load_banner_image(path: str, target_height: int, max_width: int = 0, transparent_white: bool = False):
            if not path:
                return None
            if HAVE_PIL and Image and ImageTk:
                try:
                    pil_img = Image.open(path)
                    if transparent_white:
                        rgba = pil_img.convert("RGBA")
                        px = list(rgba.getdata())
                        cleaned = []
                        for r, g, b, a in px:
                            # Remove near-white background so logos blend with blue banner.
                            if r >= 245 and g >= 245 and b >= 245:
                                cleaned.append((r, g, b, 0))
                            else:
                                cleaned.append((r, g, b, a))
                        rgba.putdata(cleaned)
                        pil_img = rgba
                    w, h = pil_img.size
                    if w <= 0 or h <= 0:
                        return None
                    scale = min(1.0, float(target_height) / float(h))
                    if max_width > 0:
                        scale = min(scale, float(max_width) / float(w))
                    if scale < 1.0:
                        new_size = (max(1, int(w * scale)), max(1, int(h * scale)))
                        if hasattr(Image, "Resampling"):
                            pil_img = pil_img.resize(new_size, Image.Resampling.LANCZOS)
                        else:
                            pil_img = pil_img.resize(new_size, Image.LANCZOS)
                    tk_img = ImageTk.PhotoImage(pil_img)
                    self._banner_images.append(tk_img)
                    return tk_img
                except Exception:
                    return None
            try:
                tk_img = tk.PhotoImage(file=path)
                self._banner_images.append(tk_img)
                return tk_img
            except Exception:
                return None

        def _add_banner_logo(
            parent: tk.Widget,
            path: str,
            target_height: int = 56,
            max_width: int = 360,
            transparent_white: bool = True,
        ) -> bool:
            if not path:
                return False
            logo = _load_banner_image(
                path,
                target_height=target_height,
                max_width=max(max_width, target_height * 3),
                transparent_white=transparent_white,
            )
            if not logo:
                return False
            tk.Label(
                parent,
                image=logo,
                bg=banner_bg,
                bd=0,
                highlightthickness=0,
            ).pack(side="left")
            return True

        radiance_logo_path = resolve_asset_path("photos/Radiance-copilot-icon.png", "Radiance-copilot-icon.png")
        if not _add_banner_logo(center_logo_wrap, radiance_logo_path, target_height=58, max_width=210):
            tk.Label(
                center_logo_wrap,
                text="Radiance",
                font=("Segoe UI Semibold", 10),
                fg="#FFFFFF",
                bg=banner_bg,
                anchor="w",
            ).pack(side="left")

        tk.Label(
            title_wrap,
            text="Government of Canada | CBSA",
            font=("Segoe UI Semibold", 14),
            fg="#FFFFFF",
            bg=banner_bg,
            anchor="w",
        ).pack(anchor="w", pady=(10, 0))
        role  = getattr(self, "profile_role",  "BSO") or "BSO"
        badge = getattr(self, "profile_badge", "")   or ""
        profile_line = f"Radiance Copilot   |   Profile: {role}"
        if role == "BSO" and badge:
            profile_line += f"  ({badge})"
        tk.Label(
            title_wrap,
            text=profile_line,
            font=("Segoe UI", 11),
            fg="#DCE8F4",
            bg=banner_bg,
            anchor="w",
        ).pack(anchor="w", pady=(0, 10))

        # Top-right actions — placed in grid column 1 so it never overlaps the title.
        top_actions = tk.Frame(banner, bg=banner_bg, bd=0, highlightthickness=0)
        top_actions.grid(row=0, column=1, sticky="ne", padx=(0, 12), pady=(10, 0))
        self.btnHelp = ttk.Button(
            top_actions,
            text="Help",
            command=self.on_open_help,
            style="Help.TButton",
        )
        self.btnHelp.pack(side="right")
        self.btnConfig = ttk.Button(
            top_actions,
            text="Configurations",
            command=self.on_browse_templates_folder,
            style="HeaderPrimary.TButton",
        )
        self.btnConfig.pack(side="right", padx=(0, 8))
        self.btnChangeProfile = ttk.Button(
            top_actions,
            text="Change Profile",
            command=self.on_change_profile,
            style="Help.TButton",
        )
        self.btnChangeProfile.pack(side="right", padx=(0, 8))
        self.btnExit = ttk.Button(
            top_actions,
            text="Exit",
            command=self.on_exit_session,
            style="Help.TButton",
        )
        self.btnExit.pack(side="right", padx=(0, 8))

        # --- Instruction & Feedback strip (moved to top, compact) ---
        progress_box = ttk.LabelFrame(frm, text="Instruction & Feedback", padding=6)
        progress_box.grid(row=1, column=0, sticky="ew", pady=(10, 0))
        progress_box.columnconfigure(0, weight=1)
        progress_box.rowconfigure(1, weight=1)

        progress_row = ttk.Frame(progress_box)
        progress_row.grid(row=0, column=0, sticky="ew", pady=(0, 4))
        progress_row.columnconfigure(0, weight=1)
        self.varProgressStatus = tk.StringVar(
            value="Configurations: Click the Configurations button (top-right) and set templates folder + badge number."
        )
        ttk.Label(progress_row, textvariable=self.varProgressStatus).grid(row=0, column=0, sticky="w")
        self.progress = ttk.Progressbar(progress_row, mode="indeterminate", length=140)
        self.progress.grid(row=0, column=1, sticky="e")

        self.txt = tk.Text(progress_box, height=6, width=100)
        self.txt.configure(
            bg="#0E1621",
            fg="#E8EDF3",
            insertbackground="#E8EDF3",
            relief="flat",
            padx=8,
            pady=6,
            font=("Consolas", 9),
        )
        self.txt.grid(row=1, column=0, sticky="nsew")

        # Case folder strip directly under instruction banner.
        case_strip = ttk.Frame(frm)
        case_strip.grid(row=2, column=0, sticky="ew", pady=(8, 0))
        case_strip.columnconfigure(1, weight=1)
        ttk.Label(
            case_strip,
            text="Active Case Folder:",
            font=("Segoe UI Semibold", 10),
        ).grid(row=0, column=0, sticky="w", padx=(2, 6))
        self.varCaseFolderBanner = tk.StringVar(value="-")
        ttk.Label(
            case_strip,
            textvariable=self.varCaseFolderBanner,
            font=("Segoe UI Semibold", 10),
        ).grid(row=0, column=1, sticky="w")
        self.btnOpenCaseFolder = ttk.Button(
            case_strip,
            text="Open Folder",
            command=self.on_open_active_case_folder,
            style="Secondary.TButton",
        )
        self.btnOpenCaseFolder.grid(row=0, column=2, sticky="e")

        # --- Configurations summary ---
        cfg_row = ttk.Frame(frm)
        cfg_row.grid(row=3, column=0, sticky="ew", pady=(8, 0))
        cfg_row.columnconfigure(1, weight=1)
        self.varBadgeNumber = tk.StringVar(value="")
        ttk.Label(
            cfg_row,
            text="Configurations:",
            font=("Segoe UI Semibold", 9),
        ).grid(row=0, column=0, sticky="w", padx=(0, 8))
        self.varConfigSummary = tk.StringVar(
            value="Not configured. Use the Configurations button (top-right)."
        )
        ttk.Label(
            cfg_row,
            textvariable=self.varConfigSummary,
            font=("TkDefaultFont", 8),
        ).grid(row=0, column=1, sticky="w")
        
        # --- SAISIE PDF/Word/Image file selector ---
        self.boxSaisieFile = ttk.LabelFrame(frm, text="Select Saisie à Faire", padding=8)
        self.boxSaisieFile.grid(row=4, column=0, sticky="ew", pady=(8, 0))
        self.boxSaisieFile.columnconfigure(0, weight=1)
        self.entSaisieFile = ttk.Entry(self.boxSaisieFile)
        self.entSaisieFile.grid(row=0, column=0, sticky="ew", padx=(0, 8))
        ttk.Button(self.boxSaisieFile, text="Browse...", command=self.on_browse_saisie_file, style="Secondary.TButton").grid(row=0, column=1)
        ttk.Label(self.boxSaisieFile, text="Drag & drop your Saisie à Faire (PDF or Word) here, or use Browse to select it.",
                 font=("TkDefaultFont", 8)).grid(row=1, column=0, columnspan=2, sticky="w", pady=(5, 0))

        # Form type var (selector moved into Saisie D'affaire tab — see below)
        self.form_type_var = tk.StringVar(value=self.state.form_type)

        # Extraction is now automatic — no manual button needed.

        # --- Tabs (Agenda + Saisie d'interet + K138) ---
        self.tabs = ttk.Notebook(frm)
        self.tabs.grid(row=7, column=0, sticky="nsew", pady=(10, 0))
        frm.rowconfigure(7, weight=1)

        self.tabSelectFolder = ttk.Frame(self.tabs)
        self.tabSaisieAffaire = ttk.Frame(self.tabs)
        self.tabAgenda = ttk.Frame(self.tabs)
        self.tabK138 = ttk.Frame(self.tabs)
        self.tabNarrative = ttk.Frame(self.tabs)
        self.tabSaisieInteret = ttk.Frame(self.tabs)
        self.tabs.add(self.tabSelectFolder, text="1: Select Folder")
        self.tabs.add(self.tabSaisieAffaire, text="2: Saisie D'affaire")
        self.tabs.add(self.tabAgenda, text="3: Agenda")
        self.tabs.add(self.tabK138, text="4: K138")
        self.tabs.add(self.tabNarrative, text="5: Narrative")
        self.tabs.add(self.tabSaisieInteret, text="6: Saisie d'interet")
        self.tabs.tab(self.tabSaisieAffaire, state="disabled")
        self.tabs.tab(self.tabAgenda, state="disabled")
        self.tabs.tab(self.tabK138, state="disabled")
        self.tabs.tab(self.tabNarrative, state="disabled")
        self.tabs.tab(self.tabSaisieInteret, state="disabled")

        # Select Folder tab
        self.tabSelectFolder.columnconfigure(1, weight=1)
        self.varSelectedCaseFolder = tk.StringVar(value="-")
        ttk.Label(self.tabSelectFolder, text="Active case folder:").grid(row=0, column=0, sticky="w", padx=(8, 8), pady=(10, 2))
        ttk.Label(self.tabSelectFolder, textvariable=self.varSelectedCaseFolder).grid(row=0, column=1, sticky="w", pady=(10, 2))
        folder_actions = ttk.Frame(self.tabSelectFolder)
        folder_actions.grid(row=1, column=0, columnspan=2, sticky="w", padx=8, pady=(8, 8))
        ttk.Button(
            folder_actions,
            text="Select  (Drag & Drop)",
            command=self.on_select_case_folder,
            style="Secondary.TButton",
        ).grid(row=0, column=0, padx=(0, 8))
        ttk.Button(
            folder_actions,
            text="New",
            command=self.on_create_case_folder,
            style="Primary.TButton",
        ).grid(row=0, column=1, padx=(0, 8))
        ttk.Button(
            folder_actions,
            text="Open Folder",
            command=self.on_open_active_case_folder,
            style="Secondary.TButton",
        ).grid(row=0, column=2)

        # Saisie D'affaire tab
        self.tabSaisieAffaire.columnconfigure(0, weight=1)
        self.varSaisieAffaireFound = tk.StringVar(value="")
        ttk.Label(
            self.tabSaisieAffaire,
            text="Saisie D'affaire (13 main fields + 4 checkboxes)   |   Tip: use ; as line separator for multi-line address fields (EXPEDITEUR, DESTINATAIRE)",
            font=("TkDefaultFont", 9),
        ).grid(row=0, column=0, sticky="w", padx=8, pady=(10, 6))
        self.frmSaisieAffaire = ttk.Frame(self.tabSaisieAffaire)
        self.frmSaisieAffaire.grid(row=1, column=0, sticky="ew", padx=8, pady=(0, 6))
        self.frmSaisieAffaire.columnconfigure(1, weight=1)
        self.frmSaisieAffaire.columnconfigure(3, weight=1)

        self.saisie_affaire_field_specs: List[Tuple[str, str]] = [
            ("bond_room_ledger", "BOND ROOM LEDGER #"),
            ("agent_badge", "INSIGNE AGENT SAISISSANT"),
            ("inventory_number", "INVENTAIRE"),
            ("country", "PAYS"),
            ("interception_location", "LIEU INTERCEPTION"),
            ("interception_datetime", "DATE / HEURE INTERCEPTION"),
            ("declaration", "DÉCLARATION"),
            ("weight_qty", "POIDS / QTÉ MARCH."),
            ("item_description", "DESCRIPTION DE L'ITEM À SAISIR"),
            ("sender", "EXPÉDITEUR"),
            ("recipient", "DESTINATAIRE"),
            ("indices", "INDICES"),
            ("notes", "NOTES"),
        ]
        self.varSaisieAffaireFields: Dict[str, tk.StringVar] = {}
        for idx, (key, label_text) in enumerate(self.saisie_affaire_field_specs):
            row = idx // 2
            col_block = idx % 2
            c0 = 0 if col_block == 0 else 2
            c1 = 1 if col_block == 0 else 3
            ttk.Label(self.frmSaisieAffaire, text=f"{label_text}:").grid(
                row=row, column=c0, sticky="w", padx=(0, 6), pady=2
            )
            initial_value = ""
            if key == "interception_datetime":
                initial_value = datetime.now().strftime("%Y-%m-%d %H:%M")
            elif key == "agent_badge":
                initial_value = re.sub(r"\D", "", self.state.badge_number or self.varBadgeNumber.get() or "")
            var = tk.StringVar(value=initial_value)
            self.varSaisieAffaireFields[key] = var
            ttk.Entry(self.frmSaisieAffaire, textvariable=var).grid(
                row=row, column=c1, sticky="ew", pady=2
            )

        checks_row = ttk.Frame(self.tabSaisieAffaire)
        checks_row.grid(row=2, column=0, sticky="w", padx=8, pady=(4, 2))
        self.varCheckSaisie = tk.BooleanVar(value=False)
        self.varCheckConfiscation = tk.BooleanVar(value=False)
        self.varCheckK9 = tk.BooleanVar(value=False)
        self.varCheckSaisieEnvergure = tk.BooleanVar(value=False)
        self.varCheckLabo = tk.BooleanVar(value=False)
        ttk.Checkbutton(checks_row, text="SAISIE", variable=self.varCheckSaisie).grid(row=0, column=0, padx=(0, 10))
        ttk.Checkbutton(checks_row, text="CONFISCATION", variable=self.varCheckConfiscation).grid(row=0, column=1, padx=(0, 10))
        ttk.Checkbutton(checks_row, text="K9", variable=self.varCheckK9).grid(row=0, column=2, padx=(0, 10))
        ttk.Checkbutton(checks_row, text="SAISIE D'ENVERGURE", variable=self.varCheckSaisieEnvergure).grid(row=0, column=3, padx=(0, 10))
        ttk.Checkbutton(checks_row, text="LABO", variable=self.varCheckLabo, command=self._on_labo_changed).grid(row=0, column=4)

        # Seizure type selector (moved from main frame so Clerk can also see it)
        self.boxFormType = ttk.LabelFrame(self.tabSaisieAffaire, text="Seizure Type (K138 / Narrative)", padding=6)
        self.boxFormType.grid(row=3, column=0, sticky="ew", padx=8, pady=(2, 4))
        ttk.Radiobutton(
            self.boxFormType, text="Cannabis / Stupéfiant",
            variable=self.form_type_var, value="Cannabis-Stupefiant",
            command=self.on_form_type_changed,
        ).grid(row=0, column=0, sticky="w", padx=(0, 20))
        ttk.Radiobutton(
            self.boxFormType, text="Armes / Arms",
            variable=self.form_type_var, value="Knives-Arms",
            command=self.on_form_type_changed,
        ).grid(row=0, column=1, sticky="w", padx=(0, 20))
        ttk.Radiobutton(
            self.boxFormType, text="Autre / Other (default)",
            variable=self.form_type_var, value="Stupefiant-Others",
            command=self.on_form_type_changed,
        ).grid(row=0, column=2, sticky="w")

        # Consistent button order: [Generate] [Update] [Complete ✓] [X]
        saisie_actions = ttk.Frame(self.tabSaisieAffaire)
        saisie_actions.grid(row=4, column=0, sticky="w", padx=8, pady=(0, 4))
        ttk.Button(
            saisie_actions,
            text="Generate",
            command=self.on_generate_saisie_affaire,
            style="Primary.TButton",
        ).grid(row=0, column=0, padx=(0, 6))
        ttk.Button(
            saisie_actions,
            text="Update",
            command=self.on_update_from_saisie_affaire,
            style="Secondary.TButton",
        ).grid(row=0, column=1, padx=(0, 6))
        self.varCompleteSaisie = tk.BooleanVar(value=False)
        ttk.Checkbutton(
            saisie_actions, text="Complete ✓",
            variable=self.varCompleteSaisie,
            command=lambda: self._on_tab_complete_changed("saisie"),
        ).grid(row=0, column=2, padx=(0, 6))
        ttk.Button(
            saisie_actions, text="✕",
            command=lambda: self._on_tab_close("saisie"),
            style="Secondary.TButton", width=3,
        ).grid(row=0, column=3)

        saisie_found_row = ttk.Frame(self.tabSaisieAffaire)
        saisie_found_row.grid(row=5, column=0, sticky="w", padx=8, pady=(0, 8))
        ttk.Label(saisie_found_row, text="Saisie D'affaire in case folder:").grid(row=0, column=0, sticky="w", padx=(0, 6))
        ttk.Label(saisie_found_row, textvariable=self.varSaisieAffaireFound).grid(row=0, column=1, sticky="w")

        # K138 panel (downstream of Agenda)
        self.tabK138.columnconfigure(1, weight=1)
        self.tabK138.rowconfigure(4, weight=1)
        self.varK138AgendaReady = tk.StringVar(value="no")
        self.varK138Sied = tk.StringVar(value="-")
        self.varK138Status = tk.StringVar(value="Process SAISIE and create Agenda first.")

        ttk.Label(self.tabK138, text="CE/CID (from Agenda):").grid(row=0, column=0, sticky="w", padx=(8, 8), pady=(10, 2))
        ttk.Label(self.tabK138, textvariable=self.varK138Sied).grid(row=0, column=1, sticky="w", pady=(10, 2))
        ttk.Label(self.tabK138, textvariable=self.varK138Status).grid(row=1, column=0, columnspan=2, sticky="w", padx=8, pady=(2, 6))

        self.lblK138Gate = ttk.Label(
            self.tabK138,
            text="K138 generation is available only after Agenda is created.",
        )
        self.lblK138Gate.grid(row=3, column=0, columnspan=2, sticky="w", padx=8, pady=(0, 8))

        self.frmK138Actions = ttk.Frame(self.tabK138)
        self.frmK138Actions.grid(row=3, column=0, columnspan=2, sticky="w", padx=8, pady=(0, 8))
        self.btnGenerateK138 = ttk.Button(
            self.frmK138Actions, text="Generate",
            command=self.on_generate_k138, style="Primary.TButton",
        )
        self.btnGenerateK138.grid(row=0, column=0, padx=(0, 6))
        self.btnUpdateK138 = ttk.Button(
            self.frmK138Actions, text="Update",
            command=self._on_refresh_agenda_from_pdf, style="Secondary.TButton",
        )
        self.btnUpdateK138.grid(row=0, column=1, padx=(0, 6))
        self.varCompleteK138 = tk.BooleanVar(value=False)
        ttk.Checkbutton(
            self.frmK138Actions, text="Complete ✓",
            variable=self.varCompleteK138,
            command=lambda: self._on_tab_complete_changed("k138"),
        ).grid(row=0, column=2, padx=(0, 6))
        ttk.Button(
            self.frmK138Actions, text="✕",
            command=lambda: self._on_tab_close("k138"),
            style="Secondary.TButton", width=3,
        ).grid(row=0, column=3)

        # Agenda status panel
        self.tabAgenda.columnconfigure(1, weight=1)
        self.varAgendaWorkingDir = tk.StringVar(value="-")
        self.varAgendaInventory = tk.StringVar(value="-")
        self.varAgendaAgent = tk.StringVar(value="-")
        self.varAgendaFound = tk.StringVar(value="no")

        ttk.Label(self.tabAgenda, text="Inventory number detected:").grid(row=0, column=0, sticky="w", padx=(8, 8), pady=(10, 2))
        ttk.Label(self.tabAgenda, textvariable=self.varAgendaInventory).grid(row=0, column=1, sticky="w", pady=(10, 2))

        ttk.Label(self.tabAgenda, text="Agent detected:").grid(row=1, column=0, sticky="w", padx=(8, 8), pady=2)
        ttk.Label(self.tabAgenda, textvariable=self.varAgendaAgent).grid(row=1, column=1, sticky="w", pady=2)

        ttk.Label(self.tabAgenda, text="Agenda found:").grid(row=2, column=0, sticky="w", padx=(8, 8), pady=2)
        ttk.Label(self.tabAgenda, textvariable=self.varAgendaFound).grid(row=2, column=1, sticky="w", pady=2)

        btnAgendaRow = ttk.Frame(self.tabAgenda)
        btnAgendaRow.grid(row=3, column=0, columnspan=2, sticky="w", padx=8, pady=(10, 8))
        self.btnGenerateAgenda = ttk.Button(
            btnAgendaRow, text="Generate",
            command=self.on_fill_agenda, style="Primary.TButton",
        )
        self.btnGenerateAgenda.grid(row=0, column=0, padx=(0, 6))

        # Update — re-reads Agenda PDF after manual officer edits, regenerates barcode
        self.btnRefreshAgenda = ttk.Button(
            btnAgendaRow, text="Update",
            command=self._on_refresh_agenda_from_pdf,
            state="disabled", style="Secondary.TButton",
        )
        self.btnRefreshAgenda.grid(row=0, column=1, padx=(0, 6))

        self.varCompleteAgenda = tk.BooleanVar(value=False)
        ttk.Checkbutton(
            btnAgendaRow, text="Complete ✓",
            variable=self.varCompleteAgenda,
            command=lambda: self._on_tab_complete_changed("agenda"),
        ).grid(row=0, column=2, padx=(0, 6))
        ttk.Button(
            btnAgendaRow, text="✕",
            command=lambda: self._on_tab_close("agenda"),
            style="Secondary.TButton", width=3,
        ).grid(row=0, column=3, padx=(0, 8))

        # Clerk-only: select case folder to auto-load Saisie D'affaire and enable Agenda + K138
        self.btnClerkSelectAgenda = ttk.Button(
            btnAgendaRow,
            text="Select Case Folder",
            command=self.on_clerk_select_case_folder,
            style="Secondary.TButton",
        )
        self.btnClerkSelectAgenda.grid(row=0, column=4, padx=(0, 8))
        # Show only for Clerk role; hidden for all others
        if getattr(self, "profile_role", "BSO") != "Clerk":
            self.btnClerkSelectAgenda.grid_remove()

        # Auto-refresh agenda status when user switches to the Agenda tab
        self.tabs.bind("<<NotebookTabChanged>>", self._on_tab_changed)

        # Narrative tab (BSO only)
        self.tabNarrative.columnconfigure(0, weight=1)
        self.tabNarrative.rowconfigure(2, weight=1)
        self.varNarrativeStatus = tk.StringVar(value="")
        ttk.Label(
            self.tabNarrative,
            text="Narrative — complete after Saisie D'affaire is processed.",
            font=("TkDefaultFont", 9),
        ).grid(row=0, column=0, sticky="w", padx=8, pady=(10, 4))

        narrative_opts = ttk.Frame(self.tabNarrative)
        narrative_opts.grid(row=1, column=0, sticky="w", padx=8, pady=(0, 6))
        ttk.Label(narrative_opts, text="Language:").grid(row=0, column=0, sticky="w", padx=(0, 6))
        self.varNarrativeLang = tk.StringVar(value="EN")
        ttk.Radiobutton(narrative_opts, text="English", variable=self.varNarrativeLang, value="EN").grid(row=0, column=1, padx=(0, 10))
        ttk.Radiobutton(narrative_opts, text="Français", variable=self.varNarrativeLang, value="FR").grid(row=0, column=2, padx=(0, 20))

        self.txtNarrative = tk.Text(self.tabNarrative, wrap="word", font=("Courier New", 10), height=12)
        self.txtNarrative.grid(row=2, column=0, sticky="nsew", padx=8, pady=(0, 4))
        narrative_scroll = ttk.Scrollbar(self.tabNarrative, orient="vertical", command=self.txtNarrative.yview)
        narrative_scroll.grid(row=2, column=1, sticky="ns")
        self.txtNarrative.configure(yscrollcommand=narrative_scroll.set)

        narrative_actions = ttk.Frame(self.tabNarrative)
        narrative_actions.grid(row=3, column=0, sticky="w", padx=8, pady=(0, 8))
        ttk.Button(
            narrative_actions, text="Generate",
            command=self.on_generate_narrative, style="Primary.TButton",
        ).grid(row=0, column=0, padx=(0, 6))
        ttk.Button(
            narrative_actions, text="Copy to Clipboard",
            command=self._on_narrative_copy, style="Secondary.TButton",
        ).grid(row=0, column=1, padx=(0, 6))
        self.varCompleteNarrative = tk.BooleanVar(value=False)
        ttk.Checkbutton(
            narrative_actions, text="Complete ✓",
            variable=self.varCompleteNarrative,
            command=lambda: self._on_tab_complete_changed("narrative"),
        ).grid(row=0, column=2, padx=(0, 6))
        ttk.Button(
            narrative_actions, text="✕",
            command=lambda: self._on_tab_close("narrative"),
            style="Secondary.TButton", width=3,
        ).grid(row=0, column=3, padx=(0, 8))
        ttk.Label(narrative_actions, textvariable=self.varNarrativeStatus).grid(row=0, column=4, padx=(6, 0))

        # Saisie d'interet panel
        self.tabSaisieInteret.columnconfigure(1, weight=1)
        self.varInteretWorkingDir = tk.StringVar(value="-")
        self.varInteretTemplate = tk.StringVar(value="-")
        self.varInteretOutput = tk.StringVar(value="-")
        self.varInteretStatus = tk.StringVar(value="Select SAISIE file and run extraction first.")

        ttk.Label(self.tabSaisieInteret, text="Active case folder:").grid(row=0, column=0, sticky="w", padx=(8, 8), pady=(10, 2))
        ttk.Label(self.tabSaisieInteret, textvariable=self.varInteretWorkingDir).grid(row=0, column=1, sticky="w", pady=(10, 2))

        ttk.Label(self.tabSaisieInteret, text="Template found:").grid(row=1, column=0, sticky="w", padx=(8, 8), pady=2)
        ttk.Label(self.tabSaisieInteret, textvariable=self.varInteretTemplate).grid(row=1, column=1, sticky="w", pady=2)

        ttk.Label(self.tabSaisieInteret, text="Latest output:").grid(row=2, column=0, sticky="w", padx=(8, 8), pady=2)
        ttk.Label(self.tabSaisieInteret, textvariable=self.varInteretOutput).grid(row=2, column=1, sticky="w", pady=2)

        ttk.Label(self.tabSaisieInteret, textvariable=self.varInteretStatus).grid(row=3, column=0, columnspan=2, sticky="w", padx=8, pady=(2, 6))

        interetRow = ttk.Frame(self.tabSaisieInteret)
        interetRow.grid(row=4, column=0, columnspan=2, sticky="w", padx=8, pady=(8, 8))
        self.btnGenerateSaisieInteret = ttk.Button(
            interetRow,
            text="Generate Saisie d'interet",
            command=self.on_generate_saisie_interet,
            state="disabled",
            style="Primary.TButton",
        )
        self.btnGenerateSaisieInteret.grid(row=0, column=0)

        # Bottom middle action (placeholder for upcoming workflow).
        bottom_actions = ttk.Frame(frm)
        bottom_actions.grid(row=8, column=0, sticky="ew", pady=(10, 0))
        bottom_actions.columnconfigure(0, weight=1)
        bottom_actions.columnconfigure(1, weight=0)
        bottom_actions.columnconfigure(2, weight=1)
        self.btnInOutInventory = ttk.Button(
            bottom_actions,
            text="In/Out Inventory",
            command=self.on_open_inout_inventory,
            style="Secondary.TButton",
        )
        self.btnInOutInventory.grid(row=0, column=1)

        # Load saved paths from config
        templates_path = get_config_path("paths", "templates_folder")
        if templates_path:
            self.state.templates_folder = templates_path
        badge_cfg = clean_value(get_config_text("user", "badge_number", ""))
        if badge_cfg:
            self.state.badge_number = badge_cfg
            self.varBadgeNumber.set(badge_cfg)
            if hasattr(self, "varSaisieAffaireFields") and ("agent_badge" in self.varSaisieAffaireFields):
                self.varSaisieAffaireFields["agent_badge"].set(badge_cfg)
        
        saisie_folder = get_config_path("paths", "saisie_folder")
        if saisie_folder:
            # Set initial directory for file dialog
            self.last_saisie_folder = saisie_folder
        else:
            self.last_saisie_folder = None

        # Restore last SAISIE file from previous session (if it still exists)
        last_saisie_file = get_config_path("paths", "last_saisie_file")
        if last_saisie_file and last_saisie_file.exists() and last_saisie_file.is_file():
            self.state.saisie_pdf_file = last_saisie_file
            self._set_entry(self.entSaisieFile, last_saisie_file)
            if not self.state.working_dir:
                wd, cn = detect_working_directory(last_saisie_file)
                self.state.working_dir = wd
                self.state.case_folder_name = cn

        self._prefill_saisie_affaire_defaults()

        self._refresh_config_summary()
        self._init_dnd()
        self._apply_role_tab_visibility()   # hide tabs the current role cannot access
        self._refresh_case_folder_banner()
        self._refresh_agenda_status()
        self._apply_step_visibility()
        self._log_missing_helper_modules_once()

    def _init_styles(self):
        style = ttk.Style(self.root)
        try:
            style.theme_use("clam")
        except Exception:
            pass

        bg = "#F5F7FA"
        card = "#FFFFFF"
        ink = "#1F2A37"
        accent = "#1F4E79"
        accent_hover = "#163A5C"
        secondary = "#E8EEF5"
        secondary_hover = "#D9E3EF"
        line = "#C7D3E0"
        self._accent_color = accent

        style.configure(".", background=bg, foreground=ink, font=("Segoe UI", 10))
        style.configure("TFrame", background=bg)
        style.configure("Card.TFrame", background=card)
        style.configure("TLabelframe", background=bg, bordercolor=line, relief="solid")
        style.configure("TLabelframe.Label", background=bg, foreground=ink, font=("Segoe UI Semibold", 10))
        style.configure("TLabel", background=bg, foreground=ink)
        style.configure("TEntry", fieldbackground="#FFFFFF", bordercolor=line, relief="solid")
        style.configure("DragOver.TEntry", fieldbackground="#E3F0FB", bordercolor="#1F4E79", relief="solid")

        style.configure(
            "Primary.TButton",
            background=accent,
            foreground="#FFFFFF",
            bordercolor=accent,
            focusthickness=1,
            focuscolor=accent,
            padding=(10, 6),
            font=("Segoe UI Semibold", 10),
        )
        style.map(
            "Primary.TButton",
            background=[("active", accent_hover), ("disabled", "#9DB4C9")],
            foreground=[("disabled", "#F4F7FA")],
            bordercolor=[("active", accent_hover), ("disabled", "#9DB4C9")],
        )

        style.configure(
            "Secondary.TButton",
            background=secondary,
            foreground=ink,
            bordercolor=line,
            focusthickness=1,
            focuscolor=line,
            padding=(10, 6),
            font=("Segoe UI", 10),
        )
        style.map(
            "Secondary.TButton",
            background=[("active", secondary_hover), ("disabled", "#EDF2F7")],
            foreground=[("disabled", "#8A97A6")],
            bordercolor=[("active", line), ("disabled", "#D5DEE8")],
        )

        # Compact button style for small header actions (e.g., Help).
        style.configure(
            "Help.TButton",
            background=secondary,
            foreground=ink,
            bordercolor=line,
            focusthickness=1,
            focuscolor=line,
            padding=(6, 3),
            font=("Segoe UI", 9),
        )
        style.map(
            "Help.TButton",
            background=[("active", secondary_hover), ("disabled", "#EDF2F7")],
            foreground=[("disabled", "#8A97A6")],
            bordercolor=[("active", line), ("disabled", "#D5DEE8")],
        )

        # Compact primary button for header actions (e.g., Configurations).
        style.configure(
            "HeaderPrimary.TButton",
            background=accent,
            foreground="#FFFFFF",
            bordercolor=accent,
            focusthickness=1,
            focuscolor=accent,
            padding=(8, 3),
            font=("Segoe UI Semibold", 9),
        )
        style.map(
            "HeaderPrimary.TButton",
            background=[("active", accent_hover), ("disabled", "#9DB4C9")],
            foreground=[("disabled", "#F4F7FA")],
            bordercolor=[("active", accent_hover), ("disabled", "#9DB4C9")],
        )

        style.configure("TButton", padding=(10, 6))

        style.configure(
            "TNotebook",
            background=bg,
            borderwidth=0,
            tabmargins=(2, 4, 2, 0),
        )
        style.configure(
            "TNotebook.Tab",
            background="#E6EDF5",
            foreground=ink,
            padding=(14, 8),
            font=("Segoe UI Semibold", 10),
        )
        style.map(
            "TNotebook.Tab",
            background=[("selected", "#FFFFFF"), ("active", "#DFE9F3"), ("disabled", "#EEF3F8")],
            foreground=[("selected", accent), ("disabled", "#8B9AA9")],
        )

    def _init_dnd(self):
        """Base class - drag-and-drop not available."""
        pass

    def _on_app_close(self):
        """Graceful shutdown to avoid lingering grab/bind artifacts between launches."""
        try:
            existing = getattr(self, "_help_window", None)
            if existing is not None and existing.winfo_exists():
                existing.destroy()
        except Exception:
            pass
        try:
            self.root.unbind_all("<MouseWheel>")
        except Exception:
            pass
        try:
            self.root.quit()
        except Exception:
            pass
        try:
            self.root.destroy()
        except Exception:
            pass

    def _set_busy(self, busy: bool, message: str = ""):
        """Show/hide lightweight progress indicator for long-running actions."""
        if busy:
            self.varProgressStatus.set(f"Status: {message or 'Working...'}")
            try:
                self.progress.start(12)
            except Exception:
                pass
            try:
                self.root.config(cursor="watch")
            except Exception:
                pass
        else:
            try:
                self.progress.stop()
            except Exception:
                pass
            try:
                self.root.config(cursor="")
            except Exception:
                pass
            if message:
                self.varProgressStatus.set(message)
            else:
                self._refresh_instruction_feedback()
        try:
            self.root.update_idletasks()
        except Exception:
            pass

    def _prefill_saisie_affaire_defaults(self):
        fields = getattr(self, "varSaisieAffaireFields", {})
        if not fields:
            return
        badge = re.sub(r"\D", "", self.state.badge_number or self.varBadgeNumber.get() or "")
        if badge and ("agent_badge" in fields) and (not clean_value(fields["agent_badge"].get())):
            fields["agent_badge"].set(badge)
        if "interception_datetime" in fields and (not clean_value(fields["interception_datetime"].get())):
            fields["interception_datetime"].set(datetime.now().strftime("%Y-%m-%d %H:%M"))
        if "bond_room_ledger" in fields and (not clean_value(fields["bond_room_ledger"].get())):
            brl = self._resolve_bond_room_ledger()
            if brl:
                fields["bond_room_ledger"].set(brl)

    def _resolve_bond_room_ledger(self) -> str:
        # 1) Current in-memory extraction
        if isinstance(self.state.last_top, dict):
            raw = clean_value(str(self.state.last_top.get("BOND ROOM LEDGER #", "") or ""))
            if raw:
                return raw
        # 2) Cached case values
        try:
            wd = self._resolve_working_dir()
            if wd and self.state.saisie_pdf_file:
                case_paths = ensure_case_structure(wd, self.state.saisie_pdf_file)
                cached = read_values_latest_json(case_paths["values_latest_json"])
                raw = clean_value(str(cached.get("bond_room_ledger", "") or ""))
                if raw:
                    return raw
        except Exception:
            pass
        return ""

    def _extract_inventory_from_case_folder_name(self, folder_name: str) -> str:
        raw = clean_value(folder_name)
        if not raw:
            return ""

        parts = [p for p in re.split(r"\s+", raw) if p]
        if len(parts) >= 3:
            third = _normalize_inventory_number(parts[2])
            if third:
                return third

        for token in parts:
            t = clean_value(token)
            if not t:
                continue
            if re.fullmatch(r"\d{4}-\d{2}-\d{2}", t):
                continue
            if re.fullmatch(r"\d{4,6}", t):
                continue
            if "AEADS" in t.upper():
                continue
            candidate = _normalize_inventory_number(t)
            if candidate:
                return candidate

        return _normalize_inventory_number(raw)

    def _prefill_inventory_from_case_folder(self, folder_name: str, force: bool = True):
        fields = getattr(self, "varSaisieAffaireFields", {})
        inv_var = fields.get("inventory_number")
        if inv_var is None:
            return
        detected_inv = self._extract_inventory_from_case_folder_name(folder_name)
        if not detected_inv:
            return
        current_inv = _normalize_inventory_number(inv_var.get() or "")
        if force or (not current_inv):
            inv_var.set(detected_inv)
        self.state.last_inventory_number = detected_inv

    def _saisie_affaire_is_complete(self, case_paths: Dict[str, Path]) -> bool:
        required_keys = (
            "agent_badge",
            "inventory_number",
            "country",
            "interception_location",
            "interception_datetime",
            "declaration",
            "weight_qty",
            "item_description",
            "sender",
            "recipient",
            "indices",
        )
        cached = read_values_latest_json(case_paths["values_latest_json"])
        manual = cached.get("saisie_affaire_manual")
        if not isinstance(manual, dict):
            return False
        fields = manual.get("fields")
        if not isinstance(fields, dict):
            return False
        return all(clean_value(str(fields.get(k, "") or "")) for k in required_keys)

    def _instruction_feedback_text(self) -> str:
        has_config = bool(self.state.templates_folder) and bool(
            self.state.badge_number or re.sub(r"\D", "", self.varBadgeNumber.get() or "")
        )
        if not has_config:
            return "Configurations: Click the Configurations button (top-right) and set templates folder + badge number."

        has_selected_case_folder = bool(self.state.case_folder_locked and self.state.working_dir)
        wd = self._resolve_working_dir()
        if not has_selected_case_folder:
            return "Select Folder: Open 'Select Folder' tab, then select or create the case folder."

        agenda_exists = False
        has_k138_base = False
        k138_exists = False
        saisie_affaire_done = bool(self.state.saisie_affaire_generated)
        try:
            case_paths = ensure_case_structure(wd, self.state.saisie_pdf_file)
            saisie_affaire_done = self._saisie_affaire_is_complete(case_paths)
            has_k138_base = bool(self._cached_k138_values(case_paths))
            agenda_path = self._agenda_existing_path(case_paths)
            agenda_exists = bool(agenda_path and agenda_path.exists())
            k138_exists = bool(
                case_paths["k138_output_pdf"].exists()
                or case_paths["k138_latest_pdf"].exists()
                or (
                    self.state.last_k138_output
                    and Path(self.state.last_k138_output).exists()
                )
            )
        except Exception:
            pass

        selected = ""
        try:
            selected = self.tabs.select()
        except Exception:
            selected = ""

        # Tab-aware guidance: message is relative to current tab.
        if selected == str(self.tabSelectFolder):
            if has_selected_case_folder:
                return "Select Folder: Case folder is selected. Next, open Saisie D'affaire tab."
            return "Select Folder: Choose an existing case folder or create a new one."

        if selected == str(self.tabSaisieAffaire):
            if not saisie_affaire_done:
                return "Saisie D'affaire: Fill required fields (badge/date auto-filled), set checkboxes, then click Generate Saisie D'affaire."
            return "Saisie D'affaire: Done. Next, open Agenda tab and run Extract Values from Saise A faire."

        if selected == str(self.tabAgenda):
            if not self.state.saisie_pdf_file:
                return "Agenda: Select SAISIE input file, then click Extract Values from Saise A faire."
            if not self.state.extraction_ran:
                return "Agenda: Click Extract Values from Saise A faire, then click Generate Agenda."
            if not agenda_exists:
                return "Agenda: Click Generate Agenda. If Agenda changes, click Update Values from Agenda."
            if not has_k138_base:
                return "Agenda: Re-run Extract Values from Saise A faire to prepare K138 base values."
            return "Agenda: Done. Next, open K138 tab."

        if selected == str(self.tabK138):
            if not self.state.saisie_pdf_file or not self.state.extraction_ran:
                return "K138: In Agenda tab, select SAISIE file and run Extract Values first."
            if not agenda_exists:
                return "K138: Generate Agenda first."
            if not has_k138_base:
                return "K138: Re-run Extract Values from Saise A faire first."
            if not k138_exists:
                return "K138: Confirm form type, then click Generate K138."
            return "K138: Generated. If Agenda values change, regenerate K138."

        if selected == str(self.tabSaisieInteret):
            if not self.state.saisie_pdf_file or not self.state.extraction_ran:
                return "Saisie d'interet: Run Extract Values from Saise A faire first."
            return "Saisie d'interet: Click Generate Saisie d'interet when needed."

        # Fallback: global next step.
        if not self.state.saisie_pdf_file:
            return "Saisie D'affaire: Fill fields, then click Generate Saisie D'affaire."
        if not saisie_affaire_done:
            return "Saisie D'affaire: Complete fields, then click Generate Saisie D'affaire."
        if not self.state.extraction_ran:
            return "Agenda: Click Extract Values from Saise A faire."
        if not agenda_exists:
            return "Agenda: Click Generate Agenda."
        if not has_k138_base:
            return "Agenda: Re-run Extract Values from Saise A faire."
        if not k138_exists:
            return "K138: Click Generate K138."
        return "Complete."

    def _refresh_instruction_feedback(self):
        self.varProgressStatus.set(self._instruction_feedback_text())

    def log(self, msg: str):
        line = normalize_output_text(str(msg))
        if HIDE_CSV_IN_FEEDBACK and ("csv" in line.lower()):
            return
        if CONCISE_PROGRESS_LOGS:
            ll = line.lower()
            keep_force = (
                ("error" in ll)
                or ("failed" in ll)
                or ("validation" in ll)
                or ("missing critical field" in ll)
            )
            if (not keep_force) and any(p in ll for p in _NOISY_PROGRESS_SUBSTRINGS):
                return
        self.txt.insert("end", line + "\n")
        self.txt.see("end")

    def _active_case_folder_path(self) -> Optional[Path]:
        wd = self._resolve_working_dir()
        return wd.resolve() if wd else None

    def _active_case_folder_name(self) -> str:
        p = self._active_case_folder_path()
        return p.name if p else "-"

    def _role_allows_tab(self, tab_name: str) -> bool:
        """Return True if the current profile role has access to the given tab."""
        role = getattr(self, "profile_role", "BSO") or "BSO"
        if role == "Supervisor":
            return True
        _access = {
            # BSO: no K138, no Saisie d'interet
            "BSO":   {"Select Folder", "Saisie D'affaire", "Agenda", "Narrative"},
            # Clerk: can see and update Saisie D'affaire and Select Folder too
            "Clerk": {"Select Folder", "Saisie D'affaire", "Agenda", "K138", "Saisie d'interet"},
        }
        return tab_name in _access.get(role, set())

    def _apply_role_tab_visibility(self):
        """
        Enforce role-based tab restrictions.
        Tabs the role cannot access are fully hidden from the tab bar (state='hidden').
        Called at startup and re-called after every status refresh so nothing can
        accidentally re-show a restricted tab.
        """
        role = getattr(self, "profile_role", "BSO") or "BSO"
        _all_tabs = {
            "Select Folder":    self.tabSelectFolder,
            "Saisie D'affaire": self.tabSaisieAffaire,
            "Agenda":           self.tabAgenda,
            "K138":             self.tabK138,
            "Narrative":        self.tabNarrative,
            "Saisie d'interet": self.tabSaisieInteret,
        }
        try:
            for tab_name, tab_widget in _all_tabs.items():
                if not self._role_allows_tab(tab_name):
                    self.tabs.tab(tab_widget, state="hidden")
            # Clerk: Select Folder and Agenda are always visible; land on Select Folder at startup.
            if role == "Clerk":
                self.tabs.tab(self.tabSelectFolder, state="normal")
                self.tabs.tab(self.tabAgenda, state="normal")
                if not getattr(self, "_clerk_initial_tab_set", False):
                    self.tabs.select(self.tabSelectFolder)
                    self._clerk_initial_tab_set = True
        except Exception:
            pass

    def _refresh_folder_dependent_tabs(self):
        has_folder = bool(self.state.case_folder_locked and self.state.working_dir)
        role = getattr(self, "profile_role", "BSO") or "BSO"
        # Saisie D'affaire found status
        if hasattr(self, "varSaisieAffaireFound"):
            active = self._active_case_folder_path()
            if active and active.exists():
                pdfs = list(active.glob("*Saisie_D_affaire.pdf"))
                if pdfs:
                    self.varSaisieAffaireFound.set(f"found: yes  ({pdfs[0].name})")
                else:
                    self.varSaisieAffaireFound.set("found: no")
            elif active:
                self.varSaisieAffaireFound.set("found: no")
            else:
                self.varSaisieAffaireFound.set("-")
        try:
            # Saisie D'affaire — only enable if role allows AND folder is selected
            if self._role_allows_tab("Saisie D'affaire"):
                self.tabs.tab(self.tabSaisieAffaire,
                              state="normal" if has_folder else "disabled")

            # Agenda — Clerk always sees it; others need a folder
            if self._role_allows_tab("Agenda"):
                if role == "Clerk":
                    self.tabs.tab(self.tabAgenda, state="normal")
                else:
                    self.tabs.tab(self.tabAgenda,
                                  state="normal" if has_folder else "disabled")

            # Select Folder — always enabled when role allows
            if self._role_allows_tab("Select Folder"):
                self.tabs.tab(self.tabSelectFolder, state="normal")

            # Narrative — BSO/Supervisor, same availability as Agenda
            if self._role_allows_tab("Narrative"):
                self.tabs.tab(self.tabNarrative,
                              state="normal" if has_folder else "disabled")

            # When no folder is selected, land on the correct default tab —
            # but only if the user isn't already on an allowed tab for their role.
            if not has_folder:
                if role == "Clerk":
                    try:
                        current = self.tabs.select()
                        clerk_tabs = {
                            str(self.tabSelectFolder), str(self.tabSaisieAffaire),
                            str(self.tabAgenda), str(self.tabK138), str(self.tabSaisieInteret),
                        }
                        if current not in clerk_tabs:
                            self.tabs.select(self.tabSelectFolder)
                    except Exception:
                        self.tabs.select(self.tabSelectFolder)
                elif self._role_allows_tab("Select Folder"):
                    self.tabs.select(self.tabSelectFolder)
        except Exception:
            pass

    def _refresh_case_folder_banner(self):
        active_path = self._active_case_folder_path()
        name = active_path.name if active_path else "-"
        self.varCaseFolderBanner.set(name)
        if hasattr(self, "varSelectedCaseFolder"):
            self.varSelectedCaseFolder.set(name)
        try:
            self.btnOpenCaseFolder.configure(state="normal" if active_path else "disabled")
        except Exception:
            pass
        self._refresh_folder_dependent_tabs()

    def _refresh_config_summary(self):
        folder_txt = self.state.templates_folder.name if self.state.templates_folder else "-"
        badge_txt = self.state.badge_number or re.sub(r"\D", "", self.varBadgeNumber.get() or "") or "-"
        files_txt = ""
        if self.state.templates_folder and self.state.templates_folder.exists():
            try:
                found = [
                    f.name for f in self.state.templates_folder.iterdir()
                    if f.is_file() and f.suffix.lower() in (".pdf", ".docx", ".xlsx", ".txt")
                ]
                if found:
                    files_txt = "  |  Files: " + ", ".join(sorted(found)[:8])
                    if len(found) > 8:
                        files_txt += f" (+{len(found) - 8} more)"
            except Exception:
                pass
        self.varConfigSummary.set(f"Folder: {folder_txt} | Badge: {badge_txt}{files_txt}")
        self._apply_step_visibility()
        self._refresh_instruction_feedback()

    def _apply_step_visibility(self):
        """
        Show only controls relevant to the current workflow step:
        - SAISIE select + Extract: visible only in Agenda tab (and only after configurations are set).
        - K138 form type: visible only when user is in K138 tab.
        """
        has_config = bool(self.state.templates_folder) and bool(self.state.badge_number)
        selected = ""
        try:
            selected = self.tabs.select()
        except Exception:
            selected = ""

        show_saisie_step = has_config and (selected == str(self.tabAgenda))
        show_k138_form_type = has_config and (selected == str(self.tabK138))

        try:
            if show_saisie_step:
                self.boxSaisieFile.grid()
            else:
                self.boxSaisieFile.grid_remove()
        except Exception:
            pass

        try:
            if show_k138_form_type:
                self.boxFormType.grid()
            else:
                self.boxFormType.grid_remove()
        except Exception:
            pass

    def _open_directory(self, path: Optional[Path]):
        if not path:
            return
        p = Path(path)
        if not p.exists():
            messagebox.showwarning("Folder Not Found", f"Folder not found:\n{p}")
            return
        try:
            if os.name == "nt":
                os.startfile(str(p))  # type: ignore[attr-defined]
            elif sys.platform == "darwin":
                subprocess.Popen(["open", str(p)])
            else:
                subprocess.Popen(["xdg-open", str(p)])
        except Exception as e:
            messagebox.showerror("Open Folder", f"Could not open folder:\n{e}")

    def on_open_active_case_folder(self):
        self._open_directory(self._active_case_folder_path())

    def on_open_templates_folder(self):
        if not self.state.templates_folder:
            messagebox.showwarning("Missing Folder", "Select Configurations folder first.")
            return
        self._open_directory(self.state.templates_folder)

    def on_save_badge_number(self):
        badge = re.sub(r"\D", "", clean_value(self.varBadgeNumber.get()))
        self.state.badge_number = badge
        self.varBadgeNumber.set(badge)
        if hasattr(self, "varSaisieAffaireFields") and ("agent_badge" in self.varSaisieAffaireFields):
            self.varSaisieAffaireFields["agent_badge"].set(badge)
        set_config_text("user", "badge_number", badge)
        self._prefill_saisie_affaire_defaults()
        self._refresh_config_summary()
        self.log(f"Badge number saved: {badge or '(blank)'}")

    def _set_active_case_folder(self, folder: Path, write_log: bool = True):
        self.state.working_dir = folder.resolve()
        self.state.case_folder_name = folder.name
        self.state.case_folder_locked = True
        self.last_saisie_folder = folder.resolve()
        set_config_path("paths", "saisie_folder", folder.resolve())
        self._prefill_inventory_from_case_folder(folder.name, force=True)
        self._refresh_case_folder_banner()
        if write_log:
            self.log(f"Changed to: {folder.resolve()}")
        self._refresh_agenda_status()

    def on_select_case_folder(self):
        initial_dir = str(self._active_case_folder_path() or self.last_saisie_folder or Path.cwd())
        selected = filedialog.askdirectory(title="Select Case Folder", initialdir=initial_dir)
        if not selected:
            return
        folder = Path(selected)
        self._set_active_case_folder(folder)

    def on_create_case_folder(self):
        active = self._active_case_folder_path()
        initial_parent = str(
            (active.parent if active else None)
            or self.last_saisie_folder
            or Path.cwd()
        )
        parent_raw = filedialog.askdirectory(
            title="Select Parent Folder for New Case Folder",
            initialdir=initial_parent,
        )
        if not parent_raw:
            return
        parent = Path(parent_raw)
        badge_clean = re.sub(r"\D", "", self.state.badge_number or self.varBadgeNumber.get() or "")
        if not badge_clean:
            messagebox.showwarning(
                "Missing Badge",
                "Configure badge number first in Configurations (top-right).",
            )
            return
        default_date = datetime.now().strftime("%Y-%m-%d")
        case_date = simpledialog.askstring(
            "New Case Folder",
            "Case date (YYYY-MM-DD):",
            parent=self.root,
            initialvalue=default_date,
        )
        if case_date is None:
            return
        inventory = simpledialog.askstring(
            "New Case Folder",
            "Inventory number:",
            parent=self.root,
            initialvalue="",
        )
        if inventory is None:
            return
        date_clean = clean_value(case_date)
        inventory_clean = clean_value(inventory).upper()
        if (not badge_clean) or (not date_clean) or (not inventory_clean):
            messagebox.showerror("Invalid Case Folder", "Badge, date, and inventory are required.")
            return
        folder_name = f"{badge_clean} {date_clean} {inventory_clean}"
        case_folder = parent / folder_name
        try:
            case_folder.mkdir(parents=True, exist_ok=True)
        except Exception as e:
            messagebox.showerror("Create Case Folder", f"Could not create case folder:\n{e}")
            return
        self._set_active_case_folder(case_folder)
        self.log(f"Created case folder: {case_folder.name}")

    def _collect_saisie_affaire_form_values(self) -> Tuple[Dict[str, str], Dict[str, bool]]:
        self._prefill_saisie_affaire_defaults()
        active_case = self._active_case_folder_path()
        if active_case:
            self._prefill_inventory_from_case_folder(active_case.name, force=False)

        field_values = {
            key: clean_value(var.get())
            for key, var in getattr(self, "varSaisieAffaireFields", {}).items()
        }
        if not clean_value(field_values.get("bond_room_ledger", "")):
            field_values["bond_room_ledger"] = self._resolve_bond_room_ledger()
        if field_values.get("agent_badge", ""):
            field_values["agent_badge"] = self._normalize_agent_id(field_values.get("agent_badge", ""))
        if field_values.get("inventory_number", ""):
            field_values["inventory_number"] = _normalize_inventory_number(field_values.get("inventory_number", ""))

        check_values = {
            "saisie": bool(self.varCheckSaisie.get()),
            "confiscation": bool(self.varCheckConfiscation.get()),
            "k9": bool(self.varCheckK9.get()),
            "saisie_denvergure": bool(self.varCheckSaisieEnvergure.get()),
        }
        return field_values, check_values

    def _reset_saisie_affaire_form(self):
        for _key, var in getattr(self, "varSaisieAffaireFields", {}).items():
            var.set("")
        self.varCheckSaisie.set(False)
        self.varCheckConfiscation.set(False)
        self.varCheckK9.set(False)
        self.varCheckSaisieEnvergure.set(False)

        # Keep operational defaults for next case/form.
        self._prefill_saisie_affaire_defaults()
        active_case = self._active_case_folder_path()
        if active_case:
            self._prefill_inventory_from_case_folder(active_case.name, force=True)

    def on_generate_saisie_affaire(self):
        if not self.state.templates_folder:
            messagebox.showwarning("Missing Folder", "Select Configurations folder first.")
            return

        active_case = self._active_case_folder_path()
        if not (self.state.case_folder_locked and active_case):
            messagebox.showwarning("Missing Case Folder", "Select or create case folder first.")
            return

        template_pdf = find_saisie_template(self.state.templates_folder)
        if not template_pdf:
            messagebox.showerror(
                "Missing SAISIE Template",
                f"Could not find SAISIE template PDF in:\n{self.state.templates_folder}\n\n"
                f"Make sure your Configurations folder contains the blank SAISIE PDF template.\n\n"
                f"Go to Configurations (top-right) and re-select the correct templates folder.",
            )
            return

        field_values, check_values = self._collect_saisie_affaire_form_values()
        _inv = _normalize_inventory_number(field_values.get("inventory_number", ""))
        _inv_prefix = f"{_inv}_" if _inv else ""
        output_pdf = active_case / f"{_inv_prefix}Saisie_D_affaire.pdf"

        self._set_busy(True, "Generating Saisie D'affaire...")
        try:
            _updated_text, _created_text, _updated_checks = fill_saisie_affaire_pdf(
                template_path=template_pdf,
                output_path=output_pdf,
                field_values=field_values,
                check_values=check_values,
            )
            self.state.saisie_affaire_generated = True
            try:
                case_paths = ensure_case_structure(active_case, output_pdf)
                update_patch: Dict[str, object] = {
                    "updated_at": _timestamp_iso(),
                    "source_file": str(output_pdf),
                    "working_directory": str(active_case),
                    "case_folder_name": active_case.name,
                    "saisie_affaire_manual": {
                        "fields": field_values,
                        "checks": check_values,
                    },
                }
                inv = _normalize_inventory_number(field_values.get("inventory_number", ""))
                badge = self._normalize_agent_id(field_values.get("agent_badge", ""))
                if inv:
                    update_patch["inventory_number"] = inv
                if badge:
                    update_patch["agent_id"] = badge
                update_values_latest_json(case_paths["values_latest_json"], update_patch)
            except Exception:
                pass
            self.log("Saisie D'affaire created.")
            self.log(f"Output: {output_pdf.name}")
            self._refresh_instruction_feedback()
            if hasattr(self, "varSaisieAffaireFound"):
                self.varSaisieAffaireFound.set(f"found: yes  ({output_pdf.name})")
            messagebox.showinfo("Saisie D'affaire", f"Generated:\n{output_pdf}")
        except Exception as e:
            self.log(f"Saisie D'affaire generation failed: {e}")
            import traceback
            self.log(traceback.format_exc())
            messagebox.showerror("Saisie D'affaire", str(e))
        finally:
            self._set_busy(False)

    def on_update_from_saisie_affaire(self):
        wd = self._resolve_working_dir()
        if not wd or not self.state.saisie_pdf_file:
            messagebox.showwarning(
                "Missing Case",
                "Select SAISIE input file first to determine active case folder.",
            )
            return

        field_values, check_values = self._collect_saisie_affaire_form_values()

        case_paths = ensure_case_structure(wd, self.state.saisie_pdf_file)
        base_vals = self._cached_k138_values(case_paths)
        if not isinstance(base_vals, dict):
            base_vals = {}

        inv = _normalize_inventory_number(field_values.get("inventory_number", ""))
        badge = self._normalize_agent_id(field_values.get("agent_badge", ""))
        decl = field_values.get("declaration", "")
        item = field_values.get("item_description", "")

        if inv:
            base_vals["description_inventory"] = inv
            self.state.last_inventory_number = inv
        if badge:
            base_vals["seizing_officer"] = badge
            self.state.last_agent_id = badge
        if decl:
            base_vals["description_declared"] = decl
        if item:
            base_vals["description_item"] = item

        update_patch: Dict[str, object] = {
            "updated_at": _timestamp_iso(),
            "source_file": str(self.state.saisie_pdf_file),
            "working_directory": str(wd),
            "case_folder_name": self.state.case_folder_name or wd.name,
            "k138_values_base": base_vals,
            "saisie_affaire_manual": {
                "fields": field_values,
                "checks": check_values,
            },
        }
        if inv:
            update_patch["inventory_number"] = inv
        if badge:
            update_patch["agent_id"] = badge

        update_values_latest_json(case_paths["values_latest_json"], update_patch)
        self._refresh_agenda_status()
        self._refresh_instruction_feedback()
        self.log("Saisie D'affaire values updated.")
        messagebox.showinfo("Saisie D'affaire", "Values updated from Saisie D'affaire form.")

    def _helper_module_missing(self, feature_name: str) -> bool:
        module_name = REQUIRED_HELPER_MODULES.get(feature_name, "")
        if not module_name:
            return False
        return not is_helper_module_available(module_name)

    def _log_missing_helper_modules_once(self):
        for feature_name, module_name in REQUIRED_HELPER_MODULES.items():
            if is_helper_module_available(module_name):
                continue
            if module_name in self._missing_helper_warnings_shown:
                continue
            self._missing_helper_warnings_shown.add(module_name)
            self.log(
                f"  ! Missing helper module for {feature_name}: {module_name}. "
                f"Expected: {helper_module_hint(module_name)}"
            )

    def on_open_inout_inventory(self):
        """Placeholder for future in/out inventory module."""
        messagebox.showinfo("In/Out Inventory", "Work in progress")

    def _help_pages(self) -> List[Tuple[str, str]]:
        """Guided help pages shown in the Help dialog."""
        return [
            (
                "Welcome",
                "This app processes one SAISIE case at a time.\n\n"
                "Use the tabs in this order:\n"
                "- Select Folder\n"
                "- Saisie D'affaire\n"
                "- Agenda\n"
                "- K138\n\n"
                "Instruction & Feedback updates based on the active tab."
            ),
            (
                "Configurations",
                "Use the Configurations button at the top-right (next to Help).\n\n"
                "Select the folder that contains templates and notes:\n"
                "- K138 template PDF(s)\n"
                "- Agenda template PDF/DOCX\n"
                "- SAISIE template PDF\n"
                "- k138_note_*.txt notice files\n\n"
                "Set your badge number here as well."
            ),
            (
                "Select Folder Tab",
                "Choose an existing case folder or create a new one.\n\n"
                "Folder name format:\n"
                "BADGE YYYY-MM-DD INVENTORY\n\n"
                "Example:\n"
                "12345 2026-03-14 W30824792"
            ),
            (
                "Saisie D'affaire Tab",
                "Fill the fields and check required boxes.\n"
                "Then click Generate Saisie D'affaire.\n\n"
                "Important:\n"
                "- DÉCLARATION and NOTES are separate fields\n"
                "- Date and badge are auto-filled when available"
            ),
            (
                "Agenda Tab",
                "Select a SAISIE input file:\n"
                "- PDF\n"
                "- DOCX\n"
                "- Image (png/jpg/tif/bmp)\n\n"
                "Then click:\n"
                "1) Extract Values from Saise A faire\n"
                "2) Generate Agenda\n\n"
                "Use Update Values from Agenda after manual Agenda edits."
            ),
            (
                "K138 Tab",
                "Generate K138 uses:\n"
                "- extracted SAISIE values as base\n"
                "- CE/CID (SIED) from Agenda when available\n"
                "- selected K138 template from Configurations folder\n\n"
                "Confirm form type, then click Generate K138."
            ),
            (
                "Drag and Drop",
                "Drag-and-drop is enabled for SAISIE input file\n"
                "(PDF or DOCX).\n\n"
                "Dropped files are kept in place.\n"
                "The app links them to the active case folder runtime."
            ),
            (
                "Saisie d'interet (Optional)",
                "Generate Saisie d'interet maps extracted values into the Excel template.\n\n"
                "Run this after extraction (and after Agenda/K138 if you want latest synced values)."
            ),
            (
                "Coming Soon",
                "Planned feature: In/Out Inventory with Case Status tracking.\n\n"
                "Goal:\n"
                "- Track inventory movement in and out of custody\n"
                "- Show a per-case status timeline (example: Extracted -> Agenda -> K138 -> Closed)\n"
                "- Improve visibility for where each case is in the workflow"
            ),
        ]

    def on_exit_session(self):
        """Save current session state and close the application."""
        try:
            if self.state.saisie_pdf_file and self.state.saisie_pdf_file.exists():
                set_config_path("paths", "last_saisie_file", self.state.saisie_pdf_file)
            if self.state.working_dir:
                set_config_path("paths", "saisie_folder", self.state.working_dir)
        except Exception:
            pass
        self.root.destroy()

    def on_change_profile(self):
        """Re-show the profile splash and restart with the new selection."""
        if messagebox.askyesno(
            "Change Profile",
            "Changing profile will restart the application.\n\nContinue?",
            parent=self.root,
        ):
            self.root._change_profile_requested = True
            self.root.destroy()

    def on_open_help(self):
        """Open a paged help dialog with Next/Back navigation."""
        existing = getattr(self, "_help_window", None)
        if existing is not None:
            try:
                if existing.winfo_exists():
                    existing.lift()
                    existing.focus_force()
                    return
            except Exception:
                pass

        pages = self._help_pages()
        if not pages:
            return

        win = tk.Toplevel(self.root)
        self._help_window = win
        win.title("Help - Radiance Copilot")
        win.transient(self.root)
        win.resizable(False, False)
        win.geometry("760x520")
        win.configure(bg="#F5F7FA")

        try:
            app_png_path = resolve_asset_path("photos/Radiance-copilot-icon.png", "Radiance-copilot-icon.png")
            if app_png_path:
                win.iconphoto(True, tk.PhotoImage(file=app_png_path))
        except Exception:
            pass

        idx = {"value": 0}
        title_var = tk.StringVar()
        page_var = tk.StringVar()
        body_var = tk.StringVar()

        root_frm = ttk.Frame(win, padding=14)
        root_frm.pack(fill="both", expand=True)

        ttk.Label(root_frm, textvariable=title_var, font=("Segoe UI Semibold", 13)).pack(anchor="w")
        ttk.Label(root_frm, textvariable=page_var, font=("Segoe UI", 9)).pack(anchor="w", pady=(2, 8))
        ttk.Separator(root_frm, orient="horizontal").pack(fill="x", pady=(0, 10))

        body_holder = tk.Frame(root_frm, bg="#FFFFFF", bd=1, relief="solid")
        body_holder.pack(fill="both", expand=True)
        body_lbl = tk.Label(
            body_holder,
            textvariable=body_var,
            justify="left",
            anchor="nw",
            bg="#FFFFFF",
            fg="#1F2A37",
            wraplength=700,
            padx=12,
            pady=12,
            font=("Segoe UI", 10),
        )
        body_lbl.pack(fill="both", expand=True)

        nav = ttk.Frame(root_frm)
        nav.pack(fill="x", pady=(10, 0))
        nav.columnconfigure(0, weight=1)
        nav.columnconfigure(1, weight=0)
        nav.columnconfigure(2, weight=0)
        nav.columnconfigure(3, weight=0)

        def _close_help():
            try:
                win.destroy()
            finally:
                self._help_window = None

        def _render():
            i = idx["value"]
            title, body = pages[i]
            title_var.set(title)
            body_var.set(body)
            page_var.set(f"Page {i + 1} / {len(pages)}")
            btn_back.configure(state="normal" if i > 0 else "disabled")
            btn_next.configure(text="Next" if i < len(pages) - 1 else "Done")

        def _prev(_event=None):
            if idx["value"] > 0:
                idx["value"] -= 1
                _render()

        def _next(_event=None):
            if idx["value"] < len(pages) - 1:
                idx["value"] += 1
                _render()
            else:
                _close_help()

        btn_back = ttk.Button(nav, text="Back", command=_prev, style="Secondary.TButton")
        btn_back.grid(row=0, column=1, padx=(0, 8))
        btn_next = ttk.Button(nav, text="Next", command=_next, style="Primary.TButton")
        btn_next.grid(row=0, column=2, padx=(0, 8))
        ttk.Button(nav, text="Close", command=_close_help, style="Secondary.TButton").grid(row=0, column=3)

        win.bind("<Left>", _prev)
        win.bind("<Right>", _next)
        win.bind("<Escape>", lambda _e: _close_help())
        win.protocol("WM_DELETE_WINDOW", _close_help)

        _render()
        win.grab_set()
        win.focus_force()

    def _reset_case_runtime_state(self, keep_case_folder: bool = True):
        """Reset transient per-case values/UI when a new SAISIE file is selected."""
        if not keep_case_folder:
            self.state.working_dir = None
            self.state.case_folder_name = ""
            self.state.case_folder_locked = False
        self.state.last_top = None
        self.state.last_bottom = None
        self.state.last_saisie_csv = None
        self.state.last_k138_values_csv = None
        self.state.last_k138_output = None
        self.state.last_inventory_number = ""
        self.state.last_agent_id = ""
        self.state.last_sied_number = ""
        self.state.last_k138_values = None
        self.state.extraction_ran = False  # Must re-extract before cached values are trusted
        self.state.saisie_affaire_generated = False

        # Reset visible status fields to baseline until extraction runs for the new case.
        self.varAgendaWorkingDir.set("-")
        self.varAgendaInventory.set("-")
        self.varAgendaAgent.set("-")
        self.varAgendaFound.set("no")
        self.varK138AgendaReady.set("no")
        self.varK138Sied.set("-")
        self.varK138Status.set("Process SAISIE and create Agenda first.")
        self.varInteretWorkingDir.set("-")
        self.varInteretTemplate.set("-")
        self.varInteretOutput.set("-")
        self.varInteretStatus.set("Select SAISIE file and run extraction first.")
        self.btnRefreshAgenda.configure(state="disabled")
        self.btnGenerateK138.configure(state="disabled")
        self.btnGenerateSaisieInteret.configure(state="disabled")
        # Never disable a tab the current role is allowed to see — disabling the
        # currently-selected tab causes the Notebook to jump to another tab.
        # The gate label/button state inside the tab communicates the unavailability.
        _k138_state = "normal" if self._role_allows_tab("K138") else "disabled"
        _si_state   = "normal" if self._role_allows_tab("Saisie d'interet") else "disabled"
        self.tabs.tab(self.tabK138, state=_k138_state)
        self.tabs.tab(self.tabSaisieInteret, state=_si_state)
        self.frmK138Actions.grid_remove()
        self.lblK138Gate.grid()
        if not keep_case_folder:
            self._refresh_case_folder_banner()

    def _set_working_directory(self, file_path: Path, refresh_status: bool = True):
        if self.state.case_folder_locked and self.state.working_dir:
            wd = self.state.working_dir.resolve()
            case_name = wd.name
        else:
            wd, case_name = detect_working_directory(file_path)
        self.state.working_dir = wd
        self.state.case_folder_name = case_name
        self.log(f"Changed to: {wd}")
        self._refresh_case_folder_banner()
        if refresh_status:
            self._refresh_agenda_status()
        else:
            self.varAgendaWorkingDir.set(wd.name)
            self.varInteretWorkingDir.set(wd.name)

    def _resolve_working_dir(self) -> Optional[Path]:
        if self.state.working_dir and (not _is_transient_upload_path(self.state.working_dir)):
            return self.state.working_dir
        if self.state.saisie_pdf_file:
            return _resolve_working_directory_for_source(self.state.saisie_pdf_file)
        return None

    def _agenda_existing_path(self, case_paths: Dict[str, Path]) -> Optional[Path]:
        # Read-only lookup: never create/copy files during status checks.
        candidates: List[Path] = []
        for key in ("agenda_output_pdf", "agenda_output_docx", "agenda_latest_pdf", "agenda_latest_docx"):
            p = case_paths.get(key)
            if p and p.exists():
                candidates.append(p)
        if candidates:
            try:
                return max(candidates, key=lambda p: p.stat().st_mtime)
            except Exception:
                return candidates[0]
        return None

    def _normalize_agent_id(self, raw: str) -> str:
        d = re.sub(r"\D", "", raw or "")
        if not d:
            return ""
        return d[-5:] if len(d) > 5 else d.zfill(5)

    def _as_bool(self, value: object) -> bool:
        if isinstance(value, bool):
            return value
        if isinstance(value, (int, float)):
            return value != 0
        txt = clean_value(str(value or "")).lower()
        return txt in {"1", "true", "yes", "y", "on"}

    def _cached_k138_values(self, case_paths: Dict[str, Path]) -> Dict[str, str]:
        cached = read_values_latest_json(case_paths["values_latest_json"])
        base = cached.get("k138_values_base", {})
        if not isinstance(base, dict):
            return {}
        out: Dict[str, str] = {}
        for k, v in base.items():
            out[str(k)] = "" if v is None else str(v)
        return out

    def _recover_agenda_core_values_from_source(self, source_path: Path) -> Dict[str, str]:
        """
        Best-effort recovery for inventory/agent directly from source SAISIE file.
        Used when cached extraction missed a required field for Agenda generation.
        """
        recovered = {"inventory_number": "", "agent_id": ""}
        if not source_path or not source_path.exists():
            return recovered

        suffix = (source_path.suffix or "").lower()
        text_candidates: List[str] = []

        try:
            template_pdf: Optional[Path] = None
            if suffix == ".pdf" and self.state.templates_folder:
                template_pdf = find_saisie_template(self.state.templates_folder)
            notice_text = (
                load_notice_text(self.state.templates_folder, self.state.form_type)
                if self.state.templates_folder
                else ""
            )
            top_ref, _bottom_ref, _mode_ref = extract_field_values(template_pdf, source_path)
            vals_ref = build_k138_values_from_saisie(top_ref, source_path, self.state.form_type, notice_text)
            recovered["inventory_number"] = _normalize_inventory_number(vals_ref.get("description_inventory", ""))
            recovered["agent_id"] = self._normalize_agent_id(vals_ref.get("seizing_officer", ""))
        except Exception as rec_err:
            self.log(f"  ! Recovery extraction warning: {rec_err}")

        try:
            if suffix == ".pdf":
                page_text = get_page_text(source_path)
                if page_text:
                    text_candidates.append(page_text)
                if HAVE_PIL and HAVE_TESSERACT:
                    ocr_text = get_ocr_page_text_pdf(source_path)
                    if ocr_text:
                        text_candidates.append(ocr_text)
            elif suffix == ".docx":
                docx_text = get_text_from_docx(source_path)
                if docx_text:
                    text_candidates.append(docx_text)
                ocr_docx_text = get_ocr_text_from_docx_images(source_path)
                if ocr_docx_text:
                    text_candidates.append(ocr_docx_text)
            elif suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}:
                img_text = get_ocr_text_from_image(source_path)
                if img_text:
                    text_candidates.append(img_text)
        except Exception as txt_err:
            self.log(f"  ! Recovery text warning: {txt_err}")

        for txt in text_candidates:
            if not recovered["inventory_number"]:
                inv = (
                    extract_inventory_by_label(txt)
                    or _extract_inventory_from_text_label_only(txt)
                    or _extract_inventory_from_text_loose(txt)
                )
                inv = _normalize_inventory_number(inv) if inv else ""
                if (not inv) and txt:
                    inv = _extract_inventory_from_text_label_only(txt)
                recovered["inventory_number"] = inv or recovered["inventory_number"]
            if not recovered["agent_id"]:
                recovered["agent_id"] = self._normalize_agent_id(extract_officer_by_label(txt))
            if recovered["inventory_number"] and recovered["agent_id"]:
                break

        return recovered

    def _refresh_saisie_interet_status(self):
        wd = self._resolve_working_dir()
        self.varInteretWorkingDir.set(wd.name if wd else "-")
        self.varInteretTemplate.set("-")
        self.varInteretOutput.set("-")
        self.varInteretStatus.set("Select SAISIE input file first.")
        self.btnGenerateSaisieInteret.configure(state="disabled")

        # Never disable the tab for roles that are allowed to see it —
        # disabling the current tab causes the notebook to jump away.
        # Hide it only for roles without access; otherwise keep it normal.
        if self._role_allows_tab("Saisie d'interet"):
            self.tabs.tab(self.tabSaisieInteret, state="normal")
        else:
            self.tabs.tab(self.tabSaisieInteret, state="hidden")

        if not wd or not self.state.saisie_pdf_file:
            return
        if not self.state.templates_folder:
            self.varInteretStatus.set("Select Configurations folder first.")
            return

        if self._helper_module_missing("Saisie d'interet"):
            module_name = REQUIRED_HELPER_MODULES["Saisie d'interet"]
            self._log_missing_helper_modules_once()
            self.varInteretStatus.set(
                f"Missing helper module: {module_name}.py. Saisie d'interet is disabled."
            )
            return

        template_xlsx = find_saisie_interet_template(self.state.templates_folder)
        if template_xlsx:
            self.varInteretTemplate.set(template_xlsx.name)
        else:
            self.varInteretTemplate.set("Not found")
            self.varInteretStatus.set(
                "No Saisie d'intérêt template found.\n"
                "Add an Excel file with 'saisie' and 'interet' in the name to your Configurations folder."
            )
            messagebox.showwarning(
                "Missing Template",
                "No Saisie d'intérêt Excel template found in the Configurations folder.\n\n"
                "Add an Excel file with 'saisie' and 'interet' in its filename to:\n"
                f"{self.state.templates_folder}",
                parent=self.root,
            )
            return

        try:
            case_paths = ensure_case_structure(wd, self.state.saisie_pdf_file)
            cached = read_values_latest_json(case_paths["values_latest_json"])
            base_vals = self._cached_k138_values(case_paths)
            if (not base_vals) and isinstance(self.state.last_k138_values, dict):
                base_vals = {
                    str(k): "" if v is None else str(v)
                    for k, v in self.state.last_k138_values.items()
                }

            out_raw = clean_value(str(cached.get("saisie_interet_output", "") or ""))
            if out_raw:
                try:
                    out_path = Path(out_raw).resolve()
                    if out_path.exists():
                        self.varInteretOutput.set(str(out_path))
                except Exception:
                    pass
            if self.varInteretOutput.get() == "-":
                fallback_out = wd / f"Saisie_interet_{self.state.saisie_pdf_file.stem}.xlsx"
                if fallback_out.exists():
                    self.varInteretOutput.set(str(fallback_out))

            if not base_vals:
                self.varInteretStatus.set("Run Extract Values from Saise A faire first.")
                return

            self.varInteretStatus.set("Ready: generate Saisie d'interet from extracted values.")
            self.btnGenerateSaisieInteret.configure(state="normal")
        except Exception as e:
            self.varInteretStatus.set(f"Status error: {e}")
            self.btnGenerateSaisieInteret.configure(state="disabled")

    def _cache_is_valid_for_current_file(self, cached: Dict[str, object]) -> bool:
        """Return True only if the cached JSON was written for the currently loaded SAISIE file."""
        if not self.state.extraction_ran:
            return False
        if not self.state.saisie_pdf_file:
            return False
        cached_source = clean_value(str(cached.get("source_file", "") or ""))
        if not cached_source:
            return False
        try:
            return Path(cached_source).resolve() == self.state.saisie_pdf_file.resolve()
        except Exception:
            return False

    def _refresh_agenda_status(self):
        wd = self._resolve_working_dir()
        inv = _normalize_inventory_number(self.state.last_inventory_number or "")
        agent = self._normalize_agent_id(self.state.last_agent_id or "")
        agenda_exists = False
        has_k138_base = False
        sied_val = ""
        self.varAgendaWorkingDir.set(wd.name if wd else "-")
        self._refresh_case_folder_banner()
        if wd and self.state.saisie_pdf_file:
            case_paths = ensure_case_structure(wd, self.state.saisie_pdf_file)
            cached = read_values_latest_json(case_paths["values_latest_json"])
            # Only use cached values if the JSON was written FOR THIS SPECIFIC FILE.
            # Prevents prior-session data from bleeding into a freshly loaded file.
            cache_valid = self._cache_is_valid_for_current_file(cached)
            base_vals = self._cached_k138_values(case_paths) if cache_valid else {}
            has_k138_base = bool(base_vals)
            if cache_valid:
                # Inventory must prefer SAISIE-extracted/base values over Agenda edits.
                extracted_inv = _normalize_inventory_number(base_vals.get("description_inventory", ""))
                cached_inv = _normalize_inventory_number(str(cached.get("inventory_number", "") or ""))
                if extracted_inv:
                    inv = extracted_inv
                elif not inv:
                    inv = cached_inv

                extracted_agent = self._normalize_agent_id(base_vals.get("seizing_officer", ""))
                cached_agent = self._normalize_agent_id(str(cached.get("agent_id", "") or ""))
                if extracted_agent:
                    agent = extracted_agent
                elif not agent:
                    agent = cached_agent
            agenda_path = self._agenda_existing_path(case_paths)
            agenda_exists = bool(agenda_path and agenda_path.exists())
            cached_sied_raw = clean_value(str(cached.get("sied_number", "") or "")) if cache_valid else ""
            cached_sied = _normalize_sied_value(cached_sied_raw) if cached_sied_raw else ""
            cached_sied_confirmed = self._as_bool(cached.get("sied_confirmed", False)) if cache_valid else False
            if cached_sied and (not _is_valid_sied_value(cached_sied)):
                cached_sied = ""
            if cached_sied and inv and (_compact_alnum_token(cached_sied) == _compact_alnum_token(inv)):
                cached_sied = ""
            sied_val = cached_sied if cached_sied_confirmed else ""
            patch: Dict[str, object] = {}
            if agenda_path:
                agenda_vals = extract_agenda_core_values(agenda_path)
                agenda_inv = _normalize_inventory_number(agenda_vals.get("inventory_number", ""))
                agenda_agent = self._normalize_agent_id(agenda_vals.get("agent_id", ""))
                # Agenda inventory is fallback only; SAISIE-extracted inventory stays authoritative.
                if agenda_inv and (not inv):
                    inv = agenda_inv
                if agenda_agent:
                    agent = agenda_agent

                agenda_sied = _normalize_sied_value(agenda_vals.get("sied_number", ""), allow_raw=True)
                if agenda_sied and (not _is_valid_sied_value(agenda_sied)):
                    agenda_sied = ""
                if agenda_sied and inv and (_compact_alnum_token(agenda_sied) == _compact_alnum_token(inv)):
                    agenda_sied = ""

                cached_inv = _normalize_inventory_number(str(cached.get("inventory_number", "") or ""))
                cached_agent = self._normalize_agent_id(str(cached.get("agent_id", "") or ""))
                if inv and (inv != cached_inv):
                    patch["inventory_number"] = inv
                if agent and (agent != cached_agent):
                    patch["agent_id"] = agent
                if patch and base_vals:
                    base_updated = dict(base_vals)
                    if "agent_id" in patch:
                        base_updated["seizing_officer"] = agent
                    patch["k138_values_base"] = base_updated

                # Keep status refresh read-only. Agenda rewrites/barcode refresh happen
                # only on explicit user actions (Generate Agenda / Update Values from Agenda).

                # Agenda file is the source of truth. Use cache only when Agenda extraction is empty.
                if agenda_sied:
                    sied_val = agenda_sied
                    if agenda_sied != cached_sied_raw:
                        patch["sied_number"] = sied_val
                        patch["sied_confirmed"] = True
                        patch["sied_source"] = "agenda-read"
                elif cached_sied_confirmed and cached_sied:
                    sied_val = cached_sied

            if patch:
                patch["updated_at"] = _timestamp_iso()
                update_values_latest_json(case_paths["values_latest_json"], patch)
            self.varAgendaFound.set("yes" if agenda_exists else "no")
        else:
            self.varAgendaFound.set("no")
        self.state.last_inventory_number = inv or self.state.last_inventory_number
        self.state.last_agent_id = agent or self.state.last_agent_id
        self.state.last_sied_number = sied_val
        # Clerk workflow: ready as soon as the Saisie D'affaire is loaded in the folder.
        # Agenda and K138 are both available in any order — no Extract Values step required.
        role = getattr(self, "profile_role", "BSO") or "BSO"
        if role == "Clerk":
            clerk_saisie = self.state.saisie_pdf_file
            if clerk_saisie and clerk_saisie.exists():
                has_k138_base = True   # Can generate K138 directly from Saisie D'affaire
            clerk_agenda = getattr(self.state, "clerk_agenda_file", None)
            if clerk_agenda and clerk_agenda.exists():
                agenda_exists = True   # An Agenda was already generated in this folder

        self.varAgendaInventory.set(inv or "-")
        self.varAgendaAgent.set(agent or "-")
        self.varK138AgendaReady.set("yes" if agenda_exists else "no")
        self.varK138Sied.set(sied_val or "-")
        self.btnRefreshAgenda.configure(state="normal" if agenda_exists else "disabled")

        # K138 gate: Clerk unlocks K138 as soon as Saisie D'affaire is loaded.
        # Non-Clerk still requires an Agenda to exist first.
        k138_unlocked = has_k138_base if (role == "Clerk") else agenda_exists
        if k138_unlocked:
            self.lblK138Gate.grid_remove()
            self.frmK138Actions.grid()
        else:
            self.frmK138Actions.grid_remove()
            self.lblK138Gate.grid()

        # K138 tab state: never disable for roles that are allowed to see it.
        # Disabling the current tab causes the notebook to jump away.
        # The gate message and button state inside the tab handle the UX instead.
        if self._role_allows_tab("K138"):
            self.tabs.tab(self.tabK138, state="normal")
            if role == "Clerk" and has_k138_base:
                self.btnGenerateK138.configure(state="normal")
                if agenda_exists:
                    self.varK138Status.set("Ready: Generate K138 using Agenda CE/CID + Saisie D'affaire values.")
                else:
                    self.varK138Status.set("Ready: Generate K138 from Saisie D'affaire. (CE/CID field will be blank — fill it after printing if needed.)")
            elif agenda_exists and has_k138_base:
                self.btnGenerateK138.configure(state="normal")
                self.varK138Status.set("Ready: Agenda exists. Generate K138 from SAISIE + Agenda CE/CID.")
            elif agenda_exists:
                self.btnGenerateK138.configure(state="disabled")
                self.varK138Status.set("Run Extract Values from Saisie A faire first to prepare SAISIE values for K138.")
            else:
                self.btnGenerateK138.configure(state="disabled")
                self.varK138Status.set("Create Agenda first. K138 is downstream of Agenda.")
        else:
            self.tabs.tab(self.tabK138, state="hidden")
            self.btnGenerateK138.configure(state="disabled")

        if self._helper_module_missing("K138"):
            module_name = REQUIRED_HELPER_MODULES["K138"]
            self._log_missing_helper_modules_once()
            self.btnGenerateK138.configure(state="disabled")
            self.varK138Status.set(
                f"Missing helper module: {module_name}.py. K138 generation is disabled."
            )

        self._refresh_saisie_interet_status()
        self._apply_role_tab_visibility()   # always re-enforce role restrictions last
        self._refresh_instruction_feedback()

    # ------------------------------------------------------------------ Narrative
    def on_generate_narrative(self):
        """Load Narrative template and fill with extracted case values."""
        wd = self._resolve_working_dir()
        if not wd:
            messagebox.showwarning("Missing Case", "Select a case folder first.", parent=self.root)
            return
        if not self.state.templates_folder:
            messagebox.showwarning("Missing Folder", "Select Configurations folder first.", parent=self.root)
            return

        # Find Narrative template in templates folder
        lang = getattr(self, "varNarrativeLang", None)
        lang_code = lang.get() if lang else "EN"
        form_type = self.state.form_type or "Stupefiant-Others"

        narrative_template = None
        # Look for template matching language and form type
        for candidate in self.state.templates_folder.iterdir() if self.state.templates_folder.exists() else []:
            n = candidate.name.upper()
            if "NARRATIVE" in n and candidate.suffix.lower() == ".txt":
                narrative_template = candidate
                break

        template_text = ""
        if narrative_template and narrative_template.exists():
            try:
                template_text = narrative_template.read_text(encoding="utf-8", errors="replace")
            except Exception:
                template_text = ""
        else:
            self.log("  ! No Narrative template (.txt) found in templates folder — showing blank.")

        # Substitute placeholders with extracted values
        if wd and self.state.saisie_pdf_file:
            try:
                case_paths = ensure_case_structure(wd, self.state.saisie_pdf_file)
                base_vals = self._cached_k138_values(case_paths)
                subs = {
                    "$DESCRIPTION_FULL":  base_vals.get("description_item_full", "") or base_vals.get("description_item", ""),
                    "$DESCRIPTION_SHORT": base_vals.get("description_item", ""),
                    "$INVENTORY":         base_vals.get("description_inventory", ""),
                    "$OFFICER":           base_vals.get("seizing_officer", ""),
                    "$DATE":              base_vals.get("seizure_date_line", ""),
                    "$LOCATION":          base_vals.get("lieu_interception", ""),
                    "$DECLARED":          base_vals.get("description_declared", ""),
                    "$SEIZURE_TYPE":      form_type,
                    "$LANG":              lang_code,
                }
                for placeholder, value in subs.items():
                    template_text = template_text.replace(placeholder, value or "")
            except Exception as e:
                self.log(f"  ! Narrative substitution error: {e}")

        if hasattr(self, "txtNarrative"):
            self.txtNarrative.delete("1.0", "end")
            self.txtNarrative.insert("1.0", template_text)
        if hasattr(self, "varNarrativeStatus"):
            self.varNarrativeStatus.set("Generated." if template_text else "No template found.")

    def _on_narrative_copy(self):
        """Copy Narrative text to clipboard."""
        if not hasattr(self, "txtNarrative"):
            return
        text = self.txtNarrative.get("1.0", "end").strip()
        if text:
            self.root.clipboard_clear()
            self.root.clipboard_append(text)
            if hasattr(self, "varNarrativeStatus"):
                self.varNarrativeStatus.set("Copied to clipboard.")

    def on_generate_saisie_interet(self):
        wd = self._resolve_working_dir()
        if not wd:
            messagebox.showwarning("Missing Case", "Select SAISIE input file first to determine active case folder.")
            return
        if not self.state.saisie_pdf_file:
            messagebox.showwarning("Missing File", "Select SAISIE input file first.")
            return
        if not self.state.templates_folder:
            messagebox.showwarning("Missing Folder", "Select Configurations folder first.")
            return

        if self._helper_module_missing("Saisie d'interet"):
            module_name = REQUIRED_HELPER_MODULES["Saisie d'interet"]
            self._log_missing_helper_modules_once()
            messagebox.showerror(
                "Missing Module",
                f"Missing helper module: {module_name}.py\n\nExpected: {helper_module_hint(module_name)}",
            )
            return

        try:
            from fill_saisie_interet import fill_saisie_interet
        except ModuleNotFoundError as import_err:
            missing = (getattr(import_err, "name", "") or "").strip()
            if missing == "openpyxl":
                messagebox.showerror(
                    "Missing Dependency",
                    "openpyxl is not installed.\n\nRun:\npython -m pip install openpyxl",
                )
            elif missing == "fill_saisie_interet":
                messagebox.showerror(
                    "Missing Module",
                    f"Missing helper module: fill_saisie_interet.py\n\nExpected: {helper_module_hint('fill_saisie_interet')}",
                )
            else:
                messagebox.showerror(
                    "Missing Dependency",
                    f"Could not import required module: {missing or import_err}",
                )
            return
        except Exception as import_err:
            messagebox.showerror(
                "Missing Module",
                f"Could not load fill_saisie_interet.\n\n{import_err}",
            )
            return

        self._set_busy(True, "Generating Saisie d'interet...")
        try:
            case_paths = ensure_case_structure(wd, self.state.saisie_pdf_file)
            cached = read_values_latest_json(case_paths["values_latest_json"])
            base_values = self._cached_k138_values(case_paths)
            if (not base_values) and isinstance(self.state.last_k138_values, dict):
                base_values = {
                    str(k): "" if v is None else str(v)
                    for k, v in self.state.last_k138_values.items()
                }
            if not base_values:
                messagebox.showwarning(
                    "Missing SAISIE Data",
                    "Run Extract Values from Saise A faire first before generating Saisie d'interet.",
                )
                return

            source_raw = clean_value(str(cached.get("source_file", "") or ""))
            source_path = Path(source_raw) if source_raw else self.state.saisie_pdf_file
            if not source_path.exists():
                source_path = self.state.saisie_pdf_file
            if not source_path or (not source_path.exists()):
                messagebox.showerror("Missing Source", "Could not locate source SAISIE file for generation.")
                return

            notice_text = load_notice_text(self.state.templates_folder, self.state.form_type)
            saisie_template: Optional[Path] = None
            if (source_path.suffix or "").lower() == ".pdf":
                saisie_template = find_saisie_template(self.state.templates_folder)

            top, _bottom, extract_mode = extract_field_values(saisie_template, source_path)
            extracted_values = build_k138_values_from_saisie(top, source_path, self.state.form_type, notice_text)
            merged_values = dict(extracted_values)
            merged_values.update(base_values)

            rows = build_saisie_interet_rows(
                top=top,
                values=merged_values,
                case_folder_name=self.state.case_folder_name or self.state.saisie_pdf_file.stem,
            )
            if not rows:
                messagebox.showwarning(
                    "No Values",
                    "Could not map any values for Saisie d'interet. Check extraction first.",
                )
                return

            template_xlsx = find_saisie_interet_template(self.state.templates_folder)
            if not template_xlsx:
                messagebox.showerror(
                    "Missing Template",
                    f"Could not find Saisie d'interet template (.xlsx) in:\n{self.state.templates_folder}",
                )
                return

            output_xlsx = wd / f"Saisie_interet_{self.state.saisie_pdf_file.stem}.xlsx"
            fill_saisie_interet(
                template_xlsx=template_xlsx,
                output_xlsx=output_xlsx,
                data_by_row=rows,
                sheet_name=None,
            )

            update_values_latest_json(
                case_paths["values_latest_json"],
                {
                    "updated_at": _timestamp_iso(),
                    "source_file": str(source_path),
                    "working_directory": str(wd),
                    "case_folder_name": self.state.case_folder_name or wd.name,
                    "extract_mode": extract_mode,
                    "saisie_interet_output": str(output_xlsx),
                    "saisie_interet_template": str(template_xlsx),
                },
            )
            self.log(f"Saisie d'interet generated: {output_xlsx}")
            self._refresh_saisie_interet_status()
            messagebox.showinfo("Saisie d'interet", f"Generated:\n{output_xlsx}")
        except Exception as e:
            self.log(f"Saisie d'interet generation failed: {e}")
            import traceback
            self.log(traceback.format_exc())
            messagebox.showerror("Saisie d'interet Error", str(e))
        finally:
            self._set_busy(False)

    def on_clerk_select_case_folder(self):
        """Clerk workflow: select case folder — auto-finds Saisie D'affaire inside it.
        Enables Generate Agenda and Generate K138 immediately, in any order.
        """
        initial_dir = str(self.last_saisie_folder or Path.cwd())
        folder_raw = filedialog.askdirectory(
            title="Select Case Folder (containing Saisie D'affaire PDF)",
            initialdir=initial_dir,
        )
        if not folder_raw:
            return
        case_folder = Path(folder_raw).resolve()

        # Auto-find the Saisie D'affaire PDF in the folder
        saisie_pdfs = sorted(case_folder.glob("*Saisie_D_affaire.pdf"))
        if not saisie_pdfs:
            # Fallback: any PDF with "saisie" in the name
            saisie_pdfs = sorted(
                p for p in case_folder.glob("*.pdf")
                if "saisie" in p.name.lower() and "agenda" not in p.name.lower()
            )
        if not saisie_pdfs:
            messagebox.showwarning(
                "No Saisie D'affaire Found",
                f"No Saisie D'affaire PDF found in:\n{case_folder}\n\n"
                "Make sure the BSO has placed the Saisie D'affaire PDF in this folder.",
            )
            return

        saisie_pdf = saisie_pdfs[0]
        if len(saisie_pdfs) > 1:
            self.log(f"  i Multiple Saisie D'affaire PDFs found; using: {saisie_pdf.name}")

        # Set working dir and source file
        self.state.working_dir = case_folder
        self.state.case_folder_locked = True
        self.last_saisie_folder = case_folder
        set_config_path("paths", "saisie_folder", case_folder)

        self.state.saisie_pdf_file = saisie_pdf
        self._set_entry(self.entSaisieFile, saisie_pdf)
        set_config_path("paths", "last_saisie_file", saisie_pdf)

        # Check if an Agenda already exists (for K138 to use later)
        existing_agenda = self._find_agenda_in_folder(case_folder)
        self.state.clerk_agenda_file = existing_agenda

        self.log(f"Clerk: Case folder: {case_folder.name}")
        self.log(f"  Saisie D'affaire: {saisie_pdf.name}")
        if existing_agenda:
            self.log(f"  Agenda already exists: {existing_agenda.name}")

        self._refresh_case_folder_banner()
        self._refresh_agenda_status()

    def _find_agenda_in_folder(self, folder: Path) -> Optional[Path]:
        """Return the most recently modified Agenda PDF in a case folder, or None."""
        try:
            candidates = sorted(
                (p for p in folder.glob("*.pdf") if "agenda" in p.name.lower()),
                key=lambda p: p.stat().st_mtime,
                reverse=True,
            )
            return candidates[0] if candidates else None
        except Exception:
            return None

    # Keep backward-compat alias (called from DnD / older code paths)
    def on_clerk_select_agenda(self):
        self.on_clerk_select_case_folder()

    def on_fill_agenda(self):
        wd = self._resolve_working_dir()
        if not wd:
            messagebox.showwarning("Missing Case", "Select SAISIE input file first to determine active case folder.")
            return
        if not self.state.saisie_pdf_file:
            messagebox.showwarning("Missing File", "Select SAISIE input file first.")
            return
        if not self.state.templates_folder:
            messagebox.showwarning("Missing Folder", "Select Configurations folder first.")
            return
        self._set_busy(True, "Generating Agenda...")
        try:
            case_paths = ensure_case_structure(wd, self.state.saisie_pdf_file)

            # Bug fix: detect existing Agenda before creating a new one.
            existing_agenda = self._agenda_existing_path(case_paths)
            if existing_agenda and existing_agenda.exists():
                if not messagebox.askyesno(
                    "Agenda Already Exists",
                    f"An Agenda was already found:\n{existing_agenda.name}\n\nDo you want to regenerate it?",
                    parent=self.root,
                ):
                    return

            cached_vals = read_values_latest_json(case_paths["values_latest_json"])
            base_vals = self._cached_k138_values(case_paths)
            inventory_number = _normalize_inventory_number(self.state.last_inventory_number or "")
            if not inventory_number:
                inventory_number = _normalize_inventory_number(str(cached_vals.get("inventory_number", "") or ""))
            if not inventory_number:
                inventory_number = _normalize_inventory_number(base_vals.get("description_inventory", ""))
            agent_id = self._normalize_agent_id(self.state.last_agent_id or "")
            if not agent_id:
                agent_id = self._normalize_agent_id(str(cached_vals.get("agent_id", "") or ""))
            if not agent_id:
                agent_id = self._normalize_agent_id(base_vals.get("seizing_officer", ""))

            if (not inventory_number) or (not agent_id):
                source_file_raw = str(self.state.saisie_pdf_file) or str(cached_vals.get("source_file", "") or "")
                source_path = Path(source_file_raw) if source_file_raw else None
                if source_path and source_path.exists():
                    recovered = self._recover_agenda_core_values_from_source(source_path)
                    recovered_inv_raw = clean_value(recovered.get("inventory_number", ""))
                    recovered_inv = _normalize_inventory_number(recovered_inv_raw)
                    if not recovered_inv:
                        recovered_compact = re.sub(r"[^A-Za-z0-9]+", "", recovered_inv_raw).upper()
                        if len(recovered_compact) >= 10 and re.search(r"\d", recovered_compact):
                            recovered_inv = recovered_compact
                    recovered_agent = self._normalize_agent_id(recovered.get("agent_id", ""))

                    if (not inventory_number) and recovered_inv:
                        inventory_number = recovered_inv
                        self.log(f"  i Recovered inventory from source: {inventory_number}")
                    if (not agent_id) and recovered_agent:
                        agent_id = recovered_agent
                        self.log(f"  i Recovered agent from source: {agent_id}")

                    if recovered_inv or recovered_agent:
                        merged_vals = dict(base_vals)
                        if recovered_inv:
                            merged_vals["description_inventory"] = recovered_inv
                        if recovered_agent:
                            merged_vals["seizing_officer"] = recovered_agent
                        update_patch: Dict[str, object] = {
                            "updated_at": _timestamp_iso(),
                            "k138_values_base": merged_vals,
                        }
                        if recovered_inv:
                            update_patch["inventory_number"] = recovered_inv
                        if recovered_agent:
                            update_patch["agent_id"] = recovered_agent
                        update_values_latest_json(case_paths["values_latest_json"], update_patch)

            self.state.last_inventory_number = inventory_number
            self.state.last_agent_id = agent_id
            self._refresh_agenda_status()

            if not inventory_number:
                manual_inventory = simpledialog.askstring(
                    "Inventory Number",
                    "Inventory number was not detected.\n\nEnter inventory number to generate barcode and Agenda:",
                    parent=self.root,
                    initialvalue=str(cached_vals.get("inventory_number", "") or ""),
                )
                if manual_inventory is None:
                    return
                manual_inventory = clean_value(manual_inventory)
                manual_norm = _normalize_inventory_number(manual_inventory)
                if not manual_norm:
                    compact = re.sub(r"[^A-Za-z0-9]+", "", manual_inventory).upper()
                    if len(compact) >= 8 and re.search(r"[A-Z]", compact) and re.search(r"\d", compact):
                        manual_norm = compact
                if not manual_norm:
                    messagebox.showerror(
                        "Invalid Inventory",
                        "Inventory format was not recognized. Please enter a valid inventory number.",
                    )
                    return
                inventory_number = manual_norm
                self.state.last_inventory_number = inventory_number
                self.log(f"  i Using manual inventory for Agenda/barcode: {inventory_number}")
            if not agent_id:
                messagebox.showerror("Missing Agent", "Could not detect agent ID from processed SAISIE values.")
                return

            barcode_path = generate_barcode(inventory_number, case_paths["barcode_png"])
            # Agenda SIED/CE-CID must start blank; officers fill it after agenda creation.
            sied_number = ""
            agenda_template = find_agenda_template(self.state.templates_folder)
            if not agenda_template:
                messagebox.showerror("Missing Agenda Template", f"Could not find Agenda template in:\n{self.state.templates_folder}")
                return

            agenda_internal_out = fill_agenda(
                template_path=agenda_template,
                case_paths=case_paths,
                agent_id=agent_id,
                inventory_number=inventory_number,
                sied_number=sied_number,
                barcode_png=barcode_path,
            )
            if agenda_internal_out.suffix.lower() == ".pdf":
                agenda_client_out = case_paths["agenda_output_pdf"]
            else:
                agenda_client_out = case_paths["agenda_output_docx"]

            agenda_client_out.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(agenda_internal_out, agenda_client_out)

            update_values_latest_json(
                case_paths["values_latest_json"],
                {
                    "updated_at": _timestamp_iso(),
                    "source_file": str(self.state.saisie_pdf_file),
                    "working_directory": str(wd),
                    "case_folder_name": self.state.case_folder_name or wd.name,
                    "form_type": self.state.form_type,
                    "inventory_number": inventory_number,
                    "agent_id": agent_id,
                    "sied_number": sied_number,
                    "sied_confirmed": False,
                    "sied_source": "",
                    "agenda_output": str(agenda_client_out),
                    "agenda_latest": str(agenda_internal_out),
                }
            )

            self.log(f"Agenda internal latest: {agenda_internal_out}")
            self.log(f"Agenda client output: {agenda_client_out}")
            print(f"Agenda internal latest: {agenda_internal_out}")
            print(f"Agenda client output: {agenda_client_out}")
            # For Clerk: track the generated Agenda so K138 can use it (with CE/CID filled later)
            _role = getattr(self, "profile_role", "BSO") or "BSO"
            if _role == "Clerk" and agenda_client_out.exists():
                self.state.clerk_agenda_file = agenda_client_out
            self._refresh_agenda_status()
            messagebox.showinfo("Agenda", f"Agenda generated:\n{agenda_client_out}")
        except Exception as e:
            self.log(f"Agenda fill failed: {e}")
            import traceback
            self.log(traceback.format_exc())
            messagebox.showerror("Agenda Error", str(e))
        finally:
            self._set_busy(False)

    def on_update_agenda(self):
        wd = self._resolve_working_dir()
        if not wd:
            messagebox.showwarning("Missing Case", "Select SAISIE input file first to determine active case folder.")
            return
        if not self.state.saisie_pdf_file:
            messagebox.showwarning("Missing File", "Select SAISIE input file first.")
            return
        try:
            case_paths = ensure_case_structure(wd, self.state.saisie_pdf_file)
            agenda_path = self._agenda_existing_path(case_paths)
            if not agenda_path:
                messagebox.showwarning("Missing Agenda", "Create Agenda first before editing values.")
                return

            cached = read_values_latest_json(case_paths["values_latest_json"])
            agenda_vals = extract_agenda_core_values(agenda_path)

            current_inventory = _normalize_inventory_number(agenda_vals.get("inventory_number", ""))
            if not current_inventory:
                current_inventory = _normalize_inventory_number(self.state.last_inventory_number or "")
            if not current_inventory:
                current_inventory = _normalize_inventory_number(str(cached.get("inventory_number", "") or ""))

            current_agent = self._normalize_agent_id(agenda_vals.get("agent_id", ""))
            if not current_agent:
                current_agent = self._normalize_agent_id(self.state.last_agent_id or "")
            if not current_agent:
                current_agent = self._normalize_agent_id(str(cached.get("agent_id", "") or ""))

            cached_sied_raw = clean_value(str(cached.get("sied_number", "") or ""))
            cached_sied = _normalize_sied_value(cached_sied_raw) if cached_sied_raw else ""
            if cached_sied and (not _is_valid_sied_value(cached_sied)):
                cached_sied = ""
            current_sied = _normalize_sied_value(agenda_vals.get("sied_number", ""), allow_raw=True) or cached_sied
            if current_sied and (not _is_valid_sied_value(current_sied)):
                current_sied = ""

            entered_inventory = simpledialog.askstring(
                "Edit Agenda",
                "Inventory number (used for barcode):",
                parent=self.root,
                initialvalue=current_inventory or "",
            )
            if entered_inventory is None:
                return
            inventory_source = clean_value(entered_inventory) or current_inventory
            inventory_number = _normalize_inventory_number(inventory_source)
            if not inventory_number:
                messagebox.showerror(
                    "Invalid Inventory",
                    "Inventory number is required to regenerate the barcode.",
                )
                return

            entered_agent = simpledialog.askstring(
                "Edit Agenda",
                "Agent badge number:",
                parent=self.root,
                initialvalue=current_agent or "",
            )
            if entered_agent is None:
                return
            agent_source = clean_value(entered_agent) or current_agent
            agent_id = self._normalize_agent_id(agent_source)
            if not agent_id:
                messagebox.showerror("Invalid Agent", "Agent badge number is required.")
                return

            entered_sied = simpledialog.askstring(
                "Edit Agenda CE/CID",
                "Enter CE/CID (SIED) value (optional):",
                parent=self.root,
                initialvalue=current_sied or "",
            )
            if entered_sied is None:
                return
            sied_number = _normalize_sied_value(entered_sied, allow_raw=True)

            if sied_number and (_compact_alnum_token(sied_number) == _compact_alnum_token(inventory_number)):
                self.log("  i Ignoring CE/CID because it matches inventory number.")
                sied_number = ""

            self._set_busy(True, "Updating Agenda...")
            try:
                barcode_path = generate_barcode(inventory_number, case_paths["barcode_png"])
                update_agenda(
                    agenda_path=agenda_path,
                    agent_id=agent_id,
                    inventory_number=inventory_number,
                    sied_number=sied_number,
                    barcode_png=barcode_path,
                )

                # Keep internal/client agenda files in sync.
                if agenda_path.suffix.lower() == ".pdf":
                    sync_targets = [case_paths["agenda_output_pdf"], case_paths["agenda_latest_pdf"]]
                else:
                    sync_targets = [case_paths["agenda_output_docx"], case_paths["agenda_latest_docx"]]
                for dst in sync_targets:
                    try:
                        if agenda_path.resolve() != dst.resolve():
                            dst.parent.mkdir(parents=True, exist_ok=True)
                            shutil.copy2(agenda_path, dst)
                    except Exception as sync_err:
                        self.log(f"  ! Agenda sync warning ({dst.name}): {sync_err}")

                self.state.last_inventory_number = inventory_number
                self.state.last_agent_id = agent_id
                self.state.last_sied_number = sied_number
                base_vals = self._cached_k138_values(case_paths)
                if base_vals:
                    base_vals["seizing_officer"] = agent_id
                update_patch: Dict[str, object] = {
                    "updated_at": _timestamp_iso(),
                    "source_file": str(self.state.saisie_pdf_file),
                    "working_directory": str(wd),
                    "case_folder_name": self.state.case_folder_name or wd.name,
                    "inventory_number": inventory_number,
                    "agent_id": agent_id,
                    "sied_number": sied_number,
                    "sied_confirmed": bool(clean_value(sied_number)),
                    "sied_source": "agenda-update" if clean_value(sied_number) else "",
                    "agenda_output": str(sync_targets[0]),
                    "agenda_latest": str(sync_targets[1]),
                }
                if base_vals:
                    update_patch["k138_values_base"] = base_vals
                update_values_latest_json(case_paths["values_latest_json"], update_patch)

                self.log(
                    "Agenda updated: "
                    f"inventory={inventory_number}, "
                    f"agent={agent_id}, "
                    f"CE/CID={sied_number or '(blank)'}"
                )
                self._refresh_agenda_status()
                messagebox.showinfo("Agenda", f"Agenda updated:\n{sync_targets[0]}")
            finally:
                self._set_busy(False)
        except Exception as e:
            self.log(f"Agenda update failed: {e}")
            import traceback
            self.log(traceback.format_exc())
            messagebox.showerror("Agenda Error", str(e))

    def _on_tab_changed(self, event=None):
        """Refresh visible controls/status and keep instruction text in sync with active tab."""
        try:
            selected = self.tabs.select()
            self._apply_step_visibility()
            if selected in (str(self.tabAgenda), str(self.tabK138), str(self.tabSaisieInteret)):
                self._set_busy(True, "Refreshing status...")
                try:
                    self._refresh_agenda_status()
                finally:
                    self._set_busy(False)
            else:
                self._refresh_instruction_feedback()
        except Exception:
            try:
                self._refresh_instruction_feedback()
            except Exception:
                pass

    def _on_refresh_agenda_from_pdf(self):
        """Re-read the Agenda PDF, detect any inventory edits the user made in their PDF viewer,
        and regenerate the barcode + replace the PDF if the inventory number changed."""
        wd = self._resolve_working_dir()
        if not wd or not self.state.saisie_pdf_file:
            messagebox.showwarning("Missing Case", "Select SAISIE input file first.")
            return
        self._set_busy(True, "Refreshing Agenda values...")
        try:
            case_paths = ensure_case_structure(wd, self.state.saisie_pdf_file)
            agenda_path = self._agenda_existing_path(case_paths)
            if not agenda_path or not agenda_path.exists():
                messagebox.showwarning("No Agenda", "No agenda file found. Generate one first.")
                return

            # Read current values from the agenda PDF (including any user edits in their viewer)
            agenda_vals = extract_agenda_core_values(agenda_path)
            agenda_inv = _normalize_inventory_number(agenda_vals.get("inventory_number", ""))
            agenda_agent = self._normalize_agent_id(agenda_vals.get("agent_id", ""))
            # Read SIED directly — trust whatever the user typed, no format validation.
            agenda_sied = _read_sied_raw_from_agenda_pdf(agenda_path)

            cached = read_values_latest_json(case_paths["values_latest_json"])
            cached_inv = _normalize_inventory_number(str(cached.get("inventory_number", "") or ""))
            cached_agent = self._normalize_agent_id(str(cached.get("agent_id", "") or ""))
            base_vals = self._cached_k138_values(case_paths)
            extracted_inv = _normalize_inventory_number(base_vals.get("description_inventory", ""))
            base_agent = self._normalize_agent_id(base_vals.get("seizing_officer", ""))
            recovered_inv = ""
            recovered_agent = ""

            source_raw = clean_value(str(cached.get("source_file", "") or ""))
            source_path = Path(source_raw) if source_raw else self.state.saisie_pdf_file
            if source_path and source_path.exists():
                recovered = self._recover_agenda_core_values_from_source(source_path)
                recovered_inv = _normalize_inventory_number(recovered.get("inventory_number", ""))
                recovered_agent = self._normalize_agent_id(recovered.get("agent_id", ""))

            # Explicit user action ("Update Values from Agenda"):
            # Agenda edits win for the CURRENT case; SAISIE/base is fallback.
            inv = agenda_inv or recovered_inv or extracted_inv or cached_inv or _normalize_inventory_number(self.state.last_inventory_number or "")
            agent = agenda_agent or recovered_agent or self._normalize_agent_id(self.state.last_agent_id or "") or cached_agent or base_agent

            if not inv:
                messagebox.showwarning("No Inventory", "Could not resolve inventory number from SAISIE/base values.")
                return

            # Always regenerate barcode and rewrite the agenda so it matches current inventory
            self.log(f"Refresh from Agenda PDF: inventory={inv}, agent={agent}")
            barcode_path = generate_barcode(inv, case_paths["barcode_png"])
            update_agenda(
                agenda_path=agenda_path,
                agent_id=agent,
                inventory_number=inv,
                sied_number=agenda_sied or "",
                barcode_png=barcode_path,
            )
            # Sync both client and internal copies
            sync_agenda_files(case_paths)

            # Persist updated values to JSON
            base_vals_patch = dict(base_vals) if base_vals else {}
            if inv:
                # Preserve explicit agenda update for this case so K138/barcode stay in sync.
                base_vals_patch["description_inventory"] = inv
            if base_vals_patch and agent:
                base_vals_patch["seizing_officer"] = agent
            patch: Dict[str, object] = {
                "updated_at": _timestamp_iso(),
                "inventory_number": inv,
            }
            if agenda_inv:
                patch["inventory_source"] = "agenda-refresh"
            if agent:
                patch["agent_id"] = agent
            if agenda_sied:
                patch["sied_number"] = agenda_sied
                patch["sied_confirmed"] = True
            if base_vals_patch:
                patch["k138_values_base"] = base_vals_patch
            update_values_latest_json(case_paths["values_latest_json"], patch)

            self.state.last_inventory_number = inv
            if agent:
                self.state.last_agent_id = agent
            self._refresh_agenda_status()

            sied_msg = f"\nSeizure number (CE/CID): {agenda_sied}" if agenda_sied else ""
            messagebox.showinfo("Refreshed", f"Agenda refreshed.\nInventory: {inv}{sied_msg}\nBarcode regenerated.\n\n{agenda_path}")

            # If K138 was previously generated, regenerate it now so the seizure number
            # (and any other updated Agenda values) are reflected in the K138 PDF.
            k138_prior = resolve_latest_k138_pdf(wd, case_paths)
            if k138_prior and k138_prior.exists():
                self.log("  i K138 previously generated — regenerating with updated Agenda values...")
                self._run_k138_from_agenda(agenda_path, "K138 updated from Agenda refresh")
        except Exception as e:
            self.log(f"Refresh from Agenda failed: {e}")
            import traceback
            self.log(traceback.format_exc())
            messagebox.showerror("Refresh Error", str(e))
        finally:
            self._set_busy(False)

    def _run_k138_from_agenda(self, agenda_path: Path, mode_label: str):
        role = getattr(self, "profile_role", "BSO") or "BSO"
        is_clerk = (role == "Clerk")

        wd = self._resolve_working_dir()
        if not wd:
            messagebox.showwarning("Missing Case", "Select a case folder or agenda first.")
            return
        # Non-clerk requires a SAISIE file
        if not is_clerk and not self.state.saisie_pdf_file:
            messagebox.showwarning("Missing Case", "Select a SAISIE file and run Extract Values from Saisie A faire first.")
            return
        if not self.state.templates_folder:
            messagebox.showwarning("Missing Folder", "Select Configurations folder first.")
            return
        if self._helper_module_missing("K138"):
            module_name = REQUIRED_HELPER_MODULES["K138"]
            self._log_missing_helper_modules_once()
            messagebox.showerror(
                "Missing Module",
                f"Missing helper module: {module_name}.py\n\nExpected: {helper_module_hint(module_name)}",
            )
            return
        if not agenda_path.exists():
            messagebox.showerror("Agenda Missing", f"Agenda file not found:\n{agenda_path}")
            return

        # For Clerk: use agenda's parent as working dir; for others use normal saisie_pdf_file path
        source_file = self.state.saisie_pdf_file or agenda_path
        # Track whether we were given a real Agenda or just the Saisie D'affaire as fallback.
        # Used later to avoid saving the Saisie D'affaire path as "agenda_output".
        _original_agenda_path = agenda_path
        _passed_saisie_as_agenda = source_file and (_original_agenda_path.resolve() == source_file.resolve())
        case_paths = ensure_case_structure(wd, source_file)
        # Only sync when we already have a real Agenda (not the Saisie-as-fallback path).
        # If we sync while _passed_saisie_as_agenda is True, sync_agenda_files may copy a
        # previously-generated hidden Agenda into the client folder, making it appear as
        # though an Agenda was "generated" during a K138-only operation (the reported bug).
        if not _passed_saisie_as_agenda:
            synced_agenda = sync_agenda_files(case_paths)
            if synced_agenda and synced_agenda.exists():
                agenda_path = synced_agenda
        base_values = self._cached_k138_values(case_paths)

        # Clerk fallback: search all hidden case folders in the working dir for any k138 base data
        if not base_values and is_clerk:
            try:
                hidden_root = get_hidden_data_dir(wd)
                if hidden_root.exists():
                    for case_dir in hidden_root.iterdir():
                        if not case_dir.is_dir():
                            continue
                        vjson = case_dir / "values_latest.json"
                        if vjson.exists():
                            try:
                                data = json.loads(vjson.read_text(encoding="utf-8"))
                                b = data.get("k138_values_base", {})
                                if isinstance(b, dict) and b:
                                    base_values = {str(k): "" if v is None else str(v) for k, v in b.items()}
                                    self.log("  i Clerk: loaded K138 base values from case data")
                                    break
                            except Exception:
                                pass
            except Exception:
                pass

        # For Clerk with no base values: build minimal values from agenda directly
        if not base_values and is_clerk:
            agenda_core = extract_agenda_core_values(agenda_path)
            base_values = {
                "description_inventory": _normalize_inventory_number(agenda_core.get("inventory_number", "")),
                "seizing_officer": self._normalize_agent_id(agenda_core.get("agent_id", "")),
            }
            self.log("  i Clerk: generating K138 from Agenda values only (no prior extraction found)")

        if not base_values:
            messagebox.showerror("Missing SAISIE Data", "Run Extract Values from Saisie A faire first to prepare K138 base values.")
            return

        # Inventory should stay SAISIE-extracted/base; Agenda values are fallback only.
        agenda_core = extract_agenda_core_values(agenda_path)
        if (not _normalize_inventory_number(base_values.get("description_inventory", ""))) and _normalize_inventory_number(agenda_core.get("inventory_number", "")):
            base_values["description_inventory"] = _normalize_inventory_number(agenda_core["inventory_number"])
        if self._normalize_agent_id(agenda_core.get("agent_id", "")):
            base_values["seizing_officer"] = self._normalize_agent_id(agenda_core["agent_id"])

        if not clean_value(base_values.get("description_item", "")):
            source_file_raw = (
                str(self.state.saisie_pdf_file)
                or str(read_values_latest_json(case_paths["values_latest_json"]).get("source_file", ""))
            )
            src = Path(source_file_raw) if source_file_raw else None
            if src and src.exists():
                try:
                    template_pdf: Optional[Path] = None
                    if (src.suffix or "").lower() == ".pdf":
                        template_pdf = find_saisie_template(self.state.templates_folder) if self.state.templates_folder else None
                    notice_text = load_notice_text(self.state.templates_folder, self.state.form_type) if self.state.templates_folder else ""
                    top_ref, _bottom_ref, _mode_ref = extract_field_values(template_pdf, src)
                    refreshed = build_k138_values_from_saisie(top_ref, src, self.state.form_type, notice_text)
                    if clean_value(refreshed.get("description_item", "")):
                        base_values.update(refreshed)
                        update_values_latest_json(
                            case_paths["values_latest_json"],
                            {"updated_at": _timestamp_iso(), "k138_values_base": base_values},
                        )
                        self.log("  i Refreshed K138 values from source file (recovered description item).")
                except Exception as refresh_err:
                    self.log(f"  ! K138 base refresh warning: {refresh_err}")

        cached = read_values_latest_json(case_paths["values_latest_json"])
        source_name_lower = (self.state.saisie_pdf_file.name or "").lower()
        cached_source_name = Path(clean_value(str(cached.get("source_file", "") or ""))).name.lower()
        use_saisie_affaire_manual = (
            source_name_lower.startswith("saisie_d_affaire")
            or cached_source_name.startswith("saisie_d_affaire")
        )
        manual_fields: Dict[str, str] = {}
        if use_saisie_affaire_manual:
            manual_obj = cached.get("saisie_affaire_manual")
            if isinstance(manual_obj, dict):
                fields_obj = manual_obj.get("fields")
                if isinstance(fields_obj, dict):
                    manual_fields = {
                        str(k): clean_value(str(v or ""))
                        for k, v in fields_obj.items()
                    }
            if not manual_fields:
                source_path_raw = clean_value(str(cached.get("source_file", "") or "")) or str(self.state.saisie_pdf_file)
                source_path = Path(source_path_raw) if source_path_raw else None
                if source_path and source_path.exists():
                    manual_fields = extract_saisie_affaire_manual_fields_from_pdf(source_path)
            if manual_fields:
                apply_saisie_affaire_manual_to_k138_values(base_values, manual_fields)

        cached_sied_raw = clean_value(str(cached.get("sied_number", "") or ""))
        # allow_raw=True so user-confirmed values (e.g. set by Update Values from Agenda)
        # are preserved even when the format is unconventional.
        cached_sied = _normalize_sied_value(cached_sied_raw, allow_raw=True) if cached_sied_raw else ""
        cached_sied_confirmed = self._as_bool(cached.get("sied_confirmed", False))
        # Only apply noise filter to unconfirmed cached values; trust explicit user confirmations.
        if cached_sied and (not cached_sied_confirmed) and (not _is_valid_sied_value(cached_sied)):
            cached_sied = ""
        agenda_sied = _normalize_sied_value(agenda_core.get("sied_number", ""), allow_raw=True) or _normalize_sied_value(extract_sied_from_agenda(agenda_path))
        sied_source = ""
        if agenda_sied and (not _is_valid_sied_value(agenda_sied)):
            agenda_sied = ""
        inv_guard = _normalize_inventory_number(base_values.get("description_inventory", "")) or _normalize_inventory_number(str(cached.get("inventory_number", "") or ""))
        if agenda_sied and inv_guard:
            if _compact_alnum_token(agenda_sied) == _compact_alnum_token(inv_guard):
                self.log("  i Ignoring CE/CID detected from Agenda because it matches inventory number.")
                agenda_sied = ""
        if agenda_sied:
            sied_source = "agenda-read"
        if (not agenda_sied) and cached_sied_confirmed and cached_sied_raw:
            # User explicitly confirmed this value via "Update Values from Agenda" —
            # use it directly, no format guardrails.
            agenda_sied = cached_sied_raw
            sied_source = "cache-confirmed"
        if not agenda_sied:
            self.log("  i CE/CID not found in Agenda; leaving seizure number blank.")

        k138_values = dict(base_values)
        # Default blank unless CE/CID was explicitly confirmed.
        k138_values["description_seizure_number"] = ""
        if agenda_sied:
            # CE/CID (Agenda) drives seizure number on K138.
            k138_values["description_seizure_number"] = agenda_sied
            self.state.last_sied_number = agenda_sied

        validation_errors, validation_warnings = validate_k138_values(k138_values)
        if validation_warnings:
            self.log("  ! K138 validation warnings:")
            for w in validation_warnings:
                self.log(f"    - {w}")
        if validation_errors:
            self.log("  X K138 validation errors:")
            for e in validation_errors:
                self.log(f"    - {e}")
            messagebox.showerror(
                "Validation Error",
                "Cannot generate K138 due to missing critical values.\n\n" + "\n".join(validation_errors),
            )
            return

        k138_template = find_k138_template(self.state.templates_folder, self.state.form_type)
        if not k138_template:
            messagebox.showerror(
                "Missing K138 Template",
                f"Could not find K138 template PDF for form type '{self.state.form_type}' in:\n{self.state.templates_folder}",
            )
            return

        base_tmp = get_temp_dir()
        base_tmp.mkdir(parents=True, exist_ok=True)
        with tempfile.TemporaryDirectory(prefix="radiance_copilot_", dir=str(base_tmp)) as temp_dir:
            temp_path = Path(temp_dir)
            out_vals = temp_path / "k138_values.csv"
            write_k138_values_csv(out_vals, k138_values)
            k138_output = case_paths["k138_output_pdf"]
            ran = try_run_k138_filler(k138_template, out_vals, k138_output, self.log)

        if not ran:
            messagebox.showerror("K138 Error", "Failed to generate K138. Check log for details.")
            return

        self.state.last_k138_output = k138_output
        try:
            shutil.copy2(k138_output, case_paths["k138_latest_pdf"])
        except Exception as copy_err:
            self.log(f"  ! Could not copy k138_latest.pdf: {copy_err}")

        _k138_patch: Dict[str, object] = {
            "updated_at": _timestamp_iso(),
            "source_file": str(self.state.saisie_pdf_file),
            "working_directory": str(wd),
            "case_folder_name": self.state.case_folder_name or wd.name,
            "form_type": self.state.form_type,
            "inventory_number": _normalize_inventory_number(k138_values.get("description_inventory", "")),
            "agent_id": self._normalize_agent_id(k138_values.get("seizing_officer", "")),
            "sied_number": agenda_sied,
            "sied_confirmed": bool(clean_value(agenda_sied)),
            "sied_source": sied_source if clean_value(agenda_sied) else "",
            "k138_output": str(k138_output),
            "k138_latest_pdf": str(case_paths["k138_latest_pdf"]),
            "k138_values_base": base_values,
        }
        # Only record agenda_output when we actually used a real Agenda PDF.
        # Avoids falsely marking an Agenda as generated when K138 was built
        # directly from the Saisie D'affaire (Clerk workflow, no prior Agenda).
        if not _passed_saisie_as_agenda:
            _k138_patch["agenda_output"] = str(agenda_path)
        update_values_latest_json(case_paths["values_latest_json"], _k138_patch)
        self.log(f"  ✓ {mode_label}: {k138_output}")
        self._refresh_agenda_status()
        messagebox.showinfo("K138", f"{mode_label}:\n{k138_output}")

    def on_generate_k138(self):
        role = getattr(self, "profile_role", "BSO") or "BSO"

        # Clerk path: use Agenda if it already exists in the folder; otherwise use
        # the Saisie D'affaire directly (CE/CID field will be blank — officer fills it in).
        if role == "Clerk":
            wd = self._resolve_working_dir()
            if not wd or not self.state.saisie_pdf_file:
                messagebox.showwarning("Missing Case", "Select a case folder first (Agenda tab → Select Case Folder).")
                return
            # Prefer an existing Agenda (has CE/CID filled in by officer)
            existing_agenda = getattr(self.state, "clerk_agenda_file", None)
            if not (existing_agenda and existing_agenda.exists()):
                existing_agenda = self._find_agenda_in_folder(wd)
                if existing_agenda:
                    self.state.clerk_agenda_file = existing_agenda
            source_for_k138 = (existing_agenda if (existing_agenda and existing_agenda.exists())
                               else self.state.saisie_pdf_file)
            self._set_busy(True, "Generating K138...")
            try:
                self._run_k138_from_agenda(source_for_k138, "K138 generated")
            finally:
                self._set_busy(False)
            return

        wd = self._resolve_working_dir()
        if not wd or not self.state.saisie_pdf_file:
            messagebox.showwarning("Missing Case", "Select a SAISIE file and run Extract Values from Saisie A faire first.")
            return
        case_paths = ensure_case_structure(wd, self.state.saisie_pdf_file)
        agenda_path = self._agenda_existing_path(case_paths)
        if not agenda_path:
            messagebox.showwarning("Missing Agenda", "Create Agenda first. K138 is generated from SAISIE + Agenda.")
            return
        self._set_busy(True, "Generating K138...")
        try:
            self._run_k138_from_agenda(agenda_path, "K138 generated")
        finally:
            self._set_busy(False)

    def on_update_k138(self):
        initial_dir = str(self.last_saisie_folder) if self.last_saisie_folder else None
        file_path = filedialog.askopenfilename(
            title="Select Agenda file for K138 update",
            filetypes=[
                ("Agenda files", "*.pdf;*.docx"),
                ("PDF files", "*.pdf"),
                ("Word documents", "*.docx"),
                ("All files", "*.*"),
            ],
            initialdir=initial_dir,
        )
        if not file_path:
            return
        agenda_path = Path(file_path)
        self.last_saisie_folder = agenda_path.parent
        self._set_busy(True, "Updating K138...")
        try:
            self._run_k138_from_agenda(agenda_path, "K138 updated")
        finally:
            self._set_busy(False)

    def _pick_folder(self) -> Optional[Path]:
        p = filedialog.askdirectory(title="Select Folder")
        return Path(p) if p else None

    def _set_entry(self, ent: ttk.Entry, p: Optional[Path]):
        ent.delete(0, "end")
        if p:
            ent.insert(0, str(p))

    def on_browse_templates_folder(self):
        initial_dir = str(self.state.templates_folder or self.last_saisie_folder or Path.cwd())
        p_raw = filedialog.askdirectory(
            title="Select templates folder",
            initialdir=initial_dir,
        )
        if not p_raw:
            return
        p = Path(p_raw)
        self.state.templates_folder = p
        set_config_path("paths", "templates_folder", p)
        self._prefill_saisie_affaire_defaults()
        self._refresh_config_summary()
        self.log(f"Configurations updated: folder={p.name}")
        self._refresh_agenda_status()
        self._auto_extract_if_ready()
    
    def on_browse_saisie_file(self):
        # Default to working_case_example if no saved folder, or use saved folder
        if self.last_saisie_folder:
            initial_dir = str(self.last_saisie_folder)
        else:
            # Try to default to working_case_example folder if it exists
            default_working = Path("working_case_example")
            if default_working.exists():
                initial_dir = str(default_working.resolve())
            else:
                initial_dir = None
        file_path = filedialog.askopenfilename(
            title="Select SAISIE PDF, Word, or image file",
            filetypes=[
                ("Supported files", "*.pdf;*.docx;*.png;*.jpg;*.jpeg;*.tif;*.tiff;*.bmp"),
                ("PDF files", "*.pdf"),
                ("Word documents", "*.docx"),
                ("Image files", "*.png;*.jpg;*.jpeg;*.tif;*.tiff;*.bmp"),
                ("All files", "*.*")
            ],
            initialdir=initial_dir
        )
        if file_path:
            p = Path(file_path)
            self._reset_case_runtime_state(keep_case_folder=True)
            self.state.saisie_pdf_file = p
            self._set_entry(self.entSaisieFile, p)
            # Save the folder (not the file) for next time
            if not _is_transient_upload_path(p.parent):
                set_config_path("paths", "saisie_folder", p.parent)
                self.last_saisie_folder = p.parent
            set_config_path("paths", "last_saisie_file", p)
            self._set_working_directory(p, refresh_status=False)
            self.log(f"Input File: {p.name}")
            self._auto_extract_if_ready()

    def _auto_extract_if_ready(self):
        """Auto-trigger value extraction when both templates and SAISIE file are set (BSO/Supervisor only)."""
        role = getattr(self, "profile_role", "BSO") or "BSO"
        if role == "Clerk":
            return  # Clerk never needs extraction
        if not self.state.templates_folder or not self.state.saisie_pdf_file:
            return
        if not self.state.saisie_pdf_file.exists():
            return
        if getattr(self, "_busy", False):
            return
        self.on_process_pdf()

    def on_form_type_changed(self):
        self.state.form_type = self.form_type_var.get()
        self.log(f"Form type: {self.state.form_type}")

    def _on_tab_complete_changed(self, tab: str):
        """Mark a tab as complete — lock (read-only) or unlock the associated PDF."""
        import stat as _stat
        var_map = {
            "saisie":    getattr(self, "varCompleteSaisie",    None),
            "agenda":    getattr(self, "varCompleteAgenda",    None),
            "k138":      getattr(self, "varCompleteK138",      None),
            "narrative": getattr(self, "varCompleteNarrative", None),
        }
        var = var_map.get(tab)
        is_complete = var.get() if var else False

        pdf = self._resolve_tab_pdf(tab)
        if pdf and pdf.exists():
            try:
                if is_complete:
                    # Remove write permission — lock PDF for editing
                    current = pdf.stat().st_mode
                    pdf.chmod(current & ~(_stat.S_IWRITE | _stat.S_IWGRP | _stat.S_IWOTH))
                    self.log(f"  ✓ {tab.capitalize()} marked complete — PDF locked: {pdf.name}")
                else:
                    # Restore write permission
                    current = pdf.stat().st_mode
                    pdf.chmod(current | _stat.S_IWRITE)
                    self.log(f"  i {tab.capitalize()} re-opened — PDF unlocked: {pdf.name}")
            except Exception as e:
                self.log(f"  ! Could not change PDF permissions: {e}")
        else:
            status = "complete" if is_complete else "in progress"
            self.log(f"  i {tab.capitalize()} marked {status} (no PDF found to lock)")

        # Persist complete state in values JSON
        wd = self._resolve_working_dir()
        if wd and self.state.saisie_pdf_file:
            try:
                case_paths = ensure_case_structure(wd, self.state.saisie_pdf_file)
                update_values_latest_json(case_paths["values_latest_json"], {
                    f"complete_{tab}": is_complete,
                    "updated_at": _timestamp_iso(),
                })
            except Exception:
                pass

    def _on_tab_close(self, tab: str):
        """[X] button — uncheck Complete and clear the tab's status/output fields."""
        var_map = {
            "saisie":    getattr(self, "varCompleteSaisie",    None),
            "agenda":    getattr(self, "varCompleteAgenda",    None),
            "k138":      getattr(self, "varCompleteK138",      None),
            "narrative": getattr(self, "varCompleteNarrative", None),
        }
        var = var_map.get(tab)
        if var:
            if var.get():
                var.set(False)
                self._on_tab_complete_changed(tab)   # unlock PDF
        self.log(f"  i {tab.capitalize()} tab reset.")

    def _resolve_tab_pdf(self, tab: str) -> Optional[Path]:
        """Return the output PDF path for the given tab, or None if not available."""
        wd = self._resolve_working_dir()
        if not wd or not self.state.saisie_pdf_file:
            return None
        try:
            case_paths = ensure_case_structure(wd, self.state.saisie_pdf_file)
            if tab == "saisie":
                active = self._active_case_folder_path()
                if active:
                    pdfs = sorted(active.glob("*Saisie_D_affaire.pdf"), key=lambda p: p.stat().st_mtime, reverse=True)
                    return pdfs[0] if pdfs else None
            elif tab == "agenda":
                return self._agenda_existing_path(case_paths)
            elif tab == "k138":
                p = case_paths.get("client_k138_pdf")
                return p if p and p.exists() else None
        except Exception:
            pass
        return None

    def _on_labo_changed(self):
        """LABO checkbox toggled — paste 'LABO' into the notes field as AEADS (MTL) marker.
        NOTE: Confirm with Dmitry which exact Saisie D'affaire field maps to AESD (mtl) box."""
        if hasattr(self, "varSaisieAffaireFields") and "notes" in self.varSaisieAffaireFields:
            if self.varCheckLabo.get():
                current = self.varSaisieAffaireFields["notes"].get()
                if "LABO" not in current.upper():
                    self.varSaisieAffaireFields["notes"].set(("LABO  " + current).strip())
            else:
                current = self.varSaisieAffaireFields["notes"].get()
                self.varSaisieAffaireFields["notes"].set(
                    re.sub(r"LABO\s*", "", current, flags=re.IGNORECASE).strip()
                )

    def on_process_pdf(self):
        """Process individual SAISIE input file (PDF, DOCX, image)."""
        # Validate templates folder
        if not self.state.templates_folder:
            messagebox.showwarning("Missing Folder", "Select Configurations folder first.")
            return
        
        # Validate SAISIE input file
        if not self.state.saisie_pdf_file:
            messagebox.showwarning("Missing File", "Select SAISIE input file first.")
            return
        
        if not self.state.saisie_pdf_file.exists():
            messagebox.showerror("File Not Found", f"SAISIE file not found:\n{self.state.saisie_pdf_file}")
            return

        # Sanity check: warn if SAISIE PDF is under templates folder (single-folder setup)
        templates_dir = self.state.templates_folder.resolve()
        saisie_path = self.state.saisie_pdf_file.resolve()
        is_inside_templates = False
        try:
            saisie_path.relative_to(templates_dir)
            is_inside_templates = True
        except ValueError:
            is_inside_templates = False

        if is_inside_templates:
            self.log("  [WARN] SAISIE file is inside the configurations folder — outputs will be saved there too.")

        # working_dir = folder where input Saisie file lives (output goes here)
        if self.state.working_dir:
            working_dir = self.state.working_dir.resolve()
        else:
            working_dir, case_name = detect_working_directory(self.state.saisie_pdf_file)
            self.state.working_dir = working_dir
            self.state.case_folder_name = case_name
        case_paths = ensure_case_structure(working_dir, self.state.saisie_pdf_file)
        source_name_lower = (self.state.saisie_pdf_file.name or "").lower()
        use_saisie_affaire_manual = source_name_lower.startswith("saisie_d_affaire")
        cached_existing: Dict[str, object] = {}
        manual_fields: Dict[str, str] = {}
        manual_checks: Dict[str, bool] = {}
        if use_saisie_affaire_manual:
            try:
                cached_existing = read_values_latest_json(case_paths["values_latest_json"])
                manual_obj = cached_existing.get("saisie_affaire_manual")
                if isinstance(manual_obj, dict):
                    fields_obj = manual_obj.get("fields")
                    if isinstance(fields_obj, dict):
                        manual_fields = {
                            str(k): clean_value(str(v or ""))
                            for k, v in fields_obj.items()
                        }
                    checks_obj = manual_obj.get("checks")
                    if isinstance(checks_obj, dict):
                        manual_checks = {
                            str(k): bool(v)
                            for k, v in checks_obj.items()
                        }
            except Exception:
                manual_fields = {}
                manual_checks = {}

        def _apply_saisie_affaire_manual_overrides(values_dict: Dict[str, str]) -> None:
            if not (use_saisie_affaire_manual and manual_fields):
                return
            apply_saisie_affaire_manual_to_k138_values(values_dict, manual_fields)

        self._set_busy(True, "Extracting values from SAISIE...")
        try:
            self.log(f"Configurations folder: {self.state.templates_folder}")
            self.log(f"Input file: {self.state.saisie_pdf_file.name}")
            self.log(f"Active case folder: {working_dir.name}")
            self.log(f"Case storage root: {case_paths['case_root'].name}")
            try:
                original_source = self.state.saisie_pdf_file.resolve()
                case_source, copied = ensure_case_source_file(self.state.saisie_pdf_file, case_paths)
                self.state.saisie_pdf_file = case_source
                self._set_entry(self.entSaisieFile, case_source)
                if case_source.resolve() == original_source:
                    self.log("  ✓ Source file already in case folder")
                elif copied:
                    self.log("  ✓ Copied source file into case folder (original kept)")
                else:
                    self.log("  ! Could not copy source file into case folder; using original file")
                self.log(f"Case source file: {case_source.name}")
                if use_saisie_affaire_manual and not manual_fields:
                    manual_fields = extract_saisie_affaire_manual_fields_from_pdf(case_source)
            except Exception as cp_err:
                self.log(f"  ! Could not place source file into case folder: {cp_err}")
            self.log("Searching for templates...")
            is_docx = (self.state.saisie_pdf_file.suffix or "").lower() == ".docx"
            is_image = (self.state.saisie_pdf_file.suffix or "").lower() in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}
            saisie_template: Optional[Path] = None
            if not is_docx and not is_image:
                saisie_template = find_saisie_template(self.state.templates_folder)
                if not saisie_template:
                    messagebox.showerror(
                        "Missing SAISIE Template",
                        f"Could not find SAISIE template PDF in:\n{self.state.templates_folder}\n\n"
                        f"Make sure your Configurations folder contains the blank SAISIE PDF template "
                        f"(e.g. 'SAISIE À FAIRE_francompact 2025.pdf').\n\n"
                        f"Go to Configurations (top-right) and re-select the correct templates folder.",
                    )
                    return
                self.log(f"  [OK] Found SAISIE template: {saisie_template.name}")
            elif is_image:
                self.log("  [OK] Input is image - extracting via full-image OCR + label regex")
            else:
                self.log("  [OK] Input is Word (.docx) - extracting via label regex (no template diff)")

            self.log(f"Form type: {self.state.form_type}")

            # Load notice text from .txt file
            notice_text = load_notice_text(self.state.templates_folder, self.state.form_type)
            if notice_text:
                self.log("  [OK] Loaded notice text from file")
            else:
                self.log("  [WARN] No notice text file found (will use empty notice)")

            self.log("\n" + "=" * 60)
            self.log(f"Processing: {self.state.saisie_pdf_file.name}")
            self.log(f"Output folder: {case_paths['case_root'].name}")

            # Extract data from SAISIE (once; text-diff or OCR fallback)
            top, bottom, extract_mode = extract_field_values(saisie_template, self.state.saisie_pdf_file)
            self.log(f"  [OK] Extraction mode: {extract_mode}")

            # Build K138 values (with notice text)
            values = build_k138_values_from_saisie(top, self.state.saisie_pdf_file, self.state.form_type, notice_text)
            _apply_saisie_affaire_manual_overrides(values)
            if not clean_value(values.get("description_inventory", "")) and (self.state.saisie_pdf_file.suffix or "").lower() == ".pdf":
                if not (HAVE_PIL and HAVE_TESSERACT):
                    self.log("  ! Inventory not detected and OCR is unavailable (install/enable Tesseract OCR).")
            self.state.last_inventory_number = _normalize_inventory_number(values.get("description_inventory", ""))
            self.state.last_agent_id = self._normalize_agent_id(values.get("seizing_officer", ""))
            self.state.last_k138_values = dict(values)
            self.state.last_top = dict(top)
            self.state.last_bottom = list(bottom)
            self.state.extraction_ran = True  # Cache is now trustworthy for this file
            case_meta = build_internal_case_meta(top, values)
            case_paths = ensure_case_structure(working_dir, self.state.saisie_pdf_file, case_meta=case_meta)
            self.log(f"Internal case folder: {case_paths['internal_case_root'].name}")
            values_latest_payload = {
                "updated_at": _timestamp_iso(),
                "source_file": str(self.state.saisie_pdf_file),
                "working_directory": str(working_dir),
                "case_folder_name": self.state.case_folder_name or working_dir.name,
                "form_type": self.state.form_type,
                "inventory_number": self.state.last_inventory_number,
                "inventory_source": "saisie-extract",
                "agent_id": self.state.last_agent_id,
                "bond_room_ledger": clean_value(str(top.get("BOND ROOM LEDGER #", "") or "")),
                "extract_mode": extract_mode,
                "k138_values_base": values,
            }
            if use_saisie_affaire_manual and manual_fields:
                values_latest_payload["saisie_affaire_manual"] = {
                    "fields": manual_fields,
                    "checks": manual_checks,
                }
            update_values_latest_json(case_paths["values_latest_json"], values_latest_payload)
            self._refresh_agenda_status()
            validation_errors, validation_warnings = validate_k138_values(values)
            # If text-diff missed critical fields, retry using OCR extraction before blocking.
            if validation_errors and extract_mode == "text-diff" and (not is_docx) and (not is_image) and saisie_template and HAVE_PIL and HAVE_TESSERACT:
                self.log("  ! text-diff missed critical fields; retrying with OCR fallback...")
                try:
                    top_ocr, _bottom_ocr = extract_field_values_ocr(saisie_template, self.state.saisie_pdf_file)
                    page_text_ocr = get_ocr_page_text_pdf(self.state.saisie_pdf_file)
                    if page_text_ocr:
                        top_page = _build_top_from_text(page_text_ocr)
                        if _top_quality_score(top_page) >= _top_quality_score(top_ocr):
                            top_ocr = top_page
                    _refine_top_with_label_regex(top_ocr, self.state.saisie_pdf_file)
                    values_ocr = build_k138_values_from_saisie(top_ocr, self.state.saisie_pdf_file, self.state.form_type, notice_text)
                    _apply_saisie_affaire_manual_overrides(values_ocr)
                    errs_ocr, warns_ocr = validate_k138_values(values_ocr)
                    prev_crit = {e for e in validation_errors if e.startswith("Missing critical field:")}
                    new_crit = {e for e in errs_ocr if e.startswith("Missing critical field:")}
                    crit_improved = len(new_crit) < len(prev_crit)
                    if len(errs_ocr) < len(validation_errors) or crit_improved:
                        top = top_ocr
                        values = values_ocr
                        self.state.last_inventory_number = _normalize_inventory_number(values.get("description_inventory", ""))
                        self.state.last_agent_id = self._normalize_agent_id(values.get("seizing_officer", ""))
                        self.state.last_k138_values = dict(values)
                        self.state.last_top = dict(top)
                        self.state.last_bottom = list(_bottom_ocr)
                        case_meta = build_internal_case_meta(top, values)
                        case_paths = ensure_case_structure(working_dir, self.state.saisie_pdf_file, case_meta=case_meta)
                        self.log(f"Internal case folder: {case_paths['internal_case_root'].name}")
                        values_latest_payload = {
                            "updated_at": _timestamp_iso(),
                            "source_file": str(self.state.saisie_pdf_file),
                            "working_directory": str(working_dir),
                            "case_folder_name": self.state.case_folder_name or working_dir.name,
                            "form_type": self.state.form_type,
                            "inventory_number": self.state.last_inventory_number,
                            "inventory_source": "saisie-extract",
                            "agent_id": self.state.last_agent_id,
                            "bond_room_ledger": clean_value(str(top.get("BOND ROOM LEDGER #", "") or "")),
                            "extract_mode": "text-diff+ocr-retry",
                            "k138_values_base": values,
                        }
                        if use_saisie_affaire_manual and manual_fields:
                            values_latest_payload["saisie_affaire_manual"] = {
                                "fields": manual_fields,
                                "checks": manual_checks,
                            }
                        update_values_latest_json(case_paths["values_latest_json"], values_latest_payload)
                        self._refresh_agenda_status()
                        validation_errors = errs_ocr
                        validation_warnings = warns_ocr
                        extract_mode = "text-diff+ocr-retry"
                        self.log(f"  [OK] OCR retry improved validation (remaining errors: {len(validation_errors)})")
                except Exception as e:
                    self.log(f"  ! OCR retry failed: {e}")
            if validation_warnings:
                self.log("  ! Validation warnings:")
                for w in validation_warnings:
                    self.log(f"    - {w}")
            if validation_errors:
                self.log("  X Validation errors:")
                for e in validation_errors:
                    self.log(f"    - {e}")

            # Hidden folder: save all extracted values for analysis/aggregation under working dir.
            hidden_dir = get_hidden_data_dir(working_dir)
            ts = _timestamp_compact()
            stem = self.state.saisie_pdf_file.stem
            hidden_csv = hidden_dir / f"k138_values_{stem}_{ts}.csv"
            write_k138_values_csv(hidden_csv, values)
            hidden_saisie = hidden_dir / f"saisie_extract_{stem}_{ts}.csv"
            write_saisie_csv(hidden_saisie, top, bottom)
            latest_civ = build_latest_civ_values(top, values, self.state.saisie_pdf_file, extract_mode)
            latest_civ_csv = hidden_dir / "latest_civ.csv"
            write_latest_civ_csv(latest_civ_csv, latest_civ)
            all_values_csv = hidden_dir / "all_values.csv"
            append_all_values_csv(all_values_csv, top, values, self.state.saisie_pdf_file, extract_mode)
            review_queue_csv = hidden_dir / "review_queue.csv"
            append_review_queue_csv(review_queue_csv, latest_civ, validation_errors, validation_warnings)
            case_index_csv = hidden_dir / "case_index.csv"
            duplicate_cases_csv = hidden_dir / "duplicate_cases.csv"
            is_duplicate_case = update_case_tracking_csv(case_index_csv, duplicate_cases_csv, latest_civ)
            self.log(f"  [INFO] Saved to hidden folder: {hidden_dir.name}/")
            self.log(f"  [INFO] Updated: {latest_civ_csv.name}")
            self.log(f"  [INFO] Updated: {all_values_csv.name}")
            self.log(f"  [INFO] Updated: {review_queue_csv.name}")
            self.log(f"  [INFO] Updated: {case_index_csv.name}")
            if is_duplicate_case:
                self.log(f"  ! Duplicate case detected (case_key={latest_civ.get('case_key', '')})")

            update_values_latest_json(
                case_paths["values_latest_json"],
                {
                    "updated_at": _timestamp_iso(),
                    "source_file": str(self.state.saisie_pdf_file),
                    "working_directory": str(working_dir),
                    "case_folder_name": self.state.case_folder_name or working_dir.name,
                    "form_type": self.state.form_type,
                    "inventory_number": self.state.last_inventory_number,
                    "inventory_source": "saisie-extract",
                    "agent_id": self.state.last_agent_id,
                    "extract_mode": extract_mode,
                    "k138_values_base": values,
                },
            )
            self._refresh_agenda_status()
            self.log("\n" + "=" * 60)
            self.log("Processing complete.")
            self.log("Next step: Generate Agenda. After Agenda is saved, generate K138 from the K138 tab.")
            if validation_errors:
                messagebox.showwarning(
                    "Missing Fields Detected",
                    "SAISIE was processed, but some K138 fields could not be read:\n\n"
                    + "\n".join(validation_errors)
                    + "\n\nYou can still generate Agenda and K138 — review missing fields before finalising.",
                )

        except Exception as e:
            self.log(f"ERROR: {e}")
            import traceback
            self.log(traceback.format_exc())
            messagebox.showerror("Error", str(e))
        finally:
            self._set_busy(False)


class AppDnD(AppBase):
    def _init_dnd(self):
        if not HAVE_DND:
            return

        # ── Saisie file entry ────────────────────────────────────────
        try:
            self.entSaisieFile.drop_target_register(DND_FILES)
            self.entSaisieFile.dnd_bind('<<Drop>>',      self.on_drop_saisie_file)
            self.entSaisieFile.dnd_bind('<<DragEnter>>', self._on_drag_enter_saisie)
            self.entSaisieFile.dnd_bind('<<DragLeave>>', self._on_drag_leave_saisie)
        except Exception:
            pass

        # ── Working Folder (Select Folder) tab ──────────────────────
        try:
            self.tabSelectFolder.drop_target_register(DND_FILES)
            self.tabSelectFolder.dnd_bind('<<Drop>>',      self.on_drop_select_folder)
            self.tabSelectFolder.dnd_bind('<<DragEnter>>', self._on_drag_enter_select_folder)
            self.tabSelectFolder.dnd_bind('<<DragLeave>>', self._on_drag_leave_select_folder)
        except Exception:
            pass

        # ── Agenda tab ───────────────────────────────────────────────
        try:
            self.tabAgenda.drop_target_register(DND_FILES)
            self.tabAgenda.dnd_bind('<<Drop>>',      self.on_drop_agenda_folder)
            self.tabAgenda.dnd_bind('<<DragEnter>>', self._on_drag_enter_agenda)
            self.tabAgenda.dnd_bind('<<DragLeave>>', self._on_drag_leave_agenda)
        except Exception:
            pass

    # ── Working Folder drag-drop handlers ───────────────────────────
    def on_drop_select_folder(self, event):
        self._on_drag_leave_select_folder(None)
        try:
            items = self.root.tk.splitlist(event.data)
        except Exception:
            items = []
        if items:
            self.root.after(1, lambda: self._apply_dropped_select_folder(items[0]))

    def _apply_dropped_select_folder(self, raw_path: str):
        try:
            p = Path(raw_path)
            folder = p if p.is_dir() else p.parent
            if folder.exists():
                self._set_active_case_folder(folder)
        except Exception as e:
            self.log(f"Error handling Working Folder drag-and-drop: {e}")

    def _on_drag_enter_select_folder(self, event):
        try:
            if not hasattr(self, "_select_folder_drop_lbl"):
                self._select_folder_drop_lbl = tk.Label(
                    self.tabSelectFolder,
                    text="↓  Drop folder here to set as Active Case Folder  ↓",
                    font=("Segoe UI", 11),
                    fg="#1F4E79",
                    bg="#E3F0FB",
                    relief="solid",
                    bd=1,
                    pady=14,
                )
            self._select_folder_drop_lbl.place(relx=0, rely=0, relwidth=1, relheight=1)
            self._select_folder_drop_lbl.lift()
        except Exception:
            pass

    def _on_drag_leave_select_folder(self, *_):
        try:
            if hasattr(self, "_select_folder_drop_lbl"):
                self._select_folder_drop_lbl.place_forget()
        except Exception:
            pass

    # ── Drag hover helpers ───────────────────────────────────────────
    def _on_drag_enter_saisie(self, event):
        try:
            self.boxSaisieFile.configure(text="↓  Drop Saisie à Faire here  ↓")
            self.entSaisieFile.configure(style="DragOver.TEntry")
        except Exception:
            pass

    def _on_drag_leave_saisie(self, event):
        try:
            self.boxSaisieFile.configure(text="Select Saisie à Faire")
            self.entSaisieFile.configure(style="TEntry")
        except Exception:
            pass

    def _on_drag_enter_agenda(self, event):
        try:
            if not hasattr(self, "_agenda_drop_lbl"):
                self._agenda_drop_lbl = tk.Label(
                    self.tabAgenda,
                    text="↓  Drop file or folder here to set case directory  ↓",
                    font=("Segoe UI", 11),
                    fg="#1F4E79",
                    bg="#E3F0FB",
                    relief="solid",
                    bd=1,
                    pady=14,
                )
            self._agenda_drop_lbl.place(relx=0, rely=0, relwidth=1, relheight=1)
            self._agenda_drop_lbl.lift()
        except Exception:
            pass

    def _on_drag_leave_agenda(self, event):
        try:
            if hasattr(self, "_agenda_drop_lbl"):
                self._agenda_drop_lbl.place_forget()
        except Exception:
            pass

    def on_drop_agenda_folder(self, event):
        """Handle drag-and-drop on the Agenda tab — sets the active case folder."""
        self._on_drag_leave_agenda(None)  # hide overlay immediately
        try:
            items = self.root.tk.splitlist(event.data)
        except Exception:
            items = []
        if items:
            self.root.after(1, lambda: self._apply_dropped_agenda_folder(items[0]))

    def _apply_dropped_agenda_folder(self, raw_path: str):
        """Apply a dropped path on the Agenda tab (deferred to avoid Windows DnD message conflicts)."""
        try:
            p = Path(raw_path)
            # Folder dropped → use it directly as case folder
            if p.is_dir():
                self._set_active_case_folder(p)
                return
            # File dropped → use its parent directory as case folder
            if p.exists():
                parent = p.parent
                self._set_active_case_folder(parent)
                # If it looks like a Saisie file, also load it
                if p.suffix.lower() in ('.pdf', '.docx'):
                    self._reset_case_runtime_state(keep_case_folder=True)
                    self.state.saisie_pdf_file = p
                    self._set_entry(self.entSaisieFile, p)
                    self.log(f"Dropped: {p.name}")
        except Exception as e:
            self.log(f"Error handling Agenda drag-and-drop: {e}")
    
    def on_drop_saisie_file(self, event):
        """Handle drag-and-drop of SAISIE input file."""
        self._on_drag_leave_saisie(None)  # restore label immediately
        try:
            files = self.root.tk.splitlist(event.data)
        except Exception:
            files = []
        if files:
            self.root.after(1, lambda: self._apply_dropped_saisie_file(files[0]))

    def _apply_dropped_saisie_file(self, raw_path: str):
        """Apply a dropped SAISIE file path (deferred to avoid Windows DnD message conflicts)."""
        try:
            file_path = Path(raw_path)
            if file_path.exists() and file_path.suffix.lower() in ('.pdf', '.docx'):
                self._reset_case_runtime_state(keep_case_folder=True)
                self.state.saisie_pdf_file = file_path
                self._set_entry(self.entSaisieFile, file_path)
                if not _is_transient_upload_path(file_path.parent):
                    set_config_path("paths", "saisie_folder", file_path.parent)
                    self.last_saisie_folder = file_path.parent
                set_config_path("paths", "last_saisie_file", file_path)
                self._set_working_directory(file_path, refresh_status=False)
                self.log(f"Dropped: {file_path.name}")
                self._auto_extract_if_ready()
            else:
                self.log(f"Dropped file must be PDF or Word (.docx): {Path(raw_path).name}")
        except Exception as e:
            self.log(f"Error handling drag-and-drop: {e}")


def _show_profile_splash() -> tuple:
    """
    Show a clean profile selection splash screen.
    Returns (profile_role, profile_badge).
    BSO role shows a badge number entry; other roles need no extra input.
    Remembers last selection per machine.
    """
    ACCENT        = "#1F4E79"
    CARD_BG       = "#FFFFFF"
    SELECTED_BG   = "#E3F0FB"
    SELECTED_BORDER = "#1F4E79"
    NORMAL_BORDER = "#D0D7DE"
    # Fallback circle colours when image files are missing
    AVATAR_COLORS    = {"BSO": "#1F4E79", "Clerk": "#217346", "Supervisor": "#5B2C8D"}
    AVATAR_INITIALS  = {"BSO": "B",       "Clerk": "C",       "Supervisor": "S"}
    # Image file inside assets/
    AVATAR_IMAGES    = {
        "BSO":        "assets/avatar_bso.png",
        "Clerk":      "assets/avatar_clerk1.png",
        "Supervisor": "assets/avatar_supervisor1.png",
    }
    ROLE_DESC = {
        "BSO":        "Working Folder\nSaisie d'affaire\nAgenda",
        "Clerk":      "Working Folder\nAgenda  ·  K138\nSaisie d'intérêt",
        "Supervisor": "Full access\nAll tabs",
    }

    saved_role  = get_config_text("user", "profile_role", "BSO") or "BSO"
    saved_badge = get_config_text("user", "badge_number", "")    or ""

    result = {"role": saved_role, "badge": saved_badge, "confirmed": False}

    splash = tk.Tk()
    splash.title("Radiance Copilot — Select Profile")
    splash.configure(bg=ACCENT)
    splash.resizable(False, False)

    splash.update_idletasks()
    w, h = 580, 440
    sw, sh = splash.winfo_screenwidth(), splash.winfo_screenheight()
    splash.geometry(f"{w}x{h}+{(sw - w) // 2}+{(sh - h) // 2}")

    app_ico_path = resolve_asset_path("photos/Radiance-copilot-icon.ico", "Radiance-copilot-icon.ico")
    try:
        if app_ico_path:
            splash.iconbitmap(app_ico_path)
    except Exception:
        pass

    # Pre-load avatar images (kept in list to prevent GC)
    _avatar_tk_images: dict = {}
    AVATAR_SIZE = 80  # px — displayed size inside card
    for role, rel_path in AVATAR_IMAGES.items():
        img_path = resolve_asset_path(rel_path)
        if img_path and HAVE_PIL and Image and ImageTk:
            try:
                pil_img = Image.open(img_path).convert("RGBA")
                pil_img = pil_img.resize((AVATAR_SIZE, AVATAR_SIZE), Image.LANCZOS)
                _avatar_tk_images[role] = ImageTk.PhotoImage(pil_img)
            except Exception:
                pass

    # ── Header ──────────────────────────────────────────────────
    header = tk.Frame(splash, bg=ACCENT)
    header.pack(fill="x")
    tk.Label(header, text="Radiance Copilot",
             font=("Segoe UI Semibold", 16), fg="#FFFFFF", bg=ACCENT).pack(pady=(18, 2))
    tk.Label(header, text="Select your user profile to continue",
             font=("Segoe UI", 10), fg="#DCE8F4", bg=ACCENT).pack(pady=(0, 14))

    # ── Bottom area (badge entry) — created BEFORE cards so _select_role can reference it ──
    bottom = tk.Frame(splash, bg="#F0F4F8")
    bottom.pack(fill="x", pady=0)

    badge_var = tk.StringVar(value=saved_badge)
    badge_row = tk.Frame(bottom, bg="#F0F4F8")
    tk.Label(badge_row, text="Badge number:", font=("Segoe UI", 10), bg="#F0F4F8").grid(
        row=0, column=0, padx=(0, 8))
    badge_entry = ttk.Entry(badge_row, textvariable=badge_var, width=20)
    badge_entry.grid(row=0, column=1)

    if saved_role == "BSO":
        badge_row.pack(pady=(10, 4))

    # ── Role cards ───────────────────────────────────────────────
    cards_frame = tk.Frame(splash, bg="#F0F4F8")
    cards_frame.pack(fill="x", pady=0)

    card_frames: dict = {}

    def _set_bg_recursive(widget, bg):
        try:
            # Don't recolor Label widgets that hold images — they look odd with a tinted bg
            if not isinstance(widget, tk.Label) or not widget.cget("image"):
                widget.configure(bg=bg)
        except Exception:
            pass
        for child in widget.winfo_children():
            _set_bg_recursive(child, bg)

    def _refresh_cards(selected_role):
        for role, card in card_frames.items():
            is_sel = (role == selected_role)
            card.configure(
                bg=SELECTED_BG if is_sel else CARD_BG,
                highlightbackground=SELECTED_BORDER if is_sel else NORMAL_BORDER,
                highlightthickness=2 if is_sel else 1,
            )
            for child in card.winfo_children():
                _set_bg_recursive(child, SELECTED_BG if is_sel else CARD_BG)

    def _select_role(role):
        result["role"] = role
        _refresh_cards(role)
        if role == "BSO":
            badge_row.pack(pady=(10, 4))
            badge_entry.focus_set()
        else:
            badge_row.pack_forget()

    for i, role in enumerate(["BSO", "Clerk", "Supervisor"]):
        card = tk.Frame(
            cards_frame, bg=CARD_BG,
            highlightbackground=NORMAL_BORDER, highlightthickness=1, cursor="hand2",
        )
        card.grid(row=0, column=i, padx=12, pady=16, ipadx=12, ipady=10, sticky="nsew")
        cards_frame.columnconfigure(i, weight=1)
        card_frames[role] = card

        tk_img = _avatar_tk_images.get(role)
        if tk_img:
            # Real photo — use a Label with image
            img_lbl = tk.Label(card, image=tk_img, bg=CARD_BG, cursor="hand2")
            img_lbl.pack(pady=(12, 6))
        else:
            # Fallback: draw a coloured circle with an initial
            cv = tk.Canvas(card, width=AVATAR_SIZE, height=AVATAR_SIZE,
                           bg=CARD_BG, highlightthickness=0)
            cv.pack(pady=(12, 6))
            cv.create_oval(4, 4, AVATAR_SIZE - 4, AVATAR_SIZE - 4,
                           fill=AVATAR_COLORS[role], outline=AVATAR_COLORS[role])
            cv.create_text(AVATAR_SIZE // 2, AVATAR_SIZE // 2,
                           text=AVATAR_INITIALS[role],
                           font=("Segoe UI Semibold", 26), fill="#FFFFFF")

        tk.Label(card, text=role, font=("Segoe UI Semibold", 12), bg=CARD_BG).pack()
        tk.Label(card, text=ROLE_DESC[role], font=("Segoe UI", 8), fg="#555555",
                 bg=CARD_BG, justify="center").pack(pady=(4, 10))

        for widget in [card] + list(card.winfo_children()):
            try:
                widget.bind("<Button-1>", lambda e, r=role: _select_role(r))
            except Exception:
                pass

    _refresh_cards(result["role"])

    # ── Continue button ──────────────────────────────────────────
    btn_row = tk.Frame(splash, bg=ACCENT)
    btn_row.pack(fill="x", side="bottom")

    def _confirm():
        badge = re.sub(r"\D", "", badge_var.get().strip())
        result["badge"] = badge
        result["confirmed"] = True
        set_config_text("user", "profile_role", result["role"])
        set_config_text("user", "badge_number", badge)
        splash.destroy()

    tk.Button(
        btn_row, text="Continue  →",
        font=("Segoe UI Semibold", 11),
        bg="#2E86DE", fg="#FFFFFF",
        activebackground="#1A6BBD", activeforeground="#FFFFFF",
        relief="flat", bd=0, padx=24, pady=10, cursor="hand2",
        command=_confirm,
    ).pack(pady=12)

    # Closing the window (X button) = user cancelled — entire app should not open
    splash.protocol("WM_DELETE_WINDOW", splash.destroy)
    splash.bind("<Return>", lambda e: _confirm())
    if saved_role == "BSO":
        badge_entry.focus_set()
    splash.mainloop()

    if not result.get("confirmed"):
        return None, None  # Signal to main() that user cancelled
    return result.get("role", "BSO"), result.get("badge", "")


def main():
    lock = acquire_single_instance_lock()
    if lock is None:
        show_single_instance_warning()
        return

    try:
        restart = True
        while restart:
            restart = False
            profile_role, profile_badge = _show_profile_splash()
            if profile_role is None:
                break  # User closed splash — exit completely

            if TkinterDnD is not None and DND_FILES is not None:
                root = TkinterDnD.Tk()
                AppDnD(root, profile_role=profile_role, profile_badge=profile_badge)
            else:
                root = tk.Tk()
                AppBase(root, profile_role=profile_role, profile_badge=profile_badge)

            root.mainloop()

            # If "Change Profile" was clicked, root was tagged before destroy — loop.
            if getattr(root, "_change_profile_requested", False):
                restart = True
    finally:
        release_single_instance_lock(lock)


if __name__ == "__main__":
    main()

# Build EXE (Windows, example):
#   python -m PyInstaller --onefile --windowed --clean --icon photos/Radiance-copilot-icon.ico --add-data "photos;photos" --add-data "assets;assets" --add-data "fill_k138_notice.py;." --add-data "fill_saisie_interet.py;." --hidden-import fill_k138_notice --hidden-import fill_saisie_interet saisie_a_faire_extractor.py
#endregion
